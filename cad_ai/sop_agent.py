from __future__ import annotations

import json
import re
import urllib.error
import urllib.request
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document
from pydantic import BaseModel, Field

from .sop_visual_template import (
    IE_TIME_STUDY_FIELDS,
    SOP_FLOWCHART_SHAPE_POLICY,
    _build_sop_word_document,
    _render_center_flowchart_shape_image,
    _write_word_format_check_json,
    build_process_flow_page,
    build_work_instruction_page,
    classify_sop_flow_node_shape,
)


SOP_AGENT_VERSION = "0.1.0"
SOP_GENERATION_SEQUENCE = [
    "parse_requirement",
    "fill_word_tables",
    "build_structured_flowchart",
    "render_center_flowchart_png",
    "insert_png_into_process_flow_body_cell",
    "validate_docx",
]
SOP_STATUS_DRAFT = "demo_not_for_release"
SOP_RIGHT_SECTION_TITLES = ["作业标准", "设备/工具", "辅助材料", "注意事项", "变更内容", "物料表"]
SOP_BOTTOM_SECTION_TITLES = ["批准", "审核", "制作", "材料环保要求", "管制文件（印章处）", "图号"]


class SopBomItem(BaseModel):
    item_no: str = ""
    material_code: str = ""
    name: str
    specification: str = ""
    quantity: str = ""
    note: str = ""


class SopRoutingStep(BaseModel):
    name: str
    description: str = ""
    station: str = ""
    type: str = "process"
    machine_model: str = ""
    visual_type: str = "process"
    average_observed_time_s: float | None = None
    standard_time_s: float | None = None


class SopGenerateRequest(BaseModel):
    product_name: str
    part_no: str
    document_no: str
    drawing_no: str = ""
    station: str = ""
    requirement_text: str = ""
    bom_items: list[SopBomItem] = Field(default_factory=list)
    routing_steps: list[SopRoutingStep] = Field(default_factory=list)
    machine_hints: list[str] = Field(default_factory=list)
    run_id: str | None = None
    out_dir: Path = Path("outputs/sop_agent_api")
    use_model: bool = False
    strict_model: bool = False
    model_base_url: str = "http://127.0.0.1:8081/v1"
    model_name: str = "qwen3.6-35b"
    api_key: str = "local"
    model_timeout_sec: float = 240.0
    model_max_tokens: int = 1900


class SopArtifactPaths(BaseModel):
    document_docx: Path
    center_flowchart_png: Path
    manifest_json: Path
    format_check_json: Path
    parsed_sop_json: Path
    model_prompt_txt: Path | None = None
    model_raw_response_txt: Path | None = None
    model_response_json: Path | None = None


class SopGenerateResponse(BaseModel):
    status: Literal["demo_not_for_release", "failed"]
    run_id: str
    product_name: str
    part_no: str
    document_no: str
    station: str
    generation_sequence: list[str] = Field(default_factory=lambda: list(SOP_GENERATION_SEQUENCE))
    tables_filled_before_flowchart: bool = True
    center_flowchart_target: str = "process_flow_body_table_cell_0_0"
    shape_policy: str = SOP_FLOWCHART_SHAPE_POLICY
    model: dict[str, Any] = Field(default_factory=dict)
    validation: dict[str, Any] = Field(default_factory=dict)
    artifacts: SopArtifactPaths
    api_contract: dict[str, Any] = Field(default_factory=dict)
    warnings: list[str] = Field(default_factory=list)
    ai_boundary: str = (
        "demo_not_for_release; no real IE measured time, signoff, equipment status, "
        "EHS record, OEE, yield, trial result, or released SOP conclusion is fabricated"
    )


def generate_sop_package(request: SopGenerateRequest | dict[str, Any]) -> SopGenerateResponse:
    typed_request = request if isinstance(request, SopGenerateRequest) else SopGenerateRequest.model_validate(request)
    run_id = _safe_run_id(typed_request.run_id or f"{typed_request.product_name}-{datetime.now().strftime('%Y%m%d%H%M%S')}")
    run_dir = typed_request.out_dir / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    warnings: list[str] = []
    model_artifacts: dict[str, Path] = {}
    model_info: dict[str, Any] = {
        "enabled": typed_request.use_model,
        "base_url": typed_request.model_base_url,
        "model": typed_request.model_name,
        "strict_model": typed_request.strict_model,
        "content_source": "structured_request",
    }

    if typed_request.use_model:
        try:
            model_data, artifacts = _generate_content_with_model(typed_request, run_dir)
            model_artifacts.update(artifacts)
            model_info["content_source"] = "local_35b_openai_compatible"
        except Exception as exc:
            if typed_request.strict_model:
                raise
            warnings.append(f"model_generation_failed_fallback_to_structured:{exc}")
            model_info["content_source"] = "structured_fallback_after_model_error"
            model_data = _build_structured_sop_data(typed_request)
    else:
        model_data = _build_structured_sop_data(typed_request)

    sop_data = _expand_sop_data(typed_request, model_data)
    parsed_sop_json = run_dir / "parsed_sop.json"
    parsed_sop_json.write_text(json.dumps(sop_data, ensure_ascii=False, indent=2), encoding="utf-8")

    flow_page, work_page, normalization = _build_pages(typed_request, sop_data)
    center_flowchart_png = run_dir / "center_flowchart.png"
    center_flowchart_png.write_bytes(_render_center_flowchart_shape_image(flow_page))

    document_docx = run_dir / "sop.docx"
    document = _build_sop_word_document(flow_page, work_page)
    document.save(document_docx)

    format_check_json = run_dir / "format_check.json"
    _write_word_format_check_json(format_check_json, flow_page, work_page)
    validation = _validate_docx_package(document_docx)

    manifest_json = run_dir / "manifest.json"
    artifact_paths = SopArtifactPaths(
        document_docx=document_docx.resolve(),
        center_flowchart_png=center_flowchart_png.resolve(),
        manifest_json=manifest_json.resolve(),
        format_check_json=format_check_json.resolve(),
        parsed_sop_json=parsed_sop_json.resolve(),
        model_prompt_txt=(model_artifacts.get("prompt") or None),
        model_raw_response_txt=(model_artifacts.get("raw_response") or None),
        model_response_json=(model_artifacts.get("response") or None),
    )
    response = SopGenerateResponse(
        status=SOP_STATUS_DRAFT,
        run_id=run_id,
        product_name=typed_request.product_name,
        part_no=typed_request.part_no,
        document_no=typed_request.document_no,
        station=typed_request.station,
        model=model_info | {"normalization": normalization},
        validation=validation,
        artifacts=artifact_paths,
        warnings=warnings,
        api_contract={
            "interface": "FastAPI",
            "server_factory": "cad_ai.sop_api:create_sop_fastapi_app",
            "routes": [
                "POST /api/sop/generate",
                "GET /api/sop/runs/{run_id}",
                "GET /api/sop/runs/{run_id}/artifacts/{artifact_key}",
                "GET /api/sop/health",
            ],
        },
    )
    manifest_json.write_text(json.dumps(response.model_dump(mode="json"), ensure_ascii=False, indent=2), encoding="utf-8")
    return response


def _generate_content_with_model(request: SopGenerateRequest, run_dir: Path) -> tuple[dict[str, Any], dict[str, Path]]:
    prompt = _build_model_prompt(request)
    prompt_path = run_dir / "model_prompt.txt"
    prompt_path.write_text(prompt, encoding="utf-8")
    payload = {
        "model": request.model_name,
        "messages": [
            {
                "role": "system",
                "content": (
                    "你是制造业SOP结构化数据生成器。只输出JSON对象。"
                    "所有发布、IE实测、审批和现场事实都必须保持草案边界。"
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
        "top_p": 0.85,
        "max_tokens": request.model_max_tokens,
        "stream": False,
        "response_format": {"type": "json_object"},
    }
    result = _post_openai_json(
        f"{request.model_base_url.rstrip('/')}/chat/completions",
        payload,
        api_key=request.api_key,
        timeout=request.model_timeout_sec,
    )
    response_path = run_dir / "model_response.json"
    response_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    choice = (result.get("choices") or [{}])[0]
    raw_text = (choice.get("message") or {}).get("content") or choice.get("text") or ""
    raw_path = run_dir / "model_raw_response.txt"
    raw_path.write_text(raw_text, encoding="utf-8")
    parsed = _extract_json_object(raw_text)
    return parsed, {"prompt": prompt_path.resolve(), "response": response_path.resolve(), "raw_response": raw_path.resolve()}


def _build_model_prompt(request: SopGenerateRequest) -> str:
    bom_text = "；".join(f"{item.name}/{item.specification}/{item.quantity}" for item in request.bom_items[:16]) or "未提供BOM明细"
    routing_text = "；".join(step.name for step in request.routing_steps[:12]) or "请从需求摘要拆解"
    machines = "；".join(request.machine_hints[:12]) or "未提供设备型号，需标待确认"
    return f"""只输出一个紧凑JSON对象，不要Markdown。你负责生成SOP内容，Word渲染器只负责排版。
产品:{request.product_name}; 料号:{request.part_no}; 文件:{request.document_no}; 图号:{request.drawing_no}; 工站:{request.station}
需求:{request.requirement_text}
BOM线索:{bom_text}
工艺路线线索:{routing_text}
设备/治具线索:{machines}
输出结构:
{{
 "status":"demo_not_for_release",
 "order":["6个主要动作名"],
 "steps":[["动作名","作业描述<=18字","visual_type","machine_model","avg_s","std_s"]],
 "sections":[["作业标准","短句1","短句2"],["设备/工具","短句1","短句2"],["辅助材料","短句1","短句2"],["注意事项","短句1","短句2"],["变更内容","短句1","短句2"],["物料表","短句1","短句2"]],
 "nodes":[["节点名","type","shape"]],
 "edges":[["OP01","OP02","next"]],
 "notes":["草案说明"]
}}
硬规则:
- status="demo_not_for_release"。
- steps正好6行；avg_s/std_s为草案估算秒数，std_s约等于avg_s*1.10，不是现场实测。
- sections按示例6个标题顺序输出，每项最多2个短句。
- nodes 8到10个；覆盖来料核对、加工/装配、检查/测试、合格判定、最终清洁/贴标/包装入库；不能以合格判定结束；节点名不要包含OP编号。
- 检查/检验/测量/AOI/ICT/FCT/EOL/合格判定=diamond; 装配/点胶/固化/清洁/包装=ellipse。
- 不写真实工厂、人员签核、EHS审批、良率、OEE、试产或发布结论。
"""


def _build_structured_sop_data(request: SopGenerateRequest) -> dict[str, Any]:
    steps = request.routing_steps or _infer_steps_from_text(request.requirement_text)
    steps = _pad_steps(steps)
    compact_steps = []
    for index, step in enumerate(steps[:6], start=1):
        average = step.average_observed_time_s if step.average_observed_time_s is not None else 8.0 + index * 2
        standard = step.standard_time_s if step.standard_time_s is not None else average * 1.10
        compact_steps.append(
            [
                step.name,
                step.description or f"按{step.name}要求执行并记录异常",
                step.visual_type or _visual_type_for_step(step),
                step.machine_model or _machine_hint(request, index),
                f"{average:.1f}",
                f"{standard:.1f}",
            ]
        )
    nodes = _build_compact_nodes(request, steps)
    return {
        "status": SOP_STATUS_DRAFT,
        "order": [step.name for step in steps[:6]],
        "steps": compact_steps,
        "sections": _default_sections(request),
        "nodes": nodes,
        "edges": [[f"OP{i:02d}", f"OP{i + 1:02d}", "next"] for i in range(1, len(nodes))],
        "notes": ["结构化需求生成；demo_not_for_release"],
    }


def _expand_sop_data(request: SopGenerateRequest, data: dict[str, Any]) -> dict[str, Any]:
    if data.get("metadata") and data.get("step_slots") and data.get("flowchart"):
        return data
    steps = list(data.get("steps") or [])
    nodes = list(data.get("nodes") or [])
    sections = list(data.get("sections") or [])
    order = list(data.get("order") or [step[0] for step in steps if isinstance(step, list) and step])
    step_slots: list[dict[str, Any]] = []
    ie_rows: list[dict[str, Any]] = []
    for index, raw_step in enumerate(steps[:6], start=1):
        title, desc, visual_type, machine, average, standard = _compact_step_fields(raw_step, index)
        step_slots.append({"slot_no": index, "title": title, "description": desc, "visual_type": visual_type})
        ie_rows.append(
            {
                "action": title,
                "machine_model": machine,
                "measurement_method": "草案估算，待现场IE实测",
                "observations": "3",
                "average_observed_time_s": average,
                "rating_factor": "1.00",
                "allowance_rate": "10%",
                "standard_time_s": standard,
                "time_source": "draft_estimate_not_for_release",
                "dynamic_adjustment": "MES实绩/IE复测/机器型号变更后更新",
            }
        )
    while len(step_slots) < 6:
        index = len(step_slots) + 1
        title = f"待补充动作{index}"
        step_slots.append({"slot_no": index, "title": title, "description": "需求信息不足，待人工补充", "visual_type": "process"})
        ie_rows.append(
            {
                "action": title,
                "machine_model": "待确认",
                "measurement_method": "待现场IE实测",
                "observations": "",
                "average_observed_time_s": "",
                "rating_factor": "",
                "allowance_rate": "",
                "standard_time_s": "",
                "time_source": "missing_site_measurement_not_for_release",
                "dynamic_adjustment": "补齐现场数据后更新",
            }
        )
    flow_nodes = _expand_compact_nodes(nodes, request.station)
    return {
        "status": data.get("status") or SOP_STATUS_DRAFT,
        "metadata": {
            "product_name": request.product_name,
            "part_no": request.part_no,
            "document_no": request.document_no,
            "drawing_no": request.drawing_no,
            "station": request.station,
            "version": "DRAFT",
        },
        "operation_order": " -> ".join(str(item) for item in order[:6]) or " -> ".join(slot["title"] for slot in step_slots),
        "step_slots": step_slots,
        "side_sections": _expand_compact_sections(sections),
        "bom_items": [item.model_dump(mode="json") for item in request.bom_items],
        "ie_time_study_rows": ie_rows,
        "flowchart": {
            "flowchart_title": f"{request.product_name} 工艺流程",
            "nodes": flow_nodes,
            "edges": _expand_compact_edges(data.get("edges") or [], flow_nodes),
        },
        "notes": data.get("notes") or ["demo_not_for_release"],
    }


def _build_pages(request: SopGenerateRequest, data: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    metadata = data["metadata"]
    chart = data["flowchart"]
    nodes, shape_normalization = _normalize_flow_nodes(chart.get("nodes") or [], request.station)
    operation_names = [str(node.get("name") or f"工序{index}") for index, node in enumerate(nodes, start=1)]
    flow_page = build_process_flow_page(
        product_name=metadata["product_name"],
        part_no=metadata["part_no"],
        document_no=metadata["document_no"],
        drawing_no=metadata["drawing_no"],
        operations=operation_names,
    )
    flow_page["render_center_flowchart"] = True
    flow_page["center_flowchart_style"] = "pdf_reference_shape_blocks"
    flow_page["center_flowchart"] = {
        "flowchart_title": chart.get("flowchart_title") or f"{metadata['product_name']} 工艺流程",
        "nodes": nodes,
        "edges": chart.get("edges") or _linear_edges(nodes),
    }
    flow_page["ie_time_study"] = _process_ie_time_study(data, nodes)

    work_page = build_work_instruction_page(
        product_name=metadata["product_name"],
        part_no=metadata["part_no"],
        station=metadata["station"],
        document_no=metadata["document_no"],
        drawing_no=metadata["drawing_no"],
        version=metadata.get("version") or "DRAFT",
        page_no=1,
        page_total=1,
    )
    work_page["operation_order"] = data.get("operation_order") or " -> ".join(operation_names[:6])
    work_page["step_slots"] = _normalize_step_slots(data.get("step_slots") or [])
    work_page["side_sections"] = _normalize_side_sections(data.get("side_sections") or [])
    work_page["ie_time_study"] = _work_ie_time_study(data)
    work_page["bottom_sections"] = [
        {"title": "批准", "value": ""},
        {"title": "审核", "value": ""},
        {"title": "制作", "value": ""},
        {"title": "材料环保要求", "value": "所有材料需符合RoHS/REACH要求；正式发布前由体系/工程确认。"},
        {"title": "管制文件（印章处）", "value": ""},
        {"title": "图号", "value": metadata["drawing_no"]},
    ]
    return flow_page, work_page, {
        "shape_policy": SOP_FLOWCHART_SHAPE_POLICY,
        "shape_normalization": shape_normalization,
        "side_sections_normalized_to_reference_order": True,
        "bottom_signoff_cells_blank": True,
        "ie_standard_time_formula_enforced": "missing or inconsistent draft standard time is rendered as avg_s x 1.10",
    }


def _normalize_flow_nodes(raw_nodes: list[dict[str, Any]], station: str) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    nodes: list[dict[str, Any]] = []
    normalization: list[dict[str, Any]] = []
    for index, raw_node in enumerate(raw_nodes, start=1):
        node = dict(raw_node)
        node["seq"] = int(node.get("seq") or index)
        node["id"] = str(node.get("id") or f"OP{index:02d}")
        node["label"] = str(node.get("label") or f"工序{index}")
        node["name"] = _clean_operation_name(str(node.get("name") or node.get("title") or f"工序{index}"))
        node["type"] = str(node.get("type") or "process")
        node["station"] = str(node.get("station") or station)
        node["note"] = str(node.get("note") or SOP_STATUS_DRAFT)
        model_shape = str(node.get("shape") or "").lower()
        rendered_shape = classify_sop_flow_node_shape(node)
        node["shape"] = rendered_shape
        normalization.append(
            {
                "id": node["id"],
                "name": node["name"],
                "model_shape": model_shape,
                "rendered_shape": rendered_shape,
                "changed": model_shape != rendered_shape,
            }
        )
        nodes.append(node)
    return nodes, normalization


def _process_ie_time_study(data: dict[str, Any], nodes: list[dict[str, Any]]) -> dict[str, Any]:
    rows = data.get("ie_time_study_rows") or []
    machine_summary = "; ".join(str(row.get("machine_model") or "") for row in rows[:4])
    row = {
        "动作": "流程工时汇总",
        "机器型号": machine_summary + ("; ..." if len(rows) > 4 else ""),
        "IE测量方法": "草案估算；正式值需现场秒表法/MES/视频分析确认",
        "观测次数": "",
        "平均观测工时(s)": "",
        "评比系数": "",
        "宽放率": "",
        "标准工时(s)": "详见标准作业指导书IE表",
        "工时来源": SOP_STATUS_DRAFT,
        "动态调整": "随MES实绩、IE复测、机器型号或工艺变更更新",
    }
    return {"title": "IE工时记录", "fields": list(IE_TIME_STUDY_FIELDS), "rows": [row], "policy": {"operation_count": len(nodes)}}


def _work_ie_time_study(data: dict[str, Any]) -> dict[str, Any]:
    rows = [_map_ie_row(row) for row in (data.get("ie_time_study_rows") or [])]
    return {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": rows[:8],
        "policy": {"release_requirement": "site_measured_or_human_locked"},
    }


def _map_ie_row(row: dict[str, Any]) -> dict[str, str]:
    average = str(row.get("average_observed_time_s") or "")
    return {
        "动作": str(row.get("action") or ""),
        "机器型号": str(row.get("machine_model") or "待确认"),
        "IE测量方法": str(row.get("measurement_method") or "草案估算，待现场IE实测"),
        "观测次数": str(row.get("observations") or "3"),
        "平均观测工时(s)": average,
        "评比系数": str(row.get("rating_factor") or "1.00"),
        "宽放率": str(row.get("allowance_rate") or "10%"),
        "标准工时(s)": _normalize_standard_time_value(average, str(row.get("standard_time_s") or "")),
        "工时来源": str(row.get("time_source") or "draft_estimate_not_for_release"),
        "动态调整": str(row.get("dynamic_adjustment") or "MES实绩/IE复测后更新"),
    }


def _normalize_step_slots(raw_slots: list[dict[str, Any]]) -> list[dict[str, Any]]:
    slots: list[dict[str, Any]] = []
    for index in range(1, 7):
        raw = next((slot for slot in raw_slots if int(slot.get("slot_no") or 0) == index), {})
        title = str(raw.get("title") or f"步骤{index}")
        description = str(raw.get("description") or "待补充动作描述")
        visual_type = str(raw.get("visual_type") or "process")
        slots.append(
            {
                "slot_no": index,
                "image_placeholder": False,
                "text_placeholder": f"{index}. {title}: {description}",
                "visual": {"type": visual_type, "title": title},
            }
        )
    return slots


def _normalize_side_sections(raw_sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_title = {str(section.get("title") or ""): section for section in raw_sections}
    return [{"title": title, "lines": [str(line) for line in (by_title.get(title, {}).get("lines") or [])]} for title in SOP_RIGHT_SECTION_TITLES]


def _expand_compact_sections(raw_sections: list[Any]) -> list[dict[str, Any]]:
    section_map: dict[str, list[str]] = {}
    for raw in raw_sections:
        if isinstance(raw, dict):
            title = str(raw.get("title") or "")
            lines = [str(line) for line in (raw.get("lines") or [])]
        elif isinstance(raw, list) and raw:
            title = str(raw[0])
            lines = [str(item) for item in raw[1:] if str(item)]
        else:
            continue
        section_map[title] = lines[:2]
    return [{"title": title, "lines": section_map.get(title, [])} for title in SOP_RIGHT_SECTION_TITLES]


def _expand_compact_nodes(raw_nodes: list[Any], station: str) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    for index, raw in enumerate(raw_nodes[:10], start=1):
        if isinstance(raw, dict):
            name = _clean_operation_name(str(raw.get("name") or raw.get("title") or f"工序{index}"))
            node_type = str(raw.get("type") or "process")
            shape = str(raw.get("shape") or "")
        else:
            values = list(raw) if isinstance(raw, list) else []
            values += [""] * (3 - len(values))
            name = _clean_operation_name(str(values[0] or f"工序{index}"))
            node_type = str(values[1] or "process")
            shape = str(values[2] or "")
        nodes.append(
            {
                "seq": index,
                "id": f"OP{index:02d}",
                "label": f"工序{index}",
                "name": name,
                "type": node_type,
                "shape": shape,
                "station": station,
                "note": SOP_STATUS_DRAFT,
            }
        )
    return nodes


def _expand_compact_edges(raw_edges: list[Any], nodes: list[dict[str, Any]]) -> list[dict[str, str]]:
    edges: list[dict[str, str]] = []
    node_ids = {str(node["id"]) for node in nodes}
    for raw in raw_edges:
        values = raw if isinstance(raw, list) else [raw.get("from"), raw.get("to"), raw.get("label", "next")] if isinstance(raw, dict) else []
        values += [""] * (3 - len(values))
        edge_from, edge_to, label = str(values[0] or ""), str(values[1] or ""), str(values[2] or "next")
        if edge_from in node_ids and edge_to in node_ids:
            edges.append({"from": edge_from, "to": edge_to, "label": label})
    return edges or _linear_edges(nodes)


def _linear_edges(nodes: list[dict[str, Any]]) -> list[dict[str, str]]:
    return [{"from": str(nodes[index]["id"]), "to": str(nodes[index + 1]["id"]), "label": "next"} for index in range(len(nodes) - 1)]


def _compact_step_fields(raw_step: Any, index: int) -> tuple[str, str, str, str, str, str]:
    if isinstance(raw_step, dict):
        title = str(raw_step.get("title") or raw_step.get("action") or f"动作{index}")
        desc = str(raw_step.get("description") or raw_step.get("desc") or "按SOP执行")
        visual_type = str(raw_step.get("visual_type") or raw_step.get("type") or "process")
        machine = str(raw_step.get("machine_model") or raw_step.get("machine") or "待确认")
        average = str(raw_step.get("avg_s") or raw_step.get("average_observed_time_s") or "")
        standard = _normalize_standard_time_value(average, str(raw_step.get("std_s") or raw_step.get("standard_time_s") or ""))
        return title, desc, visual_type, machine, average, standard
    values = list(raw_step) if isinstance(raw_step, list) else []
    values += [""] * (6 - len(values))
    title = str(values[0] or f"动作{index}")
    desc = str(values[1] or "按SOP执行")
    visual_type = str(values[2] or "process")
    machine = str(values[3] or "待确认")
    average = str(values[4] or "")
    standard = _normalize_standard_time_value(average, str(values[5] or ""))
    return title, desc, visual_type, machine, average, standard


def _normalize_standard_time_value(average: str, standard: str) -> str:
    avg = _parse_float(average)
    std = _parse_float(standard)
    if avg is None:
        return standard
    if std is None or std <= avg or std > avg * 1.5:
        return f"{avg * 1.10:.1f}"
    return standard


def _parse_float(value: str) -> float | None:
    try:
        return float(str(value).strip().replace("%", ""))
    except ValueError:
        return None


def _infer_steps_from_text(text: str) -> list[SopRoutingStep]:
    candidates = [item.strip(" ，,;；。") for item in re.split(r"->|→|,|，|;|；|、|\n", text) if item.strip()]
    names = [item for item in candidates if 2 <= len(item) <= 18][:8]
    if not names:
        names = ["来料核对", "装配作业", "过程检查", "功能测试", "合格判定", "包装入库"]
    return [SopRoutingStep(name=name, type=_node_type_for_name(name), visual_type=_visual_type_for_name(name)) for name in names]


def _pad_steps(steps: list[SopRoutingStep]) -> list[SopRoutingStep]:
    padded = list(steps)
    defaults = ["来料核对", "装配作业", "过程检查", "功能测试", "合格判定", "包装入库"]
    existing = {step.name for step in padded}
    for name in defaults:
        if len(padded) >= 6:
            break
        if name not in existing:
            padded.append(SopRoutingStep(name=name, type=_node_type_for_name(name), visual_type=_visual_type_for_name(name)))
    return padded[:6]


def _build_compact_nodes(request: SopGenerateRequest, steps: list[SopRoutingStep]) -> list[list[str]]:
    node_names = ["来料核对"] + [step.name for step in steps[:6]]
    if not any("合格" in name or "判定" in name for name in node_names):
        node_names.append("合格判定")
    if not any("包装" in name or "入库" in name for name in node_names):
        node_names.append("包装入库")
    compact = []
    seen: set[str] = set()
    for name in node_names[:10]:
        clean = _clean_operation_name(name)
        if clean in seen:
            continue
        seen.add(clean)
        node_type = _node_type_for_name(clean)
        shape = "diamond" if node_type in {"inspection", "test", "measurement", "quality", "decision"} else "ellipse"
        compact.append([clean, node_type, shape])
    return compact


def _default_sections(request: SopGenerateRequest) -> list[list[str]]:
    bom_line = "；".join(item.name for item in request.bom_items[:4]) or "BOM待确认"
    machine_line = "；".join(request.machine_hints[:4]) or "设备/治具待确认"
    return [
        ["作业标准", "按BOM与工艺路线执行", "异常品隔离并记录"],
        ["设备/工具", machine_line, "正式型号待现场确认"],
        ["辅助材料", bom_line, "标签/包装按工单确认"],
        ["注意事项", "草案工时不可发布", "签核与EHS需人工确认"],
        ["变更内容", "DRAFT/接口生成", "变更需评审"],
        ["物料表", bom_line, "用量以受控BOM为准"],
    ]


def _machine_hint(request: SopGenerateRequest, index: int) -> str:
    if index - 1 < len(request.machine_hints):
        return request.machine_hints[index - 1]
    return "待确认"


def _visual_type_for_step(step: SopRoutingStep) -> str:
    return step.visual_type or _visual_type_for_name(step.name)


def _visual_type_for_name(name: str) -> str:
    node_type = _node_type_for_name(name)
    if node_type in {"inspection", "measurement"}:
        return "inspection"
    if node_type in {"test", "quality", "decision"}:
        return "test"
    if "清洁" in name:
        return "clean"
    if "包" in name or "入库" in name:
        return "pack"
    if "来料" in name or "备料" in name:
        return "materials"
    return "process"


def _node_type_for_name(name: str) -> str:
    upper = name.upper()
    if any(token in upper for token in ["AOI", "ICT", "FCT", "EOL", "TEST"]):
        return "test"
    if any(token in name for token in ["测试", "电测", "功能"]):
        return "test"
    if any(token in name for token in ["检查", "检验", "核对", "外观"]):
        return "inspection"
    if any(token in name for token in ["测量", "尺寸"]):
        return "measurement"
    if any(token in name for token in ["合格", "判定"]):
        return "decision"
    return "process"


def _clean_operation_name(name: str) -> str:
    cleaned = re.sub(r"^OP\s*\d+\s*[_:：.-]?\s*", "", name, flags=re.IGNORECASE).strip()
    return cleaned or name


def _safe_run_id(value: str) -> str:
    slug = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_.-]+", "_", value).strip("._")
    return slug[:80] or f"sop-{datetime.now().strftime('%Y%m%d%H%M%S')}"


def _extract_json_object(text: str) -> dict[str, Any]:
    stripped = text.strip()
    stripped = re.sub(r"^```(?:json)?", "", stripped, flags=re.IGNORECASE).strip()
    stripped = re.sub(r"```$", "", stripped).strip()
    start = stripped.find("{")
    end = stripped.rfind("}")
    if start < 0 or end < start:
        raise ValueError("no JSON object boundaries")
    return json.loads(stripped[start : end + 1])


def _post_openai_json(url: str, payload: dict[str, Any], *, api_key: str, timeout: float) -> dict[str, Any]:
    native_url = _ollama_native_chat_url(url)
    if native_url:
        return _post_ollama_native_json(native_url, _ollama_native_chat_payload(payload), timeout=timeout)
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as exc:
        raise RuntimeError(f"model endpoint unavailable: {exc}") from exc


def _ollama_native_chat_url(openai_chat_url: str) -> str | None:
    normalized = openai_chat_url.rstrip("/")
    if not normalized.startswith("http://127.0.0.1:11434/") and not normalized.startswith("http://localhost:11434/"):
        return None
    if normalized.endswith("/v1/chat/completions"):
        return normalized[: -len("/v1/chat/completions")] + "/api/chat"
    return None


def _ollama_native_chat_payload(openai_payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "model": openai_payload.get("model"),
        "messages": openai_payload.get("messages") or [],
        "stream": False,
        "think": False,
        "format": "json",
        "options": {
            "temperature": openai_payload.get("temperature", 0.1),
            "top_p": openai_payload.get("top_p", 0.85),
            "num_predict": openai_payload.get("max_tokens", 1900),
        },
    }


def _post_ollama_native_json(url: str, payload: dict[str, Any], *, timeout: float) -> dict[str, Any]:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as exc:
        raise RuntimeError(f"model endpoint unavailable: {exc}") from exc
    message = data.get("message") if isinstance(data, dict) else {}
    content = (message or {}).get("content") or ""
    return {
        "id": data.get("id") or "ollama-native-chat",
        "object": "chat.completion",
        "model": data.get("model") or payload.get("model"),
        "choices": [{"index": 0, "message": {"role": "assistant", "content": content}, "finish_reason": "stop" if data.get("done") else ""}],
        "usage": {
            "prompt_tokens": data.get("prompt_eval_count", 0),
            "completion_tokens": data.get("eval_count", 0),
            "total_tokens": int(data.get("prompt_eval_count", 0) or 0) + int(data.get("eval_count", 0) or 0),
        },
        "ollama_native": True,
    }


def _validate_docx_package(path: Path) -> dict[str, Any]:
    document = Document(path)
    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
        document_xml = package.read("word/document.xml").decode("utf-8")
    all_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return {
        "sections": len(document.sections),
        "top_level_tables": len(document.tables),
        "has_png_media": any(name.startswith("word/media/") and name.endswith(".png") for name in names),
        "has_picture_reference": "pic:pic" in document_xml,
        "has_svg": "<svg" in document_xml.lower(),
        "has_vml_shape": "<v:shape" in document_xml.lower(),
        "contains_replacement_char": "\ufffd" in all_text,
        "contains_ie_time_title": "IE工时记录" in all_text,
        "contains_machine_model_field": "机器型号" in all_text,
        "contains_demo_boundary": SOP_STATUS_DRAFT in all_text,
    }
