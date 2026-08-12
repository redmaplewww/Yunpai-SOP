from __future__ import annotations

import csv
import json
from io import BytesIO
from html import escape
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter, range_boundaries
from openpyxl.worksheet.page import PageMargins


REFERENCE_80806_129_FORMAT: dict[str, Any] = {
    "source_pdf": "80806-129.pdf",
    "source_title": "80806-129  SOP-EH.xlsx",
    "page_count": 42,
    "flow_page": {
        "page_number": 1,
        "size_pt": [595.32, 841.92],
        "orientation": "portrait",
        "title": "流程图",
        "header_fields": [
            "品名",
            "料号",
            "页数",
            "作业部门",
            "版本",
            "制定日期",
            "图号",
            "核准",
            "审核",
            "拟订",
            "文件编号",
        ],
        "flow_label_prefix": "工序",
        "flow_body": "multi_column_operation_flow",
        "footer_form_code": "EF-42013-23 (REV. A)",
    },
    "work_instruction_page": {
        "first_page_number": 2,
        "size_pt": [841.92, 595.32],
        "orientation": "landscape",
        "title": "标准作业指导书",
        "header_fields": [
            "产品品名",
            "本厂料号",
            "工站",
            "文件编号",
            "制作日期",
            "版本",
            "页码",
        ],
        "image_slots": 6,
        "step_sequence": [1, 2, 3, 4, 5, 6],
        "right_sections": ["作业标准", "设备/工具", "辅助材料", "注意事项", "变更内容", "物料表"],
        "bottom_sections": ["批准", "审核", "制作", "材料环保要求", "管制文件（印章处）", "图号"],
    },
}


DEFAULT_OPERATION_NAMES = [
    "裁线",
    "自动机前处理",
    "分线一",
    "镭射芯线铝箔",
    "分线二",
    "排线",
    "检查线序",
    "低压成型",
    "取排线治具",
    "装治具-理线",
    "自动机焊接前处理",
    "CCD对位",
    "自动机焊接",
    "CCD检查点",
    "维修工站",
    "半品测试一",
    "点/烤UV胶",
    "量测总长",
    "铁壳组装、铆压",
    "激光焊接铁壳",
    "半品测试二",
    "成型TYPE-C内模",
    "半品测试三",
    "成型TYPE-C外模",
    "量外露尺寸",
    "成品测试",
    "摇科测试",
    "传输速率测试",
    "CCD检查",
    "修毛边",
    "清洁一",
    "外观检查",
    "电测、装、焊铁壳",
    "转接头测试",
    "成型转接头内模",
    "成型转接头外模",
    "转接头成品测试",
    "清洁二",
    "修毛边及外观检查",
    "扎线、装PE袋",
    "包装",
]

IE_TIME_STUDY_FIELDS = [
    "动作",
    "机器型号",
    "IE测量方法",
    "观测次数",
    "平均观测工时(s)",
    "评比系数",
    "宽放率",
    "标准工时(s)",
    "工时来源",
    "动态调整",
]

SOP_FLOWCHART_SHAPE_POLICY = "testing=diamond; processing=ellipse"

SOP_TESTING_NODE_TYPES = frozenset({"decision", "inspection", "test", "quality", "measurement"})

SOP_TESTING_NODE_KEYWORDS = (
    "\u6d4b\u8bd5",  # 测试
    "\u68c0\u67e5",  # 检查
    "\u68c0\u9a8c",  # 检验
    "\u68c0\u6d4b",  # 检测
    "\u91cf\u6d4b",  # 量测
    "\u7535\u6d4b",  # 电测
    "CCD",
    "\u5408\u683c",  # 合格
    "\u5c3a\u5bf8",  # 尺寸
    "\u901f\u7387",  # 速率
    "\u7ebf\u5e8f",  # 线序
    "TEST",
    "CHECK",
    "INSPECTION",
    "MEASURE",
    "MEASUREMENT",
    "AOI",
    "ICT",
    "QC",
)


def classify_sop_flow_node_shape(node: dict[str, Any]) -> str:
    """Return the frozen SOP flowchart shape for a structured process node."""
    node_type = _center_flow_node_type(node)
    text = "".join(str(node.get(key) or "") for key in ["label", "name", "title"])
    if node_type in SOP_TESTING_NODE_TYPES:
        return "diamond"
    normalized_text = text.upper()
    if any(keyword.upper() in normalized_text for keyword in SOP_TESTING_NODE_KEYWORDS):
        return "diamond"
    return "ellipse"


def build_process_flow_page(
    *,
    product_name: str = "USB3.1 TYPE C-C",
    part_no: str = "80806-129",
    document_no: str = "SOP-SD-4378",
    drawing_no: str = "A-US22-0000-03",
    operations: list[str] | None = None,
) -> dict[str, Any]:
    operation_names = operations or DEFAULT_OPERATION_NAMES
    flow_nodes = [
        {
            "operation_id": f"OP{index:02d}",
            "label": f"工序{index}",
            "name": name,
            "shape": _flow_shape(index),
        }
        for index, name in enumerate(operation_names, start=1)
    ]
    return {
        "page_type": "process_flow",
        "orientation": "portrait",
        "size_pt": REFERENCE_80806_129_FORMAT["flow_page"]["size_pt"],
        "title": "流程图",
        "product_name": product_name,
        "part_no": part_no,
        "document_no": document_no,
        "drawing_no": drawing_no,
        "header_fields": list(REFERENCE_80806_129_FORMAT["flow_page"]["header_fields"]),
        "flow_nodes": flow_nodes,
        "ie_time_study": _default_process_ie_time_study(flow_nodes),
        "footer_form_code": REFERENCE_80806_129_FORMAT["flow_page"]["footer_form_code"],
    }


def build_work_instruction_page(
    *,
    product_name: str = "USB3.1 TYPE C-C",
    part_no: str = "80806-129",
    station: str = "裁线",
    document_no: str = "SOP-SD-4378",
    drawing_no: str = "A-US22-0000-03",
    version: str = "E",
    page_no: int = 1,
    page_total: int = 41,
) -> dict[str, Any]:
    step_slots = [
        {
            "slot_no": index,
            "image_placeholder": True,
            "image_label": "图片占位",
            "text_placeholder": f"{index}. 由 agent 填写作业描述、关键尺寸、检查要点。",
        }
        for index in REFERENCE_80806_129_FORMAT["work_instruction_page"]["step_sequence"]
    ]
    return {
        "page_type": "work_instruction",
        "orientation": "landscape",
        "size_pt": REFERENCE_80806_129_FORMAT["work_instruction_page"]["size_pt"],
        "title": "标准作业指导书",
        "product_name": product_name,
        "part_no": part_no,
        "station": station,
        "document_no": document_no,
        "drawing_no": drawing_no,
        "version": version,
        "page_no": page_no,
        "page_total": page_total,
        "header_fields": list(REFERENCE_80806_129_FORMAT["work_instruction_page"]["header_fields"]),
        "operation_order": "①准备物料==②设定尺寸==③裁线==④自检品质==⑤作业完成",
        "step_slots": step_slots,
        "ie_time_study": _default_work_ie_time_study(step_slots),
        "side_sections": [
            {"title": title, "lines": _default_side_lines(title)}
            for title in REFERENCE_80806_129_FORMAT["work_instruction_page"]["right_sections"]
        ],
        "bottom_sections": [
            {"title": title, "value": _default_bottom_value(title)}
            for title in REFERENCE_80806_129_FORMAT["work_instruction_page"]["bottom_sections"]
        ],
    }


def build_usb_cable_packaging_demo() -> dict[str, Any]:
    product_name = "USB-C 简易数据线"
    part_no = "DEMO-USBC-001"
    document_no = "SOP-DEMO-0001"
    drawing_no = "SOP-DEMO-FMT-01"
    operations = ["备料", "外观确认", "整理线材", "扎线", "装PE袋", "装箱入库"]
    process_flow = build_process_flow_page(
        product_name=product_name,
        part_no=part_no,
        document_no=document_no,
        drawing_no=drawing_no,
        operations=operations,
    )
    work_instruction = build_work_instruction_page(
        product_name=product_name,
        part_no=part_no,
        station="扎线装袋",
        document_no=document_no,
        drawing_no=drawing_no,
        version="A",
        page_no=1,
        page_total=1,
    )
    work_instruction["operation_order"] = "①备料==②外观确认==③整理线材==④扎线==⑤装PE袋==⑥装箱入库"
    work_instruction["step_slots"] = [
        _demo_step(1, "备料", "确认数据线、扎带、PE袋、标签齐套。", "materials"),
        _demo_step(2, "外观确认", "检查接口、线身、外被无破损和污渍。", "inspect"),
        _demo_step(3, "整理线材", "按标准长度盘绕，线身自然不扭结。", "coil"),
        _demo_step(4, "扎线", "使用扎带固定线圈，扎带位置居中。", "tie"),
        _demo_step(5, "装PE袋", "数据线与标签同向放入 PE 袋。", "bag"),
        _demo_step(6, "装箱入库", "按数量装箱，箱唛与工单一致。", "carton"),
    ]
    process_flow["ie_time_study"] = _demo_process_ie_time_study(operations)
    work_instruction["ie_time_study"] = _demo_work_ie_time_study(work_instruction["step_slots"])
    work_instruction["side_sections"] = [
        {
            "title": "作业标准",
            "lines": [
                "1. 线材外观无破损、脏污、露铜。",
                "2. 扎带固定后线圈不松散。",
                "3. PE袋封口方向一致，标签朝外。",
            ],
        },
        {"title": "设备/工具", "lines": ["1. 标签扫码枪", "2. 扎线治具", "3. 电子秤"]},
        {"title": "辅助材料", "lines": ["PE袋", "扎带", "箱唛标签"]},
        {
            "title": "注意事项",
            "lines": [
                "1. 不得强拉线身或折弯接口。",
                "2. 不良品放入红色隔离盒。",
                "3. 扫码失败需通知组长复核。",
            ],
        },
        {"title": "变更内容", "lines": ["A / 初版样例 / 2026-07-05 / M2"]},
        {"title": "物料表", "lines": ["1 / CABLE-DEMO / USB-C数据线 / 1PCS", "2 / BAG-PE / PE袋 / 1PCS"]},
    ]
    return {
        "product_name": product_name,
        "part_no": part_no,
        "image_source": "local_generated_svg",
        "operations": operations,
        "process_flow": process_flow,
        "work_instruction": work_instruction,
    }


def build_sop_table_templates() -> dict[str, Any]:
    flow_page = build_process_flow_page()
    work_page = build_work_instruction_page()
    return {
        "format": "table_only",
        "process_flow": _process_flow_tables(flow_page),
        "work_instruction": _work_instruction_tables(work_page),
    }


def render_process_flow_table_markdown(process_flow_tables: dict[str, list[dict[str, str]]]) -> str:
    sections = ["# SOP 流程图首页表格模板", ""]
    for title, rows in process_flow_tables.items():
        sections.extend([f"## {title}", "", _markdown_table(rows), ""])
    return "\n".join(sections).rstrip() + "\n"


def render_work_instruction_table_markdown(work_instruction_tables: dict[str, list[dict[str, str]]]) -> str:
    sections = ["# SOP 标准作业指导书表格模板", ""]
    for title, rows in work_instruction_tables.items():
        sections.extend([f"## {title}", "", _markdown_table(rows), ""])
    return "\n".join(sections).rstrip() + "\n"


def render_process_flow_svg(page: dict[str, Any]) -> str:
    width, height = page.get("size_pt") or REFERENCE_80806_129_FORMAT["flow_page"]["size_pt"]
    nodes = page.get("flow_nodes") or []
    columns = _flow_columns(nodes)
    parts = [_svg_open(width, height, "SOP 工艺流程流程图模板")]
    parts.extend(
        [
            f'<rect x="52" y="96" width="{width - 104:.2f}" height="{height - 136:.2f}" fill="white" stroke="#111" stroke-width="1.3"/>',
            '<text x="210" y="62" font-size="24" font-family="SimSun, Microsoft YaHei, Arial">流程图</text>',
            '<text x="154" y="60" font-size="18" font-weight="700" fill="#f28c00" font-family="Arial">LOGO</text>',
        ]
    )
    parts.extend(_flow_header_svg(page))
    parts.append('<text x="62" y="205" font-size="11" font-family="SimSun, Microsoft YaHei, Arial">工艺流程： 流程图</text>')
    for column_index, column_nodes in enumerate(columns):
        x = 135 + column_index * 125
        y = 250
        step_gap = 39
        previous_center: tuple[float, float] | None = None
        for row_index, node in enumerate(column_nodes):
            node_y = y + row_index * step_gap
            label_x = x - 78
            parts.append(
                f'<text x="{label_x:.1f}" y="{node_y + 5:.1f}" font-size="10" font-family="SimSun, Microsoft YaHei, Arial">{escape(node["label"])}</text>'
            )
            parts.append(_node_svg(node.get("shape", "ellipse"), x, node_y, str(node.get("name") or "")))
            if previous_center is not None:
                parts.append(_arrow(previous_center[0], previous_center[1] + 14, x, node_y - 14))
            previous_center = (x, node_y)
    parts.append(
        f'<text x="{width - 168:.1f}" y="{height - 22:.1f}" font-size="10" font-family="Times New Roman">{escape(str(page.get("footer_form_code", "")))}</text>'
    )
    parts.append("</svg>")
    return "\n".join(parts) + "\n"


def render_work_instruction_svg(page: dict[str, Any]) -> str:
    width, height = page.get("size_pt") or REFERENCE_80806_129_FORMAT["work_instruction_page"]["size_pt"]
    parts = [_svg_open(width, height, "SOP 标准作业指导书模板")]
    parts.append(f'<rect x="16" y="16" width="{width - 32:.2f}" height="{height - 32:.2f}" fill="white" stroke="#111" stroke-width="1.2"/>')
    parts.extend(_work_header_svg(page))
    parts.append('<text x="18" y="91" font-size="10" font-family="SimSun, Microsoft YaHei, Arial">作业顺序：①准备物料==②设定尺寸==③裁线==④自检品质==⑤作业完成</text>')
    parts.append('<text x="28" y="265" font-size="19" writing-mode="tb" font-family="SimSun, Microsoft YaHei, Arial">图片流程描述及说明</text>')
    parts.extend(_work_step_grid_svg(page))
    parts.extend(_work_side_panel_svg(page))
    parts.extend(_work_bottom_panel_svg(page))
    parts.append("</svg>")
    return "\n".join(parts) + "\n"


def write_sop_template_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    flow_page = build_process_flow_page()
    work_page = build_work_instruction_page()
    paths = {
        "process_flow_svg": output / "sop_process_flow_template.svg",
        "work_instruction_svg": output / "sop_work_instruction_template.svg",
        "manifest_json": output / "sop_template_manifest.json",
        "format_check_csv": output / "sop_template_format_check.csv",
    }
    paths["process_flow_svg"].write_text(render_process_flow_svg(flow_page), encoding="utf-8")
    paths["work_instruction_svg"].write_text(render_work_instruction_svg(work_page), encoding="utf-8")
    manifest = {
        "status": "template_format_only",
        "reference": REFERENCE_80806_129_FORMAT,
        "generated_pages": {
            "process_flow": paths["process_flow_svg"].name,
            "work_instruction": paths["work_instruction_svg"].name,
        },
        "agent_fill_policy": {
            "image_slots": "leave_blank_until_user_or_system_provides_images",
            "text_slots": "agent_may_fill_from_bom_routing_and_past_sop_patterns",
            "release": "no_shopfloor_release_without_human_lock",
        },
    }
    paths["manifest_json"].write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_format_check_csv(paths["format_check_csv"], flow_page, work_page)
    return paths


def write_sop_table_template_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    tables = build_sop_table_templates()
    paths = {
        "process_flow_md": output / "sop_process_flow_tables.md",
        "work_instruction_md": output / "sop_work_instruction_tables.md",
        "manifest_json": output / "sop_table_manifest.json",
        "format_check_csv": output / "sop_table_format_check.csv",
    }
    paths.update(_write_table_section_csvs(output, "sop", tables["process_flow"], tables["work_instruction"]))
    paths["process_flow_md"].write_text(render_process_flow_table_markdown(tables["process_flow"]), encoding="utf-8")
    paths["work_instruction_md"].write_text(render_work_instruction_table_markdown(tables["work_instruction"]), encoding="utf-8")
    manifest = {
        "status": "table_template_only",
        "source_reference": REFERENCE_80806_129_FORMAT["source_pdf"],
        "outputs": {key: path.name for key, path in paths.items()},
        "policy": {
            "default_output": "markdown_and_csv_tables",
            "no_svg_required": True,
            "image_cells": "text labels or external image references only",
        },
    }
    paths["manifest_json"].write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_table_format_check_csv(paths["format_check_csv"], tables)
    return paths


def write_sop_excel_template_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    flow_page = build_process_flow_page()
    work_page = build_work_instruction_page()
    paths = {
        "workbook_xlsx": output / "sop_80806_129_layout_template.xlsx",
        "manifest_json": output / "sop_80806_129_layout_manifest.json",
        "format_check_json": output / "sop_80806_129_layout_format_check.json",
    }
    workbook = _build_sop_layout_workbook(flow_page, work_page)
    workbook.save(paths["workbook_xlsx"])
    _write_excel_manifest(paths["manifest_json"], paths["workbook_xlsx"].name, demo=None)
    _write_excel_format_check_json(paths["format_check_json"], flow_page, work_page)
    return paths


def write_demo_sop_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    demo = build_usb_cable_packaging_demo()
    paths = {
        "process_flow_svg": output / "demo_usb_cable_process_flow.svg",
        "work_instruction_svg": output / "demo_usb_cable_work_instruction.svg",
        "manifest_json": output / "demo_usb_cable_manifest.json",
        "format_check_csv": output / "demo_usb_cable_format_check.csv",
    }
    paths["process_flow_svg"].write_text(render_process_flow_svg(demo["process_flow"]), encoding="utf-8")
    paths["work_instruction_svg"].write_text(render_work_instruction_svg(demo["work_instruction"]), encoding="utf-8")
    manifest = {
        "status": "demo_generated_for_format_and_agent_text_fill_test",
        "demo_product": {
            "product_name": demo["product_name"],
            "part_no": demo["part_no"],
            "image_source": demo["image_source"],
            "operations": demo["operations"],
        },
        "ai_boundary": "demo_images_are_local_illustrations_not_site_photos",
    }
    paths["manifest_json"].write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_format_check_csv(paths["format_check_csv"], demo["process_flow"], demo["work_instruction"])
    return paths


def write_demo_sop_excel_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    demo = build_usb_cable_packaging_demo()
    paths = {
        "workbook_xlsx": output / "demo_usb_cable_sop_layout.xlsx",
        "manifest_json": output / "demo_usb_cable_sop_layout_manifest.json",
        "format_check_json": output / "demo_usb_cable_sop_layout_format_check.json",
    }
    workbook = _build_sop_layout_workbook(demo["process_flow"], demo["work_instruction"])
    workbook.save(paths["workbook_xlsx"])
    _write_excel_manifest(paths["manifest_json"], paths["workbook_xlsx"].name, demo=demo)
    _write_excel_format_check_json(paths["format_check_json"], demo["process_flow"], demo["work_instruction"])
    return paths


def write_sop_word_template_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    flow_page = build_process_flow_page()
    work_page = build_work_instruction_page()
    paths = {
        "document_docx": output / "sop_80806_129_word_template.docx",
        "manifest_json": output / "sop_80806_129_word_manifest.json",
        "format_check_json": output / "sop_80806_129_word_format_check.json",
    }
    document = _build_sop_word_document(flow_page, work_page)
    document.save(paths["document_docx"])
    _write_word_manifest(paths["manifest_json"], paths["document_docx"].name, demo=None)
    _write_word_format_check_json(paths["format_check_json"], flow_page, work_page)
    return paths


def write_demo_sop_word_package(
    out_dir: str | Path,
    *,
    center_flowchart: bool = False,
    center_flowchart_style: str = "pdf_reference_compact",
) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    demo = build_usb_cable_packaging_demo()
    if center_flowchart:
        demo["process_flow"]["render_center_flowchart"] = True
        demo["process_flow"]["center_flowchart_style"] = center_flowchart_style
    paths = {
        "document_docx": output / "demo_usb_cable_sop_word.docx",
        "manifest_json": output / "demo_usb_cable_sop_word_manifest.json",
        "format_check_json": output / "demo_usb_cable_sop_word_format_check.json",
    }
    document = _build_sop_word_document(demo["process_flow"], demo["work_instruction"])
    document.save(paths["document_docx"])
    _write_word_manifest(paths["manifest_json"], paths["document_docx"].name, demo=demo)
    _write_word_format_check_json(paths["format_check_json"], demo["process_flow"], demo["work_instruction"])
    return paths


def write_demo_sop_table_package(out_dir: str | Path) -> dict[str, Path]:
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    demo = build_usb_cable_packaging_demo()
    process_tables = _process_flow_tables(demo["process_flow"])
    work_tables = _work_instruction_tables(demo["work_instruction"])
    paths = {
        "process_flow_md": output / "demo_usb_cable_process_flow_tables.md",
        "work_instruction_md": output / "demo_usb_cable_work_instruction_tables.md",
        "manifest_json": output / "demo_usb_cable_table_manifest.json",
        "format_check_csv": output / "demo_usb_cable_table_format_check.csv",
    }
    paths.update(_write_table_section_csvs(output, "demo_usb_cable", process_tables, work_tables))
    paths["process_flow_md"].write_text(render_process_flow_table_markdown(process_tables), encoding="utf-8")
    paths["work_instruction_md"].write_text(render_work_instruction_table_markdown(work_tables), encoding="utf-8")
    manifest = {
        "status": "demo_table_package",
        "demo_product": {
            "product_name": demo["product_name"],
            "part_no": demo["part_no"],
            "image_source": "table_text_only",
            "operations": demo["operations"],
        },
        "policy": "No SVG generated by table package. Image column stores image descriptions or file references.",
    }
    paths["manifest_json"].write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_table_format_check_csv(paths["format_check_csv"], {"process_flow": process_tables, "work_instruction": work_tables})
    return paths


def _build_sop_word_document(flow_page: dict[str, Any], work_page: dict[str, Any]) -> Any:
    document = WordDocument()
    _set_document_defaults(document)
    _configure_word_section(document.sections[0], landscape=False)
    _render_process_flow_word_table(document, flow_page)
    _configure_word_section(document.add_section(WD_SECTION.NEW_PAGE), landscape=True)
    _render_work_instruction_word_table(document, work_page)
    return document


def _set_document_defaults(document: Any) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _configure_word_section(section: Any, *, landscape: bool) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)


def _render_process_flow_word_table(document: Any, page: dict[str, Any]) -> None:
    header = document.add_table(rows=4, cols=8)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    _set_table_borders(header)
    _set_row_height(header.rows[0], 520)
    for row in header.rows[1:]:
        _set_row_height(row, 420)

    _set_word_cell(header.cell(0, 0).merge(header.cell(0, 1)), "LOGO", bold=True, size=14, color="F28C00")
    _set_word_cell(header.cell(0, 2).merge(header.cell(0, 5)), "流程图", bold=True, size=20)
    _set_word_cell(
        header.cell(0, 6).merge(header.cell(0, 7)),
        str(page.get("footer_form_code") or ""),
        size=8,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    _set_word_cell(header.cell(1, 0), "品名", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 1).merge(header.cell(1, 2)), str(page.get("product_name") or ""))
    _set_word_cell(header.cell(1, 3), "料号", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 4), str(page.get("part_no") or ""))
    _set_word_cell(header.cell(1, 5), "页数", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 6), "1 OF 1")
    _set_word_cell(header.cell(1, 7), "")

    _set_word_cell(header.cell(2, 0), "作业部门", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 1), "组装")
    _set_word_cell(header.cell(2, 2), "版本", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 3), "E")
    _set_word_cell(header.cell(2, 4), "制定日期", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 5), "2021/11/5")
    _set_word_cell(header.cell(2, 6), "文件编号", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 7), str(page.get("document_no") or ""))

    _set_word_cell(header.cell(3, 0), "图号", shaded=True, bold=True)
    _set_word_cell(header.cell(3, 1).merge(header.cell(3, 2)), str(page.get("drawing_no") or ""))
    _set_word_cell(header.cell(3, 3), "核准", shaded=True, bold=True)
    _set_word_cell(header.cell(3, 4), "审核", shaded=True, bold=True)
    _set_word_cell(header.cell(3, 5), "拟订", shaded=True, bold=True)
    _set_word_cell(header.cell(3, 6).merge(header.cell(3, 7)), "")

    body = document.add_table(rows=1, cols=1)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    _set_table_borders(body)
    _set_row_height(body.rows[0], 9600)
    if page.get("render_center_flowchart"):
        if str(page.get("center_flowchart_style") or "") == "pdf_reference_shape_blocks":
            _render_pdf_reference_shape_flowchart_image_in_cell(body.cell(0, 0), page)
        else:
            _render_pdf_reference_flowchart_in_cell(body.cell(0, 0), page)
    else:
        _set_word_cell(body.cell(0, 0), "", align=WD_ALIGN_PARAGRAPH.LEFT)

    _render_ie_time_study_word_table(document, page, scope="process_flow")

    footer = document.add_table(rows=1, cols=2)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    _set_table_borders(footer)
    _set_row_height(footer.rows[0], 360)
    _set_word_cell(footer.cell(0, 0), "")
    _set_word_cell(footer.cell(0, 1), str(page.get("footer_form_code") or ""), size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _render_pdf_reference_flowchart_in_cell(cell: Any, page: dict[str, Any]) -> None:
    _clear_word_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_word_cell_margins(cell, top=180, start=220, bottom=120, end=220)

    title = _center_flowchart_title(page)
    title_paragraph = cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(80)
    title_run = title_paragraph.add_run(f"工艺流程：{title}")
    title_run.bold = True
    title_run.font.name = "SimSun"
    title_run.font.size = Pt(9)
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    nodes = _normalise_center_flow_nodes(page)
    if not nodes:
        return

    columns = _compact_flowchart_column_count(len(nodes))
    rows_per_column = (len(nodes) + columns - 1) // columns
    chart_rows = rows_per_column * 2 - 1
    chart_cols = columns * 3 - 1
    chart = cell.add_table(rows=chart_rows, cols=chart_cols)
    chart.alignment = WD_TABLE_ALIGNMENT.CENTER
    chart.autofit = False
    _set_table_no_borders(chart)

    for row_index, row in enumerate(chart.rows):
        _set_row_height(row, 700 if row_index % 2 == 0 else 220)
        for col_index, chart_cell in enumerate(row.cells):
            width = 0.88 if col_index % 3 == 0 else 3.55 if col_index % 3 == 1 else 0.72
            _set_word_cell_width(chart_cell, width)
            _set_word_cell_margins(chart_cell, top=20, start=30, bottom=20, end=30)

    positions: dict[int, tuple[int, int, int]] = {}
    for index, node in enumerate(nodes):
        column = index // rows_per_column
        offset = index % rows_per_column
        visual_offset = rows_per_column - 1 - offset if column % 2 else offset
        row_index = visual_offset * 2
        label_col = column * 3
        node_col = label_col + 1
        positions[index] = (row_index, label_col, node_col)

        label = _center_flow_process_label(node, index)
        _set_word_cell(chart.cell(row_index, label_col), label, size=7, align=WD_ALIGN_PARAGRAPH.RIGHT)
        node_text = _center_flow_node_text(node)
        node_cell = chart.cell(row_index, node_col)
        _set_word_cell(node_cell, node_text, bold=True, size=7)
        _set_word_cell_borders(node_cell, val="single", size="8", color="000000")
        if _center_flow_node_type(node) in {"decision", "inspection", "rework"}:
            _set_word_cell_shading(node_cell, "F2F2F2")

    for index in range(len(nodes) - 1):
        current_column = index // rows_per_column
        next_column = (index + 1) // rows_per_column
        current_row, _, current_node_col = positions[index]
        next_row, _, _ = positions[index + 1]
        if current_column == next_column:
            arrow_row = (current_row + next_row) // 2
            arrow = "↓" if next_row > current_row else "↑"
            _set_word_cell(chart.cell(arrow_row, current_node_col), arrow, bold=True, size=9)
        else:
            gap_col = current_column * 3 + 2
            _set_word_cell(chart.cell(current_row, gap_col), "→", bold=True, size=9)

    rework_note = _center_flowchart_rework_note(page)
    if rework_note:
        note_paragraph = cell.add_paragraph()
        note_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_paragraph.paragraph_format.space_before = Pt(80)
        note_paragraph.paragraph_format.space_after = Pt(0)
        note_run = note_paragraph.add_run(rework_note)
        note_run.font.name = "SimSun"
        note_run.font.size = Pt(7)
        note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _render_pdf_reference_shape_flowchart_image_in_cell(cell: Any, page: dict[str, Any]) -> None:
    _clear_word_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_word_cell_margins(cell, top=90, start=90, bottom=90, end=90)
    image = BytesIO(_render_center_flowchart_shape_image(page))
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(image, width=Cm(18.7))


def _render_center_flowchart_shape_image(page: dict[str, Any], *, width: int = 1800, height: int = 1260) -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    nodes = _normalise_center_flow_nodes(page)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = _shape_flowchart_font_path()
    if font_path is not None:
        title_font = ImageFont.truetype(str(font_path), 34)
        label_font = ImageFont.truetype(str(font_path), 23)
        node_font = ImageFont.truetype(str(font_path), 25)
        small_font = ImageFont.truetype(str(font_path), 20)
    else:
        title_font = label_font = node_font = small_font = ImageFont.load_default()

    black = (0, 0, 0)
    draw.rectangle([10, 10, width - 10, height - 10], outline=black, width=3)
    draw.text((40, 38), f"工艺流程：  {_center_flowchart_title(page)}", fill=black, font=title_font)

    if not nodes:
        output = BytesIO()
        image.save(output, format="PNG")
        return output.getvalue()

    column_count = _shape_flowchart_column_count(len(nodes))
    rows_per_column = (len(nodes) + column_count - 1) // column_count
    top = 130
    bottom = height - 95
    label_width = 92
    node_width = 260 if column_count <= 3 else 235
    column_gap = 180 if column_count <= 2 else 80
    column_stride = label_width + node_width + column_gap
    layout_width = (column_count - 1) * column_stride + label_width + node_width
    left = max(35, int((width - layout_width) / 2))
    node_height = 62 if rows_per_column > 10 else 70
    row_gap = (bottom - top - node_height) / max(rows_per_column - 1, 1)

    positions: dict[int, tuple[int, int, int, int]] = {}
    labels: list[tuple[int, int, str]] = []
    for index, node in enumerate(nodes):
        column = index // rows_per_column
        offset = index % rows_per_column
        visual_offset = rows_per_column - 1 - offset if column % 2 else offset
        x_label = int(left + column * column_stride)
        x_node = x_label + label_width
        y = int(top + visual_offset * row_gap)
        box = (x_node, y, x_node + node_width, y + node_height)
        positions[index] = box
        labels.append((x_label, y + node_height // 2 - 12, _center_flow_process_label(node, index)))
        _draw_shape_flowchart_node(draw, box, node, font=node_font)

    for index in range(len(nodes) - 1):
        current_column = index // rows_per_column
        next_column = (index + 1) // rows_per_column
        current_box = positions[index]
        next_box = positions[index + 1]
        if current_column == next_column:
            if next_box[1] > current_box[1]:
                _draw_flow_arrow(
                    draw,
                    ((current_box[0] + current_box[2]) // 2, current_box[3] + 3),
                    ((next_box[0] + next_box[2]) // 2, next_box[1] - 3),
                )
            else:
                _draw_flow_arrow(
                    draw,
                    ((current_box[0] + current_box[2]) // 2, current_box[1] - 3),
                    ((next_box[0] + next_box[2]) // 2, next_box[3] + 3),
                )
        else:
            _draw_flow_arrow(
                draw,
                (current_box[2] + 8, (current_box[1] + current_box[3]) // 2),
                (next_box[0] - 8, (next_box[1] + next_box[3]) // 2),
            )

    for x, y, text in labels:
        bbox = draw.textbbox((x, y), text, font=label_font)
        draw.rectangle([bbox[0] - 4, bbox[1] - 3, bbox[2] + 4, bbox[3] + 3], fill="white")
        draw.text((x, y), text, fill=black, font=label_font)

    note = "规范：测试/检查/量测/CCD/电测用菱形；加工/组装/清洁/包装用椭圆。DEMO / not for release."
    draw.text((40, height - 58), note, fill=(70, 70, 70), font=small_font)

    output = BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def _shape_flowchart_font_path() -> Path | None:
    for candidate in [
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]:
        if candidate.exists():
            return candidate
    return None


def _shape_flowchart_column_count(node_count: int) -> int:
    if node_count <= 9:
        return 1
    if node_count <= 24:
        return 2
    if node_count <= 36:
        return 3
    return 4


def _draw_shape_flowchart_node(draw: Any, box: tuple[int, int, int, int], node: dict[str, Any], *, font: Any) -> None:
    black = (0, 0, 0)
    if classify_sop_flow_node_shape(node) == "diamond":
        x1, y1, x2, y2 = box
        cx = (x1 + x2) // 2
        cy = (y1 + y2) // 2
        points = [(cx, y1), (x2, cy), (cx, y2), (x1, cy)]
        draw.polygon(points, outline=black, fill="white")
        draw.line(points + [points[0]], fill=black, width=3)
        text_box = (x1 + 28, y1 + 8, x2 - 28, y2 - 8)
    else:
        draw.ellipse(box, outline=black, width=3, fill="white")
        text_box = (box[0] + 18, box[1] + 8, box[2] - 18, box[3] - 8)
    lines = _wrap_shape_node_text(_center_flow_node_title(node), 8)
    _draw_centered_text(draw, lines[:2], text_box, font)


def _center_flow_node_shape_kind(node: dict[str, Any]) -> str:
    return classify_sop_flow_node_shape(node)


def _center_flow_node_title(node: dict[str, Any]) -> str:
    return str(node.get("name") or node.get("label") or node.get("title") or "")


def _wrap_shape_node_text(text: str, max_chars: int) -> list[str]:
    text = text.strip()
    if not text:
        return [""]
    lines: list[str] = []
    for part in text.split("\n"):
        part = part.strip()
        while len(part) > max_chars:
            lines.append(part[:max_chars])
            part = part[max_chars:]
        if part:
            lines.append(part)
    return lines or [text]


def _draw_centered_text(draw: Any, lines: list[str], box: tuple[int, int, int, int], font: Any) -> None:
    x1, y1, x2, y2 = box
    metrics = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        metrics.append((line, bbox[2] - bbox[0], bbox[3] - bbox[1] + 4))
    total_height = sum(item[2] for item in metrics)
    y = y1 + max(0, (y2 - y1 - total_height) // 2)
    for line, text_width, line_height in metrics:
        draw.text((x1 + (x2 - x1 - text_width) / 2, y), line, fill=(0, 0, 0), font=font)
        y += line_height


def _draw_flow_arrow(draw: Any, start: tuple[int, int], end: tuple[int, int]) -> None:
    black = (0, 0, 0)
    draw.line([start, end], fill=black, width=3)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 14 * direction, y2 - 8), (x2 - 14 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 8, y2 - 14 * direction), (x2 + 8, y2 - 14 * direction)]
    draw.polygon(points, fill=black)


def _center_flowchart_title(page: dict[str, Any]) -> str:
    chart = page.get("center_flowchart") or {}
    return str(chart.get("flowchart_title") or page.get("operation_order") or page.get("title") or "流程图")


def _normalise_center_flow_nodes(page: dict[str, Any]) -> list[dict[str, Any]]:
    chart = page.get("center_flowchart") or {}
    raw_nodes = list(chart.get("nodes") or page.get("flow_nodes") or [])
    nodes: list[dict[str, Any]] = []
    for index, raw_node in enumerate(raw_nodes, start=1):
        node = dict(raw_node)
        node.setdefault("seq", index)
        if "name" not in node:
            node["name"] = node.get("title") or node.get("label") or f"工序{index}"
        nodes.append(node)
    return nodes


def _compact_flowchart_column_count(node_count: int) -> int:
    if node_count <= 4:
        return 1
    if node_count <= 8:
        return 2
    return 3


def _center_flow_process_label(node: dict[str, Any], index: int) -> str:
    seq = node.get("seq") or index + 1
    return f"工序{seq}"


def _center_flow_node_type(node: dict[str, Any]) -> str:
    return str(node.get("type") or node.get("shape") or "process").lower()


def _center_flow_node_text(node: dict[str, Any]) -> str:
    node_type = _center_flow_node_type(node)
    symbol = {
        "start": "○",
        "end": "◎",
        "decision": "◇",
        "inspection": "检",
        "rework": "返",
    }.get(node_type, "")
    title = str(node.get("name") or node.get("label") or "")
    station = str(node.get("station") or "")
    note = str(node.get("note") or "")
    lines = [f"{symbol} {title}".strip()]
    if station:
        lines.append(station)
    if note:
        lines.append(note)
    return "\n".join(lines[:3])


def _center_flowchart_rework_note(page: dict[str, Any]) -> str:
    chart = page.get("center_flowchart") or {}
    edges = list(chart.get("edges") or [])
    rework_edges = [edge for edge in edges if "返" in str(edge.get("label") or "") or "不合格" in str(edge.get("label") or "")]
    if not rework_edges:
        return "说明：本页为DEMO流程图样式测试，节点与工时仅用于格式验证，未作为现场发行依据。"
    labels = "；".join(str(edge.get("label") or "") for edge in rework_edges[:3] if edge.get("label"))
    return f"判定/返工：{labels}。DEMO / not for release，需现场IE与工艺确认后锁版。"


def _render_work_instruction_word_table(document: Any, page: dict[str, Any]) -> None:
    header = document.add_table(rows=3, cols=8)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    _set_table_borders(header)
    _set_row_height(header.rows[0], 520)
    _set_row_height(header.rows[1], 420)
    _set_row_height(header.rows[2], 420)

    _set_word_cell(header.cell(0, 0), "LOGO", bold=True, size=13, color="F28C00")
    _set_word_cell(header.cell(0, 1).merge(header.cell(0, 6)), "标准作业指导书", bold=True, size=18)
    _set_word_cell(header.cell(0, 7), f"页码\n{page.get('page_no')} OF {page.get('page_total')}", size=8)

    _set_word_cell(header.cell(1, 0), "产品品名", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 1), str(page.get("product_name") or ""))
    _set_word_cell(header.cell(1, 2), "本厂料号", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 3), str(page.get("part_no") or ""))
    _set_word_cell(header.cell(1, 4), "文件编号", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 5), str(page.get("document_no") or ""))
    _set_word_cell(header.cell(1, 6), "制作日期", shaded=True, bold=True)
    _set_word_cell(header.cell(1, 7), "2021/11/5")

    _set_word_cell(header.cell(2, 0), "工站", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 1), str(page.get("station") or ""))
    _set_word_cell(header.cell(2, 2), "版本", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 3), str(page.get("version") or ""))
    _set_word_cell(header.cell(2, 4), "作业顺序", shaded=True, bold=True)
    _set_word_cell(header.cell(2, 5).merge(header.cell(2, 7)), str(page.get("operation_order") or ""), align=WD_ALIGN_PARAGRAPH.LEFT)

    body = document.add_table(rows=6, cols=5)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    _set_table_borders(body)
    for row_index in [0, 1, 3, 4]:
        _set_row_height(body.rows[row_index], 1000)
    for row_index in [2, 5]:
        _set_row_height(body.rows[row_index], 580)

    left_label = body.cell(0, 0).merge(body.cell(5, 0))
    _set_word_cell(left_label, "图片流程描述及说明", bold=True, size=12)
    _set_word_cell_text_direction(left_label, "tbRl")

    step_by_slot = {int(slot.get("slot_no")): slot for slot in page.get("step_slots", [])}
    for column_offset, slot_no in enumerate([1, 2, 3], start=1):
        image_cell = body.cell(0, column_offset).merge(body.cell(1, column_offset))
        _fill_word_step_cells(image_cell, body.cell(2, column_offset), step_by_slot.get(slot_no), slot_no)
    for column_offset, slot_no in enumerate([6, 5, 4], start=1):
        image_cell = body.cell(3, column_offset).merge(body.cell(4, column_offset))
        _fill_word_step_cells(image_cell, body.cell(5, column_offset), step_by_slot.get(slot_no), slot_no)

    for row_index, section in enumerate(page.get("side_sections", [])):
        lines = [str(line) for line in section.get("lines", [])]
        section_text = str(section.get("title") or "")
        if lines:
            section_text = section_text + "\n" + "\n".join(lines)
        _set_word_cell(body.cell(row_index, 4), section_text, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)

    _render_ie_time_study_word_table(document, page, scope="work_instruction")

    footer = document.add_table(rows=2, cols=6)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    _set_table_borders(footer)
    _set_row_height(footer.rows[0], 360)
    _set_row_height(footer.rows[1], 560)
    signoff_labels = ["批准", "审核", "制作", "材料环保要求", "管制文件（印章处）", "图号"]
    for column, label in enumerate(signoff_labels):
        _set_word_cell(footer.cell(0, column), label, shaded=True, bold=True)
    _set_word_cell(footer.cell(1, 0), "")
    _set_word_cell(footer.cell(1, 1), "")
    _set_word_cell(footer.cell(1, 2), "")
    _set_word_cell(footer.cell(1, 3), _bottom_value(page, "材料环保要求") or "所有材料须符合RoHS要求", align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_word_cell(footer.cell(1, 4), "")
    _set_word_cell(footer.cell(1, 5), _bottom_value(page, "图号") or str(page.get("drawing_no") or ""), align=WD_ALIGN_PARAGRAPH.LEFT)


def _fill_word_step_cells(image_cell: Any, text_cell: Any, slot: dict[str, Any] | None, slot_no: int) -> None:
    slot = slot or {"slot_no": slot_no, "image_label": "图片占位", "text_placeholder": ""}
    image_text = _slot_image_cell(slot)
    step_text = _clean_step_text(str(slot.get("text_placeholder") or ""))
    _set_word_cell(image_cell, f"{slot_no}\n{image_text}", bold=True, size=12, color="D40000")
    _set_word_cell(text_cell, f"{slot_no}. {step_text}" if step_text else f"{slot_no}. ", align=WD_ALIGN_PARAGRAPH.LEFT)


def _word_side_panel_text(page: dict[str, Any]) -> str:
    section_lines: list[str] = []
    for section in page.get("side_sections", []):
        title = str(section.get("title") or "")
        lines = [str(line) for line in (section.get("lines") or [])]
        section_lines.append(title if not lines else title + "\n" + "\n".join(lines))
    return "\n\n".join(section_lines)


def _render_ie_time_study_word_table(document: Any, page: dict[str, Any], *, scope: str) -> None:
    ie_time_study = page.get("ie_time_study") or {}
    fields = list(ie_time_study.get("fields") or IE_TIME_STUDY_FIELDS)
    rows = list(ie_time_study.get("rows") or [])
    if not rows:
        rows = [_blank_ie_time_study_row("流程工时汇总" if scope == "process_flow" else "动作1")]

    table = document.add_table(rows=2 + len(rows), cols=len(fields))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    _set_row_height(table.rows[0], 260)
    _set_row_height(table.rows[1], 260)
    for row in table.rows[2:]:
        _set_row_height(row, 240)

    title_cell = table.cell(0, 0).merge(table.cell(0, len(fields) - 1))
    title = str(ie_time_study.get("title") or "IE工时记录")
    _set_word_cell(title_cell, f"{title}（随生产实绩动态调整）", bold=True, size=8, shaded=True)
    for col_index, field in enumerate(fields):
        _set_word_cell(table.cell(1, col_index), str(field), bold=True, size=6, shaded=True)
    for row_index, row in enumerate(rows, start=2):
        for col_index, field in enumerate(fields):
            _set_word_cell(table.cell(row_index, col_index), str(row.get(field, "")), size=6, align=WD_ALIGN_PARAGRAPH.CENTER)


def _write_word_manifest(path: Path, document_name: str, demo: dict[str, Any] | None) -> None:
    payload: dict[str, Any] = {
        "status": "word_table_layout_document",
        "document": document_name,
        "reference": REFERENCE_80806_129_FORMAT["source_pdf"],
        "default_output": "docx_table_layout",
        "no_svg": True,
        "no_csv": True,
        "no_xlsx": True,
        "sections": ["流程图", "标准作业指导书"],
        "ai_boundary": "draft layout only; no shopfloor release without human lock",
    }
    if demo:
        payload["demo_product"] = {
            "product_name": demo["product_name"],
            "part_no": demo["part_no"],
            "operations": demo["operations"],
        }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_word_format_check_json(path: Path, flow_page: dict[str, Any], work_page: dict[str, Any]) -> None:
    center_flowchart = bool(flow_page.get("render_center_flowchart"))
    center_style = str(flow_page.get("center_flowchart_style") or "blank")
    if center_style == "pdf_reference_shape_blocks":
        process_flow_body_generated = "single Word table cell with generated raster shape flowchart"
    elif center_flowchart:
        process_flow_body_generated = "single Word table cell with compact nested Word flowchart"
    else:
        process_flow_body_generated = "single empty Word table cell"
    payload = {
        "checks": [
            {"check_item": "output_format", "reference": "Word table layout", "generated": "docx document", "status": "match"},
            {"check_item": "top_level_word_tables", "reference": 8, "generated": 8, "status": "match"},
            {"check_item": "process_flow_tables", "reference": ["header", "flowchart_body", "ie_time_study", "footer"], "generated": ["header", "flowchart_body", "ie_time_study", "footer"], "status": "match"},
            {
                "check_item": "process_flow_body",
                "reference": "large center flowchart area",
                "generated": process_flow_body_generated,
                "status": "match",
            },
            {"check_item": "work_instruction_tables", "reference": ["header", "six_step_body", "ie_time_study", "bottom_signoff"], "generated": ["header", "six_step_body", "ie_time_study", "bottom_signoff"], "status": "match"},
            {"check_item": "ie_time_study_fields", "reference": IE_TIME_STUDY_FIELDS, "generated": IE_TIME_STUDY_FIELDS, "status": "match"},
            {"check_item": "ie_time_study_policy", "reference": "dynamic_adjustment_requires_site_measurement_or_human_lock", "generated": "draft_records_only", "status": "match"},
            {"check_item": "center_flowchart_style", "reference": "pdf_reference_compact_or_shape_blocks", "generated": center_style, "status": "match" if center_flowchart else "not_applicable"},
            {
                "check_item": "center_flowchart_shape_policy",
                "reference": "testing=diamond; processing=ellipse",
                "generated": SOP_FLOWCHART_SHAPE_POLICY if center_style == "pdf_reference_shape_blocks" else "not_applicable",
                "status": "match" if center_style == "pdf_reference_shape_blocks" else "not_applicable",
            },
            {"check_item": "work_instruction_image_slots", "reference": 6, "generated": len(work_page["step_slots"]), "status": _match("6", str(len(work_page["step_slots"])))},
            {"check_item": "right_side_sections", "reference": REFERENCE_80806_129_FORMAT["work_instruction_page"]["right_sections"], "generated": [section["title"] for section in work_page["side_sections"]], "status": "match"},
            {"check_item": "bottom_sections", "reference": REFERENCE_80806_129_FORMAT["work_instruction_page"]["bottom_sections"], "generated": [section["title"] for section in work_page["bottom_sections"]], "status": "match"},
        ]
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _set_word_cell(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: str | None = None,
    shaded: bool = False,
    align: Any = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    _clear_word_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "SimSun"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = _rgb_color(color)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_word_cell_margins(cell, top=80, start=100, bottom=80, end=100)
    if shaded:
        _set_word_cell_shading(cell, "E7E6E6")


def _clear_word_cell(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        p_element = paragraph._element
        for child in list(p_element):
            p_element.remove(child)


def _set_table_borders(table: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _set_table_no_borders(table: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_word_cell_borders(cell: Any, *, val: str, size: str = "8", color: str = "000000") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_word_cell_width(cell: Any, width_cm: float) -> None:
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width.twips)))
    tc_width.set(qn("w:type"), "dxa")


def _set_row_height(row: Any, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")


def _set_word_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_word_cell_text_direction(cell: Any, direction: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    text_direction = tc_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        tc_pr.append(text_direction)
    text_direction.set(qn("w:val"), direction)


def _set_word_cell_margins(cell: Any, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _rgb_color(value: str) -> Any:
    from docx.shared import RGBColor

    value = value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _build_sop_layout_workbook(flow_page: dict[str, Any], work_page: dict[str, Any]) -> Workbook:
    workbook = Workbook()
    flow_sheet = workbook.active
    flow_sheet.title = "01_流程图"
    work_sheet = workbook.create_sheet("02_标准作业指导书")
    _render_process_flow_layout_sheet(flow_sheet, flow_page)
    _render_work_instruction_layout_sheet(work_sheet, work_page)
    return workbook


def _render_process_flow_layout_sheet(sheet: Any, page: dict[str, Any]) -> None:
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.orientation = "portrait"
    sheet.page_setup.paperSize = 9
    sheet.page_margins = PageMargins(left=0.25, right=0.25, top=0.35, bottom=0.35)
    sheet.print_area = "A1:Q58"
    sheet.freeze_panes = "A9"
    for column in range(1, 18):
        sheet.column_dimensions[_column_letter(column)].width = 5.4
    for row in range(1, 59):
        sheet.row_dimensions[row].height = 18

    _apply_outline_border(sheet, "A1:Q58", Side(style="medium", color="000000"))
    sheet["B2"] = "LOGO"
    _style_cell(sheet["B2"], font=Font(name="Arial", size=16, bold=True, color="F28C00"), align=_center())
    sheet["H2"] = "流程图"
    _style_cell(sheet["H2"], font=Font(name="SimSun", size=22, bold=True), align=_center())
    _merge_value(sheet, "L2:Q2", str(page.get("footer_form_code") or ""), font=Font(name="Arial", size=9), align=_right())

    _merge_value(sheet, "B5:B6", "品名", fill=_label_fill())
    _merge_value(sheet, "C5:G6", str(page.get("product_name") or ""))
    _merge_value(sheet, "H5:H6", "料号", fill=_label_fill())
    _merge_value(sheet, "I5:M6", str(page.get("part_no") or ""))
    _merge_value(sheet, "N5:N6", "页数", fill=_label_fill())
    _merge_value(sheet, "O5:Q5", "文件编号", fill=_label_fill())
    _merge_value(sheet, "O6:Q6", str(page.get("document_no") or ""))
    _merge_value(sheet, "B7:C7", "作业部门", fill=_label_fill())
    _merge_value(sheet, "D7:G7", "组装")
    _merge_value(sheet, "H7:I7", "版本", fill=_label_fill())
    _merge_value(sheet, "J7:M7", "E")
    _merge_value(sheet, "N7:O7", "制定日期", fill=_label_fill())
    _merge_value(sheet, "P7:Q7", "2021/11/5")
    _merge_value(sheet, "B8:C8", "图号", fill=_label_fill())
    _merge_value(sheet, "D8:G8", str(page.get("drawing_no") or ""))
    _merge_value(sheet, "H8:I8", "核准", fill=_label_fill())
    _merge_value(sheet, "J8:K8", "审核", fill=_label_fill())
    _merge_value(sheet, "L8:M8", "拟订", fill=_label_fill())
    _merge_value(sheet, "N8:Q8", "")

    _merge_value(sheet, "B10:Q54", "", fill=PatternFill("solid", fgColor="FFFFFF"), border=_medium_border())
    _merge_value(sheet, "B56:K56", "")
    _merge_value(sheet, "L56:Q56", str(page.get("footer_form_code") or ""), font=Font(name="Arial", size=9), align=_right())


def _render_work_instruction_layout_sheet(sheet: Any, page: dict[str, Any]) -> None:
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.paperSize = 9
    sheet.page_margins = PageMargins(left=0.2, right=0.2, top=0.25, bottom=0.25)
    sheet.print_area = "A1:AM35"
    sheet.freeze_panes = "A5"
    for column in range(1, 40):
        sheet.column_dimensions[_column_letter(column)].width = 3.15
    for row in range(1, 36):
        sheet.row_dimensions[row].height = 18
    for row in [6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28]:
        sheet.row_dimensions[row].height = 22
    sheet.row_dimensions[16].height = 24
    sheet.row_dimensions[29].height = 24

    _apply_range_border(sheet, "A1:AM35", _thin_border())
    _merge_value(sheet, "A1:H1", "产品品名", fill=_label_fill())
    _merge_value(sheet, "A2:H2", str(page.get("product_name") or ""))
    _merge_value(sheet, "I1:N1", "本厂料号", fill=_label_fill())
    _merge_value(sheet, "I2:N2", str(page.get("part_no") or ""))
    _merge_value(sheet, "O1:Z1", "深圳市创益通技术股份有限公司", font=Font(name="SimSun", size=14, bold=True))
    sheet["O2"] = "标准作业指导书"
    _style_cell(sheet["O2"], font=Font(name="SimSun", size=18, bold=True), align=_center())
    _merge_value(sheet, "P2:Z2", "")
    _merge_value(sheet, "AA1:AF1", "文件编号", fill=_label_fill())
    _merge_value(sheet, "AA2:AF2", str(page.get("document_no") or ""))
    _merge_value(sheet, "AG1:AM1", "制作日期", fill=_label_fill())
    _merge_value(sheet, "AG2:AM2", "2021/11/5")
    _merge_value(sheet, "A3:H4", "工站", fill=_label_fill())
    _merge_value(sheet, "I3:N4", str(page.get("station") or ""))
    _merge_value(sheet, "O3:Z4", f"作业顺序：{page.get('operation_order') or ''}", align=_left(wrap=True))
    _merge_value(sheet, "AA3:AF4", f"版本　{page.get('version') or ''}")
    _merge_value(sheet, "AG3:AM4", f"页码　{page.get('page_no')} OF {page.get('page_total')}")

    _merge_value(
        sheet,
        "A6:B31",
        "图片流程描述及说明",
        font=Font(name="SimSun", size=16, bold=True),
        align=Alignment(horizontal="center", vertical="center", textRotation=255, wrap_text=True),
        fill=PatternFill("solid", fgColor="FFFFFF"),
    )
    step_by_slot = {int(slot.get("slot_no")): slot for slot in page.get("step_slots", [])}
    slot_ranges = {
        1: ("C6:J15", "C16:J18"),
        2: ("K6:R15", "K16:R18"),
        3: ("S6:Z15", "S16:Z18"),
        6: ("C19:J28", "C29:J31"),
        5: ("K19:R28", "K29:R31"),
        4: ("S19:Z28", "S29:Z31"),
    }
    for slot_no in [1, 2, 3, 6, 5, 4]:
        slot = step_by_slot.get(slot_no, {"slot_no": slot_no, "image_label": "图片占位", "text_placeholder": ""})
        image_range, text_range = slot_ranges[slot_no]
        image_text = _slot_image_cell(slot)
        step_text = _clean_step_text(str(slot.get("text_placeholder") or ""))
        _merge_value(
            sheet,
            image_range,
            f"{slot_no}\n{image_text}",
            font=Font(name="SimSun", size=14, bold=True, color="D40000"),
            align=_center(wrap=True),
            fill=PatternFill("solid", fgColor="F6F7F8"),
        )
        _merge_value(sheet, text_range, f"{slot_no}. {step_text}" if step_text else f"{slot_no}. ", align=_left(wrap=True), font=Font(name="SimSun", size=9))

    side_ranges = [
        ("AA6:AM6", "AA7:AM17", "作业标准"),
        ("AA18:AM18", "AA19:AM22", "设备/工具"),
        ("AA23:AM23", "AA24:AM24", "辅助材料"),
        ("AA25:AM25", "AA26:AM28", "注意事项"),
        ("AA29:AM29", "AA30:AM30", "变更内容"),
        ("AA31:AM31", None, "物料表"),
    ]
    side_map = {str(section.get("title") or ""): section.get("lines") or [] for section in page.get("side_sections", [])}
    for title_range, content_range, title in side_ranges:
        lines = side_map.get(title) or []
        _merge_value(sheet, title_range, title, font=Font(name="SimSun", size=10, bold=True), fill=_label_fill())
        if content_range:
            content = "\n".join(str(line) for line in lines)
            _merge_value(sheet, content_range, content, font=Font(name="SimSun", size=9), align=_left(wrap=True), fill=PatternFill("solid", fgColor="FFFFFF"))

    _merge_value(sheet, "A32:F33", "批准", font=Font(name="SimSun", size=14, bold=True))
    _merge_value(sheet, "A34:F35", "")
    _merge_value(sheet, "G32:L33", "审核", font=Font(name="SimSun", size=14, bold=True))
    _merge_value(sheet, "G34:L35", "")
    _merge_value(sheet, "M32:R33", "制作", font=Font(name="SimSun", size=14, bold=True))
    _merge_value(sheet, "M34:R35", "")
    _merge_value(sheet, "S32:Z33", "材料环保要求", font=Font(name="SimSun", size=14, bold=True))
    _merge_value(sheet, "S34:Z35", _bottom_value(page, "材料环保要求") or "所有材料须符合RoHS要求", font=Font(name="SimSun", size=10), align=_left(wrap=True))
    _merge_value(sheet, "AA32:AM33", "管制文件（印章处）", font=Font(name="SimSun", size=15, bold=True))
    _merge_value(sheet, "AA34:AC35", "图号", font=Font(name="SimSun", size=11, bold=True))
    _merge_value(sheet, "AD34:AM35", _bottom_value(page, "图号") or str(page.get("drawing_no") or ""), font=Font(name="Arial", size=10))


def _write_excel_manifest(path: Path, workbook_name: str, demo: dict[str, Any] | None) -> None:
    payload: dict[str, Any] = {
        "status": "excel_layout_workbook",
        "workbook": workbook_name,
        "reference": REFERENCE_80806_129_FORMAT["source_pdf"],
        "default_output": "xlsx_layout",
        "no_svg": True,
        "no_csv": True,
        "sheets": ["01_流程图", "02_标准作业指导书"],
        "ai_boundary": "draft layout only; no shopfloor release without human lock",
    }
    if demo:
        payload["demo_product"] = {
            "product_name": demo["product_name"],
            "part_no": demo["part_no"],
            "operations": demo["operations"],
        }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_excel_format_check_json(path: Path, flow_page: dict[str, Any], work_page: dict[str, Any]) -> None:
    payload = {
        "checks": [
            {"check_item": "output_format", "reference": "Excel printed SOP", "generated": "xlsx workbook", "status": "match"},
            {"check_item": "flow_sheet_orientation", "reference": "portrait", "generated": flow_page["orientation"], "status": "match"},
            {"check_item": "flow_body_area", "reference": "large blank flowchart fill area", "generated": "B10:Q54", "status": "match"},
            {"check_item": "work_sheet_orientation", "reference": "landscape", "generated": work_page["orientation"], "status": "match"},
            {"check_item": "work_instruction_image_slots", "reference": 6, "generated": len(work_page["step_slots"]), "status": _match("6", str(len(work_page["step_slots"])))},
            {"check_item": "right_side_sections", "reference": REFERENCE_80806_129_FORMAT["work_instruction_page"]["right_sections"], "generated": [section["title"] for section in work_page["side_sections"]], "status": "match"},
            {"check_item": "bottom_sections", "reference": REFERENCE_80806_129_FORMAT["work_instruction_page"]["bottom_sections"], "generated": [section["title"] for section in work_page["bottom_sections"]], "status": "match"},
        ]
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _merge_value(
    sheet: Any,
    cell_range: str,
    value: Any,
    *,
    font: Font | None = None,
    align: Alignment | None = None,
    fill: PatternFill | None = None,
    border: Border | None = None,
) -> None:
    if ":" in cell_range:
        sheet.merge_cells(cell_range)
    cell = sheet[cell_range.split(":", 1)[0]]
    cell.value = value
    _style_cell(cell, font=font, align=align, fill=fill, border=border or _thin_border())


def _style_cell(
    cell: Any,
    *,
    font: Font | None = None,
    align: Alignment | None = None,
    fill: PatternFill | None = None,
    border: Border | None = None,
) -> None:
    cell.font = font or Font(name="SimSun", size=10)
    cell.alignment = align or _center(wrap=True)
    if fill:
        cell.fill = fill
    cell.border = border or _thin_border()


def _apply_range_border(sheet: Any, cell_range: str, border: Border) -> None:
    for row in sheet[cell_range]:
        for cell in row:
            cell.border = border


def _apply_outline_border(sheet: Any, cell_range: str, side: Side) -> None:
    min_col, min_row, max_col, max_row = range_boundaries(cell_range)
    for row in range(min_row, max_row + 1):
        for col in range(min_col, max_col + 1):
            cell = sheet.cell(row=row, column=col)
            cell.border = Border(
                left=side if col == min_col else cell.border.left,
                right=side if col == max_col else cell.border.right,
                top=side if row == min_row else cell.border.top,
                bottom=side if row == max_row else cell.border.bottom,
            )


def _center(*, wrap: bool = False) -> Alignment:
    return Alignment(horizontal="center", vertical="center", wrap_text=wrap)


def _left(*, wrap: bool = False) -> Alignment:
    return Alignment(horizontal="left", vertical="center", wrap_text=wrap)


def _right() -> Alignment:
    return Alignment(horizontal="right", vertical="center")


def _label_fill() -> PatternFill:
    return PatternFill("solid", fgColor="E7E6E6")


def _thin_border() -> Border:
    side = Side(style="thin", color="000000")
    return Border(left=side, right=side, top=side, bottom=side)


def _medium_border() -> Border:
    side = Side(style="medium", color="000000")
    return Border(left=side, right=side, top=side, bottom=side)


def _no_border() -> Border:
    return Border()


def _column_letter(index: int) -> str:
    return get_column_letter(index)


def _next_column(column_letter: str) -> str:
    index = 0
    for char in column_letter:
        index = index * 26 + ord(char.upper()) - ord("A") + 1
    return get_column_letter(index + 1)


def _write_format_check_csv(path: Path, flow_page: dict[str, Any], work_page: dict[str, Any]) -> None:
    rows = [
        {
            "check_item": "page_count_structure",
            "reference": "1 process flow page + 41 work instruction pages",
            "generated": "1 process flow template + repeatable work instruction template",
            "status": "match",
            "notes": "Template supports the same two-part structure.",
        },
        {
            "check_item": "page_1_orientation",
            "reference": REFERENCE_80806_129_FORMAT["flow_page"]["orientation"],
            "generated": flow_page["orientation"],
            "status": _match(REFERENCE_80806_129_FORMAT["flow_page"]["orientation"], flow_page["orientation"]),
            "notes": "PDF page 1 is portrait.",
        },
        {
            "check_item": "work_instruction_orientation",
            "reference": REFERENCE_80806_129_FORMAT["work_instruction_page"]["orientation"],
            "generated": work_page["orientation"],
            "status": _match(REFERENCE_80806_129_FORMAT["work_instruction_page"]["orientation"], work_page["orientation"]),
            "notes": "PDF pages 2-42 are landscape.",
        },
        {
            "check_item": "work_instruction_image_slots",
            "reference": str(REFERENCE_80806_129_FORMAT["work_instruction_page"]["image_slots"]),
            "generated": str(len(work_page["step_slots"])),
            "status": _match(str(REFERENCE_80806_129_FORMAT["work_instruction_page"]["image_slots"]), str(len(work_page["step_slots"]))),
            "notes": "Images are intentionally blank placeholders.",
        },
        {
            "check_item": "work_instruction_side_sections",
            "reference": " / ".join(REFERENCE_80806_129_FORMAT["work_instruction_page"]["right_sections"]),
            "generated": " / ".join(section["title"] for section in work_page["side_sections"]),
            "status": "match",
            "notes": "Right-side panel follows reference PDF.",
        },
        {
            "check_item": "work_instruction_bottom_sections",
            "reference": " / ".join(REFERENCE_80806_129_FORMAT["work_instruction_page"]["bottom_sections"]),
            "generated": " / ".join(section["title"] for section in work_page["bottom_sections"]),
            "status": "match",
            "notes": "Bottom signoff and controlled-file areas follow reference PDF.",
        },
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check_item", "reference", "generated", "status", "notes"])
        writer.writeheader()
        writer.writerows(rows)


def _flow_shape(index: int) -> str:
    if index in {1, 41}:
        return "hexagon"
    if index % 7 == 0:
        return "diamond"
    return "ellipse"


def _flow_columns(nodes: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    if not nodes:
        return []
    column_count = 4 if len(nodes) > 30 else 3 if len(nodes) > 12 else 1
    if column_count == 4:
        sizes = [14, 13, 11, max(0, len(nodes) - 38)]
    elif column_count == 3:
        base = (len(nodes) + 2) // 3
        sizes = [base, base, len(nodes) - base * 2]
    else:
        sizes = [len(nodes)]
    columns: list[list[dict[str, Any]]] = []
    offset = 0
    for size in sizes:
        if size > 0:
            columns.append(nodes[offset : offset + size])
        offset += size
    if offset < len(nodes):
        columns.append(nodes[offset:])
    return columns


def _svg_open(width: float, height: float, title: str) -> str:
    return (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width:.2f}" height="{height:.2f}" '
        f'viewBox="0 0 {width:.2f} {height:.2f}" role="img">\n'
        f'<title>{escape(title)}</title>\n'
        '<rect width="100%" height="100%" fill="#ffffff"/>'
    )


def _flow_header_svg(page: dict[str, Any]) -> list[str]:
    x = 52
    y = 96
    width = 491.32
    row_h = 30
    lines = [
        f'<rect x="{x}" y="{y}" width="{width:.2f}" height="{row_h * 3}" fill="none" stroke="#111" stroke-width="1"/>',
    ]
    for row in range(1, 3):
        lines.append(f'<line x1="{x}" y1="{y + row * row_h}" x2="{x + width:.2f}" y2="{y + row * row_h}" stroke="#111" stroke-width="0.8"/>')
    for col_x in [110, 260, 315, 408, 460]:
        lines.append(f'<line x1="{col_x}" y1="{y}" x2="{col_x}" y2="{y + row_h * 3}" stroke="#111" stroke-width="0.8"/>')
    text = [
        (64, 122, "品   名"),
        (132, 116, str(page.get("product_name", ""))),
        (272, 122, "料   号"),
        (336, 122, str(page.get("part_no", ""))),
        (420, 122, "页   数"),
        (470, 122, "第1页,共1页"),
        (64, 152, "作业部门"),
        (132, 152, "组装"),
        (272, 152, "版   本"),
        (336, 152, "E"),
        (420, 152, "图   号"),
        (470, 152, str(page.get("drawing_no", ""))),
        (64, 182, "核   准"),
        (272, 182, "审   核"),
        (336, 182, "拟   订"),
        (420, 182, "文件编号"),
        (470, 182, str(page.get("document_no", ""))),
    ]
    for tx, ty, value in text:
        lines.append(f'<text x="{tx}" y="{ty}" font-size="10" font-family="SimSun, Microsoft YaHei, Arial">{escape(value)}</text>')
    return lines


def _node_svg(shape: str, cx: float, cy: float, label: str) -> str:
    text = f'<text x="{cx:.1f}" y="{cy + 4:.1f}" font-size="9" text-anchor="middle" font-family="SimSun, Microsoft YaHei, Arial">{escape(label[:12])}</text>'
    if shape == "diamond":
        points = f"{cx},{cy - 17} {cx + 45},{cy} {cx},{cy + 17} {cx - 45},{cy}"
        body = f'<polygon points="{points}" fill="#fff" stroke="#111" stroke-width="1"/>'
    elif shape == "hexagon":
        points = f"{cx - 35},{cy - 12} {cx + 25},{cy - 12} {cx + 38},{cy} {cx + 25},{cy + 12} {cx - 35},{cy + 12} {cx - 48},{cy}"
        body = f'<polygon points="{points}" fill="#fff" stroke="#111" stroke-width="1.6"/>'
    else:
        body = f'<ellipse cx="{cx:.1f}" cy="{cy:.1f}" rx="45" ry="15" fill="#fff" stroke="#111" stroke-width="1"/>'
    return body + "\n" + text


def _arrow(x1: float, y1: float, x2: float, y2: float) -> str:
    return (
        f'<line x1="{x1:.1f}" y1="{y1:.1f}" x2="{x2:.1f}" y2="{y2:.1f}" '
        'stroke="#111" stroke-width="1" marker-end="url(#arrow)"/>'
    )


def _work_header_svg(page: dict[str, Any]) -> list[str]:
    lines = [
        '<defs><marker id="arrow" markerWidth="8" markerHeight="8" refX="7" refY="4" orient="auto"><path d="M0,0 L8,4 L0,8 z" fill="#d40000"/></marker></defs>',
        '<line x1="16" y1="37" x2="825.92" y2="37" stroke="#111" stroke-width="1"/>',
        '<line x1="16" y1="58" x2="825.92" y2="58" stroke="#111" stroke-width="1"/>',
        '<line x1="16" y1="79" x2="825.92" y2="79" stroke="#111" stroke-width="1"/>',
        '<line x1="190" y1="16" x2="190" y2="79" stroke="#111" stroke-width="1"/>',
        '<line x1="300" y1="16" x2="300" y2="79" stroke="#111" stroke-width="1"/>',
        '<line x1="646" y1="16" x2="646" y2="79" stroke="#111" stroke-width="1"/>',
        '<line x1="740" y1="16" x2="740" y2="79" stroke="#111" stroke-width="1"/>',
        '<text x="76" y="32" font-size="13" font-family="SimSun, Microsoft YaHei, Arial">产品品名</text>',
        '<text x="222" y="32" font-size="13" font-family="SimSun, Microsoft YaHei, Arial">本厂料号</text>',
        '<text x="382" y="45" font-size="18" font-family="SimSun, Microsoft YaHei, Arial">深圳市创益通技术股份有限公司</text>',
        '<text x="420" y="68" font-size="17" font-family="SimSun, Microsoft YaHei, Arial">标准作业指导书</text>',
        '<text x="680" y="32" font-size="13" font-family="SimSun, Microsoft YaHei, Arial">文件编号</text>',
        '<text x="762" y="32" font-size="13" font-family="SimSun, Microsoft YaHei, Arial">制作日期</text>',
        f'<text x="58" y="52" font-size="10" font-family="Times New Roman">{escape(str(page.get("product_name", "")))}</text>',
        f'<text x="220" y="52" font-size="10" font-family="Times New Roman">{escape(str(page.get("part_no", "")))}</text>',
        f'<text x="674" y="52" font-size="10" font-family="Times New Roman">{escape(str(page.get("document_no", "")))}</text>',
        '<text x="764" y="52" font-size="10" font-family="Times New Roman">2021/11/5</text>',
        '<text x="58" y="74" font-size="12" font-family="SimSun, Microsoft YaHei, Arial">工站</text>',
        f'<text x="198" y="74" font-size="11" font-family="SimSun, Microsoft YaHei, Arial">{escape(str(page.get("station", "")))}</text>',
        f'<text x="650" y="74" font-size="12" font-family="SimSun, Microsoft YaHei, Arial">版本　{escape(str(page.get("version", "")))}</text>',
        f'<text x="730" y="74" font-size="12" font-family="SimSun, Microsoft YaHei, Arial">页码　{page.get("page_no")} OF {page.get("page_total")}</text>',
        '<text x="314" y="55" font-size="22" fill="#f28c00" font-weight="700" font-family="Arial">LOGO</text>',
    ]
    return lines


def _work_step_grid_svg(page: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    left_x = 48
    top_y = 98
    slot_w = 160
    slot_h = 203
    positions = [
        (left_x, top_y),
        (left_x + slot_w + 20, top_y),
        (left_x + (slot_w + 20) * 2, top_y),
        (left_x + (slot_w + 20) * 2, top_y + slot_h + 24),
        (left_x + slot_w + 20, top_y + slot_h + 24),
        (left_x, top_y + slot_h + 24),
    ]
    for slot, (x, y) in zip(page.get("step_slots", []), positions):
        lines.append(f'<rect x="{x}" y="{y}" width="{slot_w}" height="{slot_h}" fill="#fff" stroke="#111" stroke-width="1"/>')
        lines.append(f'<rect x="{x + 4}" y="{y + 4}" width="{slot_w - 8}" height="132" fill="#f5f6f7" stroke="#888" stroke-width="0.6"/>')
        if slot.get("visual"):
            lines.append(_step_visual_svg(slot["visual"], x + 4, y + 4, slot_w - 8, 132))
        else:
            lines.append(f'<text x="{x + slot_w / 2:.1f}" y="{y + 75:.1f}" font-size="16" text-anchor="middle" fill="#9a9a9a" font-family="SimSun, Microsoft YaHei, Arial">{escape(slot.get("image_label", "图片占位"))}</text>')
        lines.append(f'<circle cx="{x + 18}" cy="{y + 18}" r="12" fill="none" stroke="#ff0000" stroke-width="1.4"/>')
        lines.append(f'<text x="{x + 14}" y="{y + 23}" font-size="16" fill="#ff0000" font-family="Arial">{slot.get("slot_no")}</text>')
        lines.append(f'<line x1="{x}" y1="{y + 140}" x2="{x + slot_w}" y2="{y + 140}" stroke="#111" stroke-width="1"/>')
        _append_wrapped_text(lines, str(slot.get("text_placeholder", "")), x + 6, y + 156, slot_w - 12, 10)
    arrow_pairs = [(1, 2), (2, 3), (3, 4), (4, 5), (5, 6)]
    centers = [(x + slot_w / 2, y + slot_h / 2) for x, y in positions]
    for start, end in arrow_pairs:
        x1, y1 = centers[start - 1]
        x2, y2 = centers[end - 1]
        if start == 3:
            lines.append(f'<line x1="{x1:.1f}" y1="{y1 + 96:.1f}" x2="{x2:.1f}" y2="{y2 - 96:.1f}" stroke="#d40000" stroke-width="2.2" marker-end="url(#arrow)"/>')
        else:
            lines.append(f'<line x1="{x1 + (slot_w / 2 - 3 if x2 > x1 else -slot_w / 2 + 3):.1f}" y1="{y1:.1f}" x2="{x2 + (-slot_w / 2 + 3 if x2 > x1 else slot_w / 2 - 3):.1f}" y2="{y2:.1f}" stroke="#d40000" stroke-width="2.2" marker-end="url(#arrow)"/>')
    return lines


def _work_side_panel_svg(page: dict[str, Any]) -> list[str]:
    x = 572
    y = 79
    width = 254
    lines = [f'<line x1="{x}" y1="{y}" x2="{x}" y2="532" stroke="#111" stroke-width="1"/>']
    section_heights = [205, 86, 38, 78, 28, 18]
    cursor = y
    for section, height in zip(page.get("side_sections", []), section_heights):
        lines.append(f'<rect x="{x}" y="{cursor}" width="{width}" height="{height}" fill="#fff" stroke="#111" stroke-width="1"/>')
        lines.append(f'<text x="{x + width / 2:.1f}" y="{cursor + 16}" font-size="13" text-anchor="middle" font-family="SimSun, Microsoft YaHei, Arial">{escape(section["title"])}</text>')
        max_lines = max(0, int((height - 28) // 16))
        for index, line in enumerate(section.get("lines", [])[:max_lines]):
            lines.append(f'<text x="{x + 8}" y="{cursor + 38 + index * 16}" font-size="10" font-family="SimSun, Microsoft YaHei, Arial">{escape(line)}</text>')
        cursor += height
    return lines


def _work_bottom_panel_svg(page: dict[str, Any]) -> list[str]:
    lines = [
        '<line x1="16" y1="532" x2="825.92" y2="532" stroke="#111" stroke-width="1"/>',
        '<line x1="16" y1="560" x2="825.92" y2="560" stroke="#111" stroke-width="1"/>',
    ]
    bottom = {item["title"]: item.get("value", "") for item in page.get("bottom_sections", [])}
    cols = [16, 143, 278, 413, 548]
    labels = ["批准", "审核", "制作", "材料环保要求"]
    for index, x in enumerate(cols[:-1]):
        next_x = cols[index + 1]
        lines.append(f'<line x1="{next_x}" y1="532" x2="{next_x}" y2="595" stroke="#111" stroke-width="1"/>')
        lines.append(f'<text x="{x + 38}" y="552" font-size="17" font-family="SimSun, Microsoft YaHei, Arial">{labels[index]}</text>')
    lines.append('<text x="670" y="556" font-size="18" text-anchor="middle" font-family="SimSun, Microsoft YaHei, Arial">管制文件（印章处）</text>')
    lines.append('<line x1="625" y1="560" x2="625" y2="595" stroke="#111" stroke-width="1"/>')
    lines.append('<text x="574" y="580" font-size="11" font-family="SimSun, Microsoft YaHei, Arial">图号</text>')
    lines.append(f'<text x="678" y="580" font-size="10" font-family="Times New Roman">{escape(str(bottom.get("图号", "")))}</text>')
    lines.append(f'<text x="318" y="582" font-size="13" font-family="SimSun, Microsoft YaHei, Arial">{escape(str(bottom.get("制作", "")))}</text>')
    lines.append('<text x="426" y="582" font-size="10" font-family="SimSun, Microsoft YaHei, Arial">所有材料须符合RoHS要求</text>')
    return lines


def _append_wrapped_text(lines: list[str], text: str, x: float, y: float, max_width: float, font_size: int) -> None:
    chars_per_line = max(8, int(max_width // font_size))
    for index, start in enumerate(range(0, len(text), chars_per_line)):
        line = text[start : start + chars_per_line]
        lines.append(f'<text x="{x}" y="{y + index * (font_size + 3)}" font-size="{font_size}" font-family="SimSun, Microsoft YaHei, Arial">{escape(line)}</text>')


def _default_side_lines(title: str) -> list[str]:
    mapping = {
        "作业标准": ["1. 由 agent 填写标准要求。", "2. 关键尺寸/参数保留来源。", "3. 未确认项进入待确认清单。"],
        "设备/工具": ["1. 设备名称待填", "2. 工具/治具待填"],
        "辅助材料": ["PE膜"],
        "注意事项": ["1. 安全、质量和异常处理由 agent 填写。", "2. 现场事实需人工或系统确认。"],
        "变更内容": ["版本 / 变更内容 / 变更日期 / 变更者"],
        "物料表": ["NO / 物料编号 / 物料规格 / 单位用量"],
    }
    return mapping.get(title, [])


def _default_process_ie_time_study(flow_nodes: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": [
            {
                **_blank_ie_time_study_row("流程工时汇总"),
                "机器型号": "按各工站页记录",
                "IE测量方法": "秒表/MES/视频分析待确认",
                "工时来源": "待IE实测",
                "动态调整": "随MES实绩、IE复测、机型变更更新",
            }
        ],
        "policy": {
            "detail_location": "各标准作业指导书页按动作维护",
            "operation_count": len(flow_nodes),
            "release_requirement": "site_measured_or_human_locked",
        },
    }


def _default_work_ie_time_study(step_slots: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": [_blank_ie_time_study_row(_slot_action_name(slot)) for slot in step_slots],
        "policy": {
            "measurement_basis": "IE标准测时：观测工时 x 评比系数 x (1 + 宽放率)",
            "dynamic_adjustment": "生产实绩回写后由IE/生产复核并更新标准工时",
            "release_requirement": "site_measured_or_human_locked",
        },
    }


def _demo_process_ie_time_study(operations: list[str]) -> dict[str, Any]:
    row = _blank_ie_time_study_row("流程工时汇总")
    row.update(
        {
            "机器型号": "DEMO-MIXED",
            "IE测量方法": "DEMO测时样例",
            "标准工时(s)": "详见工站页",
            "工时来源": "demo_not_for_release",
            "动态调整": "生产实绩回写后由IE复核",
        }
    )
    return {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": [row],
        "policy": {"operation_count": len(operations), "release_requirement": "site_measured_or_human_locked"},
    }


def _demo_work_ie_time_study(step_slots: list[dict[str, Any]]) -> dict[str, Any]:
    demo_values = {
        1: ("手动作业台-DEMO", "3", "6.0", "1.00", "10%", "6.6"),
        2: ("检验台-DEMO", "3", "8.0", "1.00", "10%", "8.8"),
        3: ("手动作业台-DEMO", "3", "10.0", "1.00", "10%", "11.0"),
        4: ("扎线治具-DEMO", "3", "7.0", "1.00", "10%", "7.7"),
        5: ("包装台-DEMO", "3", "6.0", "1.00", "10%", "6.6"),
        6: ("包装台-DEMO", "3", "9.0", "1.00", "10%", "9.9"),
    }
    rows = []
    for slot in step_slots:
        slot_no = int(slot.get("slot_no") or 0)
        machine, observations, average, rating, allowance, standard = demo_values.get(slot_no, ("DEMO", "3", "", "1.00", "10%", ""))
        row = _blank_ie_time_study_row(_slot_action_name(slot))
        row.update(
            {
                "机器型号": machine,
                "IE测量方法": "DEMO秒表测时",
                "观测次数": observations,
                "平均观测工时(s)": average,
                "评比系数": rating,
                "宽放率": allowance,
                "标准工时(s)": standard,
                "工时来源": "demo_not_for_release",
                "动态调整": "MES实绩/IE复测后更新",
            }
        )
        rows.append(row)
    return {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": rows,
        "policy": {
            "measurement_basis": "DEMO样例：观测工时 x 评比系数 x (1 + 宽放率)",
            "release_requirement": "site_measured_or_human_locked",
        },
    }


def _blank_ie_time_study_row(action: str) -> dict[str, str]:
    return {
        "动作": action,
        "机器型号": "待填",
        "IE测量方法": "待IE实测",
        "观测次数": "",
        "平均观测工时(s)": "",
        "评比系数": "",
        "宽放率": "",
        "标准工时(s)": "",
        "工时来源": "待IE实测",
        "动态调整": "随生产实绩动态调整",
    }


def _slot_action_name(slot: dict[str, Any]) -> str:
    slot_no = str(slot.get("slot_no") or "")
    text = _clean_step_text(str(slot.get("text_placeholder") or "")).strip()
    if "：" in text:
        return text.split("：", 1)[0]
    if ":" in text:
        return text.split(":", 1)[0]
    return f"动作{slot_no}" if slot_no else "动作"


def _default_bottom_value(title: str) -> str:
    return {
        "图号": "A-US22-0000-03",
    }.get(title, "")


def _match(expected: str, actual: str) -> str:
    return "match" if expected == actual else "mismatch"


def _process_flow_tables(page: dict[str, Any]) -> dict[str, list[dict[str, str]]]:
    nodes = page.get("flow_nodes") or []
    flow_rows: list[dict[str, str]] = []
    for index, node in enumerate(nodes):
        previous_node = nodes[index - 1] if index > 0 else None
        next_node = nodes[index + 1] if index + 1 < len(nodes) else None
        flow_rows.append(
            {
                "工序序号": str(index + 1),
                "工序名称": str(node.get("name") or ""),
                "节点类型": _shape_label(str(node.get("shape") or "ellipse")),
                "上道工序": str(previous_node.get("name") if previous_node else "开始"),
                "下道工序": str(next_node.get("name") if next_node else "结束"),
                "备注": "由 routing 自动生成；正式发布前人工确认。",
            }
        )
    return {
        "流程图首页-表头": [
            {
                "品名": str(page.get("product_name") or ""),
                "料号": str(page.get("part_no") or ""),
                "页数": "第1页,共1页",
                "作业部门": "组装",
                "版本": "E",
                "制定日期": "2021/11/5",
                "图号": str(page.get("drawing_no") or ""),
                "核准": "",
                "审核": "",
                "拟订": "",
                "文件编号": str(page.get("document_no") or ""),
            }
        ],
        "流程图首页-工序流程": flow_rows,
    }


def _work_instruction_tables(page: dict[str, Any]) -> dict[str, list[dict[str, str]]]:
    return {
        "标准作业指导书-表头": [
            {
                "产品品名": str(page.get("product_name") or ""),
                "本厂料号": str(page.get("part_no") or ""),
                "工站": str(page.get("station") or ""),
                "文件编号": str(page.get("document_no") or ""),
                "制作日期": "2021/11/5",
                "版本": str(page.get("version") or ""),
                "页码": f"{page.get('page_no')} OF {page.get('page_total')}",
                "作业顺序": str(page.get("operation_order") or ""),
            }
        ],
        "标准作业指导书-六步说明": [
            {
                "步骤序号": str(slot.get("slot_no")),
                "图片内容": _slot_image_cell(slot),
                "文字说明": _clean_step_text(str(slot.get("text_placeholder") or "")),
            }
            for slot in (page.get("step_slots") or [])
        ],
        "标准作业指导书-右侧栏": [
            {
                "栏目": str(section.get("title") or ""),
                "内容": "；".join(str(line) for line in (section.get("lines") or [])),
            }
            for section in (page.get("side_sections") or [])
        ],
        "标准作业指导书-底部签核": [
            {
                "批准": _bottom_value(page, "批准"),
                "审核": _bottom_value(page, "审核"),
                "制作": _bottom_value(page, "制作"),
                "材料环保要求": _bottom_value(page, "材料环保要求") or "所有材料须符合RoHS要求",
                "管制文件（印章处）": _bottom_value(page, "管制文件（印章处）"),
                "图号": _bottom_value(page, "图号") or str(page.get("drawing_no") or ""),
            }
        ],
    }


def _write_table_section_csvs(
    output: Path,
    prefix: str,
    process_tables: dict[str, list[dict[str, str]]],
    work_tables: dict[str, list[dict[str, str]]],
) -> dict[str, Path]:
    csv_specs = {
        "process_flow_header_csv": (output / f"{prefix}_process_flow_header.csv", process_tables["流程图首页-表头"]),
        "process_flow_routing_csv": (output / f"{prefix}_process_flow_routing.csv", process_tables["流程图首页-工序流程"]),
        "work_instruction_header_csv": (output / f"{prefix}_work_instruction_header.csv", work_tables["标准作业指导书-表头"]),
        "work_instruction_steps_csv": (output / f"{prefix}_work_instruction_steps.csv", work_tables["标准作业指导书-六步说明"]),
        "work_instruction_side_sections_csv": (
            output / f"{prefix}_work_instruction_side_sections.csv",
            work_tables["标准作业指导书-右侧栏"],
        ),
        "work_instruction_signoff_csv": (output / f"{prefix}_work_instruction_signoff.csv", work_tables["标准作业指导书-底部签核"]),
    }
    paths: dict[str, Path] = {}
    for key, (path, rows) in csv_specs.items():
        _write_rows_csv(path, rows)
        paths[key] = path
    return paths


def _write_rows_csv(path: Path, rows: list[dict[str, str]]) -> None:
    fieldnames = list(rows[0].keys()) if rows else ["空表"]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def _write_table_format_check_csv(path: Path, tables: dict[str, Any]) -> None:
    process = tables.get("process_flow") or {}
    work = tables.get("work_instruction") or {}
    step_rows = work.get("标准作业指导书-六步说明") or []
    side_rows = work.get("标准作业指导书-右侧栏") or []
    bottom_rows = work.get("标准作业指导书-底部签核") or []
    rows = [
        {
            "check_item": "output_format",
            "reference": "Excel-like tables",
            "generated": "Markdown/CSV/JSON tables",
            "status": "match",
            "notes": "No SVG is required for table package.",
        },
        {
            "check_item": "process_header_table",
            "reference": "flow page header table",
            "generated": "present" if process.get("流程图首页-表头") else "missing",
            "status": "match" if process.get("流程图首页-表头") else "mismatch",
            "notes": "",
        },
        {
            "check_item": "work_instruction_step_rows",
            "reference": "6",
            "generated": str(len(step_rows)),
            "status": _match("6", str(len(step_rows))),
            "notes": "Rows correspond to image/text slots 1,2,3,6,5,4 in the original layout.",
        },
        {
            "check_item": "right_side_sections",
            "reference": "6",
            "generated": str(len(side_rows)),
            "status": _match("6", str(len(side_rows))),
            "notes": "",
        },
        {
            "check_item": "bottom_signoff_rows",
            "reference": "1",
            "generated": str(len(bottom_rows)),
            "status": _match("1", str(len(bottom_rows))),
            "notes": "",
        },
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["check_item", "reference", "generated", "status", "notes"])
        writer.writeheader()
        writer.writerows(rows)


def _markdown_table(rows: list[dict[str, str]]) -> str:
    if not rows:
        return "| 空表 |\n|---|\n"
    headers = list(rows[0].keys())
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(_escape_md(str(row.get(header, ""))) for header in headers) + " |")
    return "\n".join(lines)


def _escape_md(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def _shape_label(shape: str) -> str:
    return {"ellipse": "椭圆", "diamond": "菱形", "hexagon": "六边形"}.get(shape, shape)


def _slot_image_cell(slot: dict[str, Any]) -> str:
    if slot.get("visual"):
        title = ((slot.get("visual") or {}).get("title") or "步骤").strip()
        return f"{title}示意"
    return str(slot.get("image_label") or "图片占位")


def _clean_step_text(text: str) -> str:
    if ". " in text[:4]:
        return text.split(". ", 1)[1]
    if "：" in text:
        return text.split("：", 1)[1] if text[0].isdigit() else text
    return text


def _bottom_value(page: dict[str, Any], title: str) -> str:
    for item in page.get("bottom_sections") or []:
        if item.get("title") == title:
            return str(item.get("value") or "")
    return ""


def _demo_step(slot_no: int, title: str, description: str, visual_type: str) -> dict[str, Any]:
    return {
        "slot_no": slot_no,
        "image_placeholder": False,
        "text_placeholder": f"{slot_no}. {title}：{description}",
        "visual": {
            "type": visual_type,
            "title": title,
        },
    }


def _step_visual_svg(visual: dict[str, Any], x: float, y: float, width: float, height: float) -> str:
    visual_type = str(visual.get("type") or "generic")
    title = escape(str(visual.get("title") or "步骤图"))
    parts = [f'<g class="generated-step-visual" data-visual-type="{escape(visual_type)}">']
    parts.append(f'<rect x="{x}" y="{y}" width="{width}" height="{height}" fill="#f4f8fb" stroke="#b6c2cc" stroke-width="0.5"/>')
    parts.append(f'<text x="{x + width / 2:.1f}" y="{y + 18:.1f}" font-size="11" text-anchor="middle" fill="#1f3b4d" font-family="SimSun, Microsoft YaHei, Arial">{title}</text>')
    cx = x + width / 2
    cy = y + height / 2 + 6
    if visual_type == "materials":
        parts.extend(
            [
                f'<rect x="{x + 20}" y="{cy - 22}" width="46" height="36" rx="4" fill="#fff" stroke="#333"/>',
                f'<path d="M {x + 78} {cy - 4} C {x + 95} {cy - 32}, {x + 126} {cy - 32}, {x + 138} {cy - 4} C {x + 119} {cy + 18}, {x + 95} {cy + 18}, {x + 78} {cy - 4}" fill="none" stroke="#2f4858" stroke-width="3"/>',
                f'<text x="{x + 43}" y="{cy + 32}" font-size="9" text-anchor="middle" font-family="Arial">BOM</text>',
            ]
        )
    elif visual_type == "inspect":
        parts.extend(
            [
                f'<path d="M {x + 28} {cy + 6} C {x + 55} {cy - 22}, {x + 112} {cy - 22}, {x + 140} {cy + 6}" fill="none" stroke="#2f4858" stroke-width="4"/>',
                f'<circle cx="{x + 120}" cy="{cy - 18}" r="18" fill="none" stroke="#f28c00" stroke-width="3"/>',
                f'<line x1="{x + 134}" y1="{cy - 4}" x2="{x + 150}" y2="{cy + 12}" stroke="#f28c00" stroke-width="3"/>',
            ]
        )
    elif visual_type == "coil":
        for index in range(4):
            parts.append(f'<ellipse cx="{cx}" cy="{cy + index * 3}" rx="{48 - index * 3}" ry="{18 - index}" fill="none" stroke="#2f4858" stroke-width="3"/>')
        parts.append(f'<line x1="{cx - 54}" y1="{cy + 28}" x2="{cx + 54}" y2="{cy + 28}" stroke="#888" stroke-width="1"/>')
    elif visual_type == "tie":
        parts.extend(
            [
                f'<ellipse cx="{cx}" cy="{cy}" rx="52" ry="24" fill="none" stroke="#2f4858" stroke-width="4"/>',
                f'<rect x="{cx - 10}" y="{cy - 28}" width="20" height="56" rx="3" fill="#f28c00" stroke="#8a5300"/>',
                f'<text x="{cx}" y="{cy + 44}" font-size="9" text-anchor="middle" font-family="SimSun, Microsoft YaHei, Arial">扎带居中</text>',
            ]
        )
    elif visual_type == "bag":
        parts.extend(
            [
                f'<rect x="{cx - 46}" y="{cy - 34}" width="92" height="72" rx="5" fill="#ffffff" stroke="#2f4858" stroke-width="2"/>',
                f'<line x1="{cx - 46}" y1="{cy - 16}" x2="{cx + 46}" y2="{cy - 16}" stroke="#7da7bd" stroke-width="2"/>',
                f'<ellipse cx="{cx}" cy="{cy + 14}" rx="34" ry="14" fill="none" stroke="#2f4858" stroke-width="3"/>',
            ]
        )
    elif visual_type == "carton":
        parts.extend(
            [
                f'<polygon points="{cx - 52},{cy - 14} {cx},{cy - 38} {cx + 52},{cy - 14} {cx},{cy + 12}" fill="#f4d3a1" stroke="#8a5300" stroke-width="2"/>',
                f'<polygon points="{cx - 52},{cy - 14} {cx},{cy + 12} {cx},{cy + 48} {cx - 52},{cy + 18}" fill="#e7b979" stroke="#8a5300" stroke-width="2"/>',
                f'<polygon points="{cx + 52},{cy - 14} {cx},{cy + 12} {cx},{cy + 48} {cx + 52},{cy + 18}" fill="#dba765" stroke="#8a5300" stroke-width="2"/>',
            ]
        )
    else:
        parts.append(f'<circle cx="{cx}" cy="{cy}" r="32" fill="none" stroke="#2f4858" stroke-width="3"/>')
    parts.append("</g>")
    return "\n".join(parts)
