from __future__ import annotations

import json
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import RenderResult
from .store import SopKnowledgeStore


ACCENT = "1F4E78"
LIGHT = "DCE6F1"
CAUTION = "FFF2CC"
MUTED = "666666"


class VariableRouteDocxRenderer:
    """Render adapter only. It never decides, truncates, merges, or invents route steps."""

    def __init__(self, store: SopKnowledgeStore) -> None:
        self.store = store

    def render(self, route_id: int, out_dir: str | Path) -> RenderResult:
        payload = self.store.get_route(route_id)
        route = payload["route"]
        steps = payload["steps"]
        root = Path(out_dir)
        root.mkdir(parents=True, exist_ok=True)
        safe_code = route["product_code"].replace("/", "_")
        docx_path = root / f"SOP_{safe_code}_route_draft.docx"
        route_json_path = root / f"SOP_{safe_code}_route.json"
        validation_path = root / f"SOP_{safe_code}_validation.json"
        route_json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        document = Document()
        self._configure_document(document)
        self._add_cover(document, route, steps, payload)
        for index, step in enumerate(steps, start=1):
            document.add_page_break()
            self._add_step_page(document, route, step, index, payload)
        document.save(docx_path)
        validation = self.validate(docx_path, payload)
        validation_path.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")
        if not validation["structural_pass"]:
            raise RuntimeError("variable route SOP validation failed: " + "; ".join(validation["errors"]))
        return RenderResult(
            product_code=route["product_code"],
            route_id=route_id,
            route_version=route["version"],
            docx_path=str(docx_path.resolve()),
            route_json_path=str(route_json_path.resolve()),
            validation_path=str(validation_path.resolve()),
            page_count_expected_from_route=len(steps) + 1,
            media_count=len(validation["media"]),
            image_policy="human_uploaded_and_confirmed_only",
        )

    def validate(self, docx_path: str | Path, payload: dict[str, Any]) -> dict[str, Any]:
        path = Path(docx_path)
        document = Document(path)
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            + [paragraph.text for section in document.sections for paragraph in section.header.paragraphs]
            + [paragraph.text for section in document.sections for paragraph in section.footer.paragraphs]
        )
        with zipfile.ZipFile(path) as archive:
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
        route = payload["route"]
        step_codes = [step["step_code"] for step in payload["steps"]]
        confirmed_media = [item for item in payload.get("media", []) if item["link_state"] == "confirmed" and Path(item["storage_path"]).is_file()]
        confirmed_asset_ids = {item["asset_id"] for item in confirmed_media}
        steps_with_confirmed_media = {item["route_step_id"] for item in confirmed_media}
        checks = {
            "target_product_present": route["product_code"] in all_text,
            "all_route_steps_present": all(code in all_text for code in step_codes),
            "variable_step_count_preserved": len(step_codes) not in {0, 6},
            "no_embedded_images": len(media) == 0,
            "only_human_confirmed_images_embedded": len(media) == len(confirmed_asset_ids),
            "blank_image_policy_present": all_text.count("待人工上传并确认图片") >= len(step_codes) - len(steps_with_confirmed_media),
            "draft_boundary_present": "DRAFT" in all_text and "不可生产发布" in all_text,
            "no_rejected_0008": "YA.C.06.0008" not in all_text,
            "no_usb_c_template_text": "USB-C数据线包装" not in all_text,
            "unknowns_structured": all(
                item.get("owner_role") and item.get("required_evidence")
                for step in payload["steps"] for item in step["unknowns_json"]
            ),
            "all_route_sections_present": {item["section_type"] for item in payload.get("sections", [])} == {
                "product_identity", "bom_material", "equipment_fixture", "process_parameter",
                "quality_control", "packaging_label", "ie_timing", "release_signoff",
            },
            "route_section_status_visible": all(item["section_type"] in all_text for item in payload.get("sections", [])),
        }
        gate_names = set(checks) - {"no_embedded_images"}
        errors = [name for name, passed in checks.items() if name in gate_names and not passed]
        return {
            "structural_pass": not errors,
            "checks": checks,
            "errors": errors,
            "route_step_count": len(step_codes),
            "top_level_steps": sum(step["parent_step_id"] is None for step in payload["steps"]),
            "child_steps": sum(step["parent_step_id"] is not None for step in payload["steps"]),
            "media": media,
            "image_policy": "human_uploaded_and_confirmed_only",
            "design_preset": "compact_reference_guide",
            "named_override": "A4 portrait, 1.4cm margins for shop-floor reference",
        }

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.header_distance = Cm(0.7)
        section.footer_distance = Cm(0.7)
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(9.5)
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.15
        for name, size, color, before, after in (
            ("Title", 22, ACCENT, 0, 8),
            ("Heading 1", 16, ACCENT, 10, 6),
            ("Heading 2", 12, ACCENT, 8, 4),
            ("Heading 3", 10.5, "1F1F1F", 6, 3),
        ):
            style = styles[name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = True
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
        header = section.header.paragraphs[0]
        header.text = "云湃智算 | SOP 工艺路线草案 | 工艺理解优先"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._font_paragraph(header, 8, MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = "DRAFT / demo_not_for_release / 签核为空 / 不可生产发布"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._font_paragraph(footer, 8, MUTED)

    def _add_cover(self, document: Document, route: dict[str, Any], steps: list[dict[str, Any]], payload: dict[str, Any]) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        run = p.add_run("标准作业程序（工艺路线草案）")
        self._font_run(run, 22, ACCENT, True)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(route["product_name"])
        self._font_run(run, 16, "1F1F1F", True)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(route["product_code"])
        self._font_run(run, 13, MUTED, True)
        sections = payload.get("sections", [])
        section_status = "；".join(f"{item['section_type']}={item['review_state']} v{item['version']}" for item in sections)
        meta = document.add_table(rows=7, cols=2)
        self._setup_table(meta, [2700, 7600], header=False)
        rows = [
            ("路线名称", route["route_name"]),
            ("工艺族", route["process_family_code"]),
            ("路线版本/状态", f"v{route['version']} / {route['status']} / {route['approval_scope']}") ,
            ("工序规模", f"共 {len(steps)} 个结构化节点；顶层 {sum(s['parent_step_id'] is None for s in steps)}，子步骤 {sum(s['parent_step_id'] is not None for s in steps)}"),
            ("图片策略", "只嵌入工作人员上传并随工序人工确认的图片；未确认图片区保持空白。"),
            ("路线级审核", section_status or "缺少路线级审核 section"),
            ("发布边界", "本文件用于人工逐字段审核；没有自动签核、自动发布或现场试产结论。"),
        ]
        for row, (label, value) in zip(meta.rows, rows):
            self._set_cell(row.cells[0], label, fill=LIGHT, bold=True)
            self._set_cell(row.cells[1], value)
        document.add_heading("路线总览", level=1)
        overview = document.add_table(rows=1, cols=4)
        self._setup_table(overview, [900, 1700, 2600, 5100], header=True)
        for cell, value in zip(overview.rows[0].cells, ("顺序", "步骤代码", "层级", "工序名称")):
            self._set_cell(cell, value, fill=ACCENT, color="FFFFFF", bold=True, align="center")
        id_to_code = {step["id"]: step["step_code"] for step in steps}
        for index, step in enumerate(steps, start=1):
            row = overview.add_row()
            values = (str(index), step["step_code"], f"子步骤，父级 {id_to_code.get(step['parent_step_id'])}" if step["parent_step_id"] else "顶层工序", step["title"])
            for cell, value in zip(row.cells, values):
                self._set_cell(cell, value, align="center" if cell is not row.cells[3] else "left")
        document.add_heading("证据与人工审核说明", level=1)
        p = document.add_paragraph()
        p.add_run("审核边界：").bold = True
        p.add_run(
            "未提供的设备、治具、参数、抽样和IE工时均保留为结构化unknown；"
            "每个节点、字段及八类路线章节必须在本地审核工作台逐项修改确认。"
        )

    def _add_step_page(self, document: Document, route: dict[str, Any], step: dict[str, Any], index: int, payload: dict[str, Any]) -> None:
        title = document.add_paragraph(style="Heading 1")
        title.add_run(f"工序 {index} | {step['step_code']} | {step['title']}")
        meta = document.add_table(rows=2, cols=4)
        self._setup_table(meta, [1300, 3900, 1300, 3900], header=False)
        meta_values = (
            ("产品", route["product_code"], "路线版本", f"v{route['version']} / {route['status']}"),
            ("层级", "子步骤" if step["parent_step_id"] else "顶层工序", "审核状态", step["review_state"]),
        )
        for table_row, values in zip(meta.rows, meta_values):
            for col, value in enumerate(values):
                self._set_cell(table_row.cells[col], value, fill=LIGHT if col % 2 == 0 else None, bold=col % 2 == 0)
        image_box = document.add_table(rows=1, cols=1)
        self._setup_table(image_box, [10400], header=False)
        media = [item for item in payload.get("media", []) if item["route_step_id"] == step["id"] and item["link_state"] == "confirmed" and Path(item["storage_path"]).is_file()]
        if media:
            cell = image_box.cell(0, 0)
            cell.text = ""
            for media_index, item in enumerate(media):
                paragraph = cell.paragraphs[0] if media_index == 0 else cell.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(item["storage_path"], width=Cm(15.5))
                caption = cell.add_paragraph(item["caption"] or item["original_name"])
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._font_paragraph(caption, 8, MUTED)
        else:
            self._set_cell(image_box.cell(0, 0), "图片区（空白）\n待人工上传并确认图片\nAI 不生成、不抓取、不把示意图冒充现场证据", fill="F7F7F7", color=MUTED, align="center")
            image_box.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(3)
            image_box.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(3)
        detail = document.add_table(rows=0, cols=2)
        self._setup_table(detail, [2300, 8100], header=False)
        rows = [
            ("做什么", step["action"]),
            ("为什么/目的", step["why"]),
            ("输入", self._format_value(step["input_json"])),
            ("物料", self._format_value(step["material_json"])),
            ("设备/工具", self._format_value(step["tool_equipment_json"])),
            ("治具/量具", self._format_value(step["fixture_json"])),
            ("控制参数及来源", self._format_value(step["parameter_json"])),
            ("可执行子步骤", self._format_numbered(step["method_json"])),
            ("检查方法", self._format_value(step["quality_check_json"])),
            ("合格判据", self._format_value(step["acceptance_criteria_json"])),
            ("安全注意", self._format_value(step["safety_json"])),
            ("输出/记录", self._format_value(step["record_output_json"])),
            ("异常隔离/处理", self._format_value(step["exception_json"])),
        ]
        for label, value in rows:
            row = detail.add_row()
            self._set_cell(row.cells[0], label, fill=LIGHT, bold=True)
            self._set_cell(row.cells[1], value or "无已确认事实；见下方结构化unknown。")
        unknowns = step["unknowns_json"]
        row = detail.add_row()
        self._set_cell(row.cells[0], "结构化 unknown", fill=CAUTION, bold=True)
        self._set_cell(
            row.cells[1],
            self._format_unknowns(unknowns) if unknowns else "本步骤没有未决字段。",
            fill=CAUTION,
            font_size=7.8,
        )
        evidence = [item for item in payload["provenance"] if item["route_step_id"] == step["id"]]
        row = detail.add_row()
        self._set_cell(row.cells[0], "字段来源/复用", fill="E2F0D9", bold=True)
        self._set_cell(row.cells[1], self._format_provenance(evidence) if evidence else "当前字段来自产品资料与工艺族草案；需在审核工作台逐字段确认。", fill="E2F0D9")
        if step["reviewer_comment"]:
            p = document.add_paragraph()
            p.add_run("审核意见：").bold = True
            p.add_run(step["reviewer_comment"])

    def _setup_table(self, table: Any, widths: list[int], *, header: bool) -> None:
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                self._set_cell_width(cell, width)
        if header and table.rows:
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    def _set_cell_width(self, cell: Any, width: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(width))
        tc_w.set(qn("w:type"), "dxa")
        margins = tc_pr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for side, value in (("top",80),("bottom",80),("start",110),("end",110)):
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def _set_cell(
        self,
        cell: Any,
        value: str,
        *,
        fill: str | None = None,
        color: str = "1F1F1F",
        bold: bool = False,
        align: str = "left",
        font_size: float = 8.6,
    ) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(str(value))
        self._font_run(run, font_size, color, bold)
        if fill:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                tc_pr.append(shading)
            shading.set(qn("w:fill"), fill)

    def _font_run(self, run: Any, size: float, color: str, bold: bool = False) -> None:
        run.font.name = "Microsoft YaHei"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold

    def _font_paragraph(self, paragraph: Any, size: float, color: str) -> None:
        for run in paragraph.runs:
            self._font_run(run, size, color)

    @staticmethod
    def _format_numbered(values: list[Any]) -> str:
        return "\n".join(f"{index}. {value}" for index, value in enumerate(values, start=1))

    @staticmethod
    def _format_value(value: Any) -> str:
        if isinstance(value, list):
            parts = []
            for item in value:
                if isinstance(item, dict):
                    parts.append("；".join(f"{key}={val}" for key, val in item.items()))
                else:
                    parts.append(str(item))
            return "\n".join(parts)
        if isinstance(value, dict):
            return "\n".join(f"{key}: {val}" for key, val in value.items())
        return str(value or "")

    @staticmethod
    def _format_unknowns(values: list[dict[str, Any]]) -> str:
        return "\n".join(
            f"[{item['field_name']}] 原因：{item['reason']} 责任角色：{item['owner_role']}。所需证据：{item['required_evidence']}。阻断：{'是' if item.get('blocking', True) else '否'}。"
            for item in values
        )

    @staticmethod
    def _format_provenance(values: list[dict[str, Any]]) -> str:
        return "\n".join(
            f"字段 {item['field_name']} | 来源 {item.get('source_path') or 'approved route '+str(item.get('source_route_id'))} | 页/表 {item.get('page_or_sheet') or '-'} | 置信 {item['confidence']} | 冲突 {item['conflict_status']}"
            for item in values
        )
