from __future__ import annotations

import argparse
import hashlib
import json
import sys
from datetime import date
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document, Document as WordDocument
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Pt, Twips

from cad_ai.sop_agent import SOP_GENERATION_SEQUENCE, _validate_docx_package
from cad_ai.sop_visual_template import (
    SOP_FLOWCHART_SHAPE_POLICY,
    _build_sop_word_document,
    _configure_word_section,
    _default_process_ie_time_study,
    _default_work_ie_time_study,
    _demo_step,
    _render_center_flowchart_shape_image,
    _render_process_flow_word_table,
    _render_work_instruction_word_table,
    _set_document_defaults,
    _set_row_height,
    _set_word_cell,
    _set_word_cell_margins,
    _write_word_format_check_json,
    build_process_flow_page,
    build_usb_cable_packaging_demo,
    build_work_instruction_page,
)


TEMPLATE_ID = "yunpai.sop.usb_c_cable_packaging.two_page.v1"
FINAL_DOCX_NAME = "SOP完整模板_USB-C数据线包装_草案.docx"
HDMI_TEMPLATE_ID = "yunpai.sop.hdmi-cable.multi-page.v1"
HDMI_FINAL_DOCX_NAME = "SOP完整模板_HDMI线制作_草案.docx"
CENTER_FLOWCHART_NAME = "center_flowchart.png"
MANIFEST_NAME = "sop_template_manifest.json"
FORMAT_CHECK_NAME = "sop_template_format_check.json"
VALIDATION_NAME = "sop_template_validation.json"
CONTENT_PROFILE_USB_C = "usb-c-packaging"
CONTENT_PROFILE_HDMI = "hdmi-cable"
CONTENT_PROFILES = (CONTENT_PROFILE_USB_C, CONTENT_PROFILE_HDMI)
HANDOFF_GENERATION_SEQUENCE = [
    *SOP_GENERATION_SEQUENCE[:-1],
    "apply_two_page_delivery_controls",
    SOP_GENERATION_SEQUENCE[-1],
]


def generate_package(
    out_dir: str | Path,
    *,
    document_date: str,
    content_profile: str = CONTENT_PROFILE_USB_C,
) -> dict[str, Any]:
    if content_profile == CONTENT_PROFILE_HDMI:
        raise ValueError(
            "HDMI SOP must use route-backed multi-page mode: provide --route-db and --route-id"
        )
    normalized_date = _normalize_date(document_date)
    display_date = normalized_date.replace("-", "/").lstrip("0").replace("/0", "/")
    profile = _resolve_content_profile(content_profile)
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    expected_names = {
        profile["document_name"],
        CENTER_FLOWCHART_NAME,
        MANIFEST_NAME,
        FORMAT_CHECK_NAME,
        VALIDATION_NAME,
    }
    unexpected_names = {path.name for path in output.iterdir()} - expected_names
    if unexpected_names:
        raise RuntimeError(
            "Output directory contains files outside the handoff contract: "
            + ", ".join(sorted(unexpected_names))
        )

    demo = profile["builder"]()
    flow_page = demo["process_flow"]
    work_page = demo["work_instruction"]
    flow_page["render_center_flowchart"] = True
    flow_page["center_flowchart_style"] = "pdf_reference_shape_blocks"

    center_flowchart = output / CENTER_FLOWCHART_NAME
    center_flowchart.write_bytes(_render_center_flowchart_shape_image(flow_page))

    document = _build_sop_word_document(flow_page, work_page)
    _apply_delivery_controls(document, normalized_date=normalized_date, display_date=display_date)
    document_path = output / profile["document_name"]
    document.save(document_path)

    format_check_path = output / FORMAT_CHECK_NAME
    _write_word_format_check_json(format_check_path, flow_page, work_page)

    validation_path = output / VALIDATION_NAME
    validation = validate_document(
        document_path,
        expected_date=display_date,
        template_id=profile["template_id"],
    )
    validation_path.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest_path = output / MANIFEST_NAME
    manifest = _build_manifest(
        output=output,
        normalized_date=normalized_date,
        document_path=document_path,
        center_flowchart=center_flowchart,
        format_check_path=format_check_path,
        validation_path=validation_path,
        validation=validation,
        profile=profile,
    )
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    if not validation["structural_pass"]:
        raise RuntimeError("SOP template structural validation failed: " + "; ".join(validation["errors"]))

    return {
        "template_id": profile["template_id"],
        "status": "demo_not_for_release",
        "document_docx": str(document_path.resolve()),
        "center_flowchart_png": str(center_flowchart.resolve()),
        "manifest_json": str(manifest_path.resolve()),
        "format_check_json": str(format_check_path.resolve()),
        "validation_json": str(validation_path.resolve()),
        "structural_pass": True,
        "visual_qa_required": True,
    }


def generate_route_package(
    out_dir: str | Path,
    *,
    document_date: str,
    db_path: str | Path,
    route_id: int,
) -> dict[str, Any]:
    """Generate page 1 flowchart plus one repeated instruction page per route step."""
    from cad_ai.sop_knowledge.store import SopKnowledgeStore

    normalized_date = _normalize_date(document_date)
    display_date = normalized_date.replace("-", "/").lstrip("0").replace("/0", "/")
    output = Path(out_dir)
    output.mkdir(parents=True, exist_ok=True)
    expected_names = {
        HDMI_FINAL_DOCX_NAME, CENTER_FLOWCHART_NAME, MANIFEST_NAME,
        FORMAT_CHECK_NAME, VALIDATION_NAME,
    }
    unexpected_names = {path.name for path in output.iterdir()} - expected_names
    if unexpected_names:
        raise RuntimeError(
            "Output directory contains files outside the handoff contract: "
            + ", ".join(sorted(unexpected_names))
        )

    store = SopKnowledgeStore(db_path)
    store.initialize()
    route = store.get_route(route_id)
    flow_page, work_pages = _route_template_pages(route, normalized_date=normalized_date)
    flow_page["render_center_flowchart"] = True
    flow_page["center_flowchart_style"] = "pdf_reference_shape_blocks"

    center_flowchart = output / CENTER_FLOWCHART_NAME
    center_flowchart.write_bytes(_render_center_flowchart_shape_image(flow_page))

    document = _build_multi_page_document(flow_page, work_pages)
    _apply_multi_page_delivery_controls(
        document,
        instruction_page_count=len(work_pages),
        normalized_date=normalized_date,
        display_date=display_date,
    )
    document_path = output / HDMI_FINAL_DOCX_NAME
    document.save(document_path)

    validation = validate_multi_page_document(
        document_path,
        expected_date=display_date,
        expected_instruction_pages=len(work_pages),
    )
    validation_path = output / VALIDATION_NAME
    validation_path.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")

    format_check = {
        "template_id": HDMI_TEMPLATE_ID,
        "layout_mode": "portrait_flow_then_repeated_landscape_work_instructions",
        "process_flow_pages": 1,
        "work_instruction_pages": len(work_pages),
        "expected_rendered_pages": 1 + len(work_pages),
        "tables_per_flow_page": 4,
        "tables_per_instruction_page": 4,
        "visual_step_order_each_instruction_page": "1,2,3 / 6,5,4",
        "images": "blank_until_human_upload_and_confirmation",
        "status": "demo_not_for_release",
    }
    format_check_path = output / FORMAT_CHECK_NAME
    format_check_path.write_text(json.dumps(format_check, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest = {
        "schema_version": "1.1",
        "template_id": HDMI_TEMPLATE_ID,
        "status": "demo_not_for_release",
        "document_date": normalized_date,
        "single_allowed_entrypoint": "python scripts/generate_sop_template_ai_handoff.py --route-db ... --route-id ...",
        "layout": {
            "first_page": "A4 portrait process flowchart",
            "following_pages": "repeated A4 landscape standard work instruction",
            "work_instruction_pages": len(work_pages),
            "expected_rendered_pages": 1 + len(work_pages),
            "step_order_each_page": "1,2,3 / 6,5,4",
        },
        "route": {
            "route_id": route_id,
            "product_code": route["route"]["product_code"],
            "route_version": route["route"]["version"],
            "route_status": route["route"]["status"],
            "instruction_step_ids": [item["id"] for item in route["steps"]],
        },
        "artifacts": {
            "document_docx": _artifact_record(document_path, output),
            "center_flowchart_png": _artifact_record(center_flowchart, output),
            "format_check_json": _artifact_record(format_check_path, output),
            "validation_json": _artifact_record(validation_path, output),
        },
        "validation": {
            "structural_pass": validation["structural_pass"],
            "visual_qa_required": True,
            "expected_rendered_pages": 1 + len(work_pages),
        },
        "guardrails": {
            "no_auto_release": True,
            "no_auto_signoff": True,
            "no_fabricated_site_photos": True,
            "no_fabricated_process_parameters_or_ie_measurement": True,
        },
    }
    manifest_path = output / MANIFEST_NAME
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    if not validation["structural_pass"]:
        raise RuntimeError("SOP multi-page template structural validation failed: " + "; ".join(validation["errors"]))
    return {
        "template_id": HDMI_TEMPLATE_ID,
        "status": "demo_not_for_release",
        "route_id": route_id,
        "document_docx": str(document_path.resolve()),
        "center_flowchart_png": str(center_flowchart.resolve()),
        "manifest_json": str(manifest_path.resolve()),
        "format_check_json": str(format_check_path.resolve()),
        "validation_json": str(validation_path.resolve()),
        "instruction_page_count": len(work_pages),
        "expected_page_count": 1 + len(work_pages),
        "structural_pass": True,
        "visual_qa_required": True,
    }
def validate_document(
    document_path: str | Path,
    *,
    expected_date: str | None = None,
    template_id: str = TEMPLATE_ID,
) -> dict[str, Any]:
    path = Path(document_path)
    errors: list[str] = []
    if not path.exists():
        return {
            "template_id": template_id,
            "document": str(path),
            "structural_pass": False,
            "errors": ["document_not_found"],
        }

    validation = _validate_docx_package(path)
    document = Document(path)
    tables = document.tables
    signoff_values: list[str] = []
    if len(tables) >= 8:
        signoff_values = [tables[7].cell(1, index).text.strip() for index in range(3)]

    expected_checks = {
        "sections": validation.get("sections") == 2,
        "top_level_tables": validation.get("top_level_tables") == 8,
        "has_png_media": validation.get("has_png_media") is True,
        "has_no_svg": validation.get("has_svg") is False,
        "has_no_vml": validation.get("has_vml_shape") is False,
        "has_no_replacement_char": validation.get("contains_replacement_char") is False,
        "contains_ie_time_title": validation.get("contains_ie_time_title") is True,
        "contains_machine_model_field": validation.get("contains_machine_model_field") is True,
        "contains_demo_boundary": validation.get("contains_demo_boundary") is True,
        "portrait_flow_section": len(document.sections) >= 1
        and document.sections[0].orientation == WD_ORIENT.PORTRAIT,
        "landscape_instruction_section": len(document.sections) >= 2
        and document.sections[1].orientation == WD_ORIENT.LANDSCAPE,
        "blank_approval_audit_author_cells": signoff_values == ["", "", ""],
        "version_is_draft": len(tables) >= 5
        and tables[0].cell(2, 3).text.strip() == "DRAFT"
        and tables[4].cell(2, 3).text.strip() == "DRAFT",
        "visual_step_order": _has_visual_step_order(tables),
        "required_sections_present": _has_required_sections(tables),
    }
    if expected_date is not None:
        expected_checks["document_date"] = (
            len(tables) >= 5
            and tables[0].cell(2, 5).text.strip() == expected_date
            and tables[4].cell(1, 7).text.strip() == expected_date
        )

    for check_name, passed in expected_checks.items():
        if not passed:
            errors.append(check_name)

    return {
        "template_id": template_id,
        "document": str(path.resolve()),
        "sha256": _sha256(path),
        "size_bytes": path.stat().st_size,
        "structural_pass": not errors,
        "checks": expected_checks,
        "module_validation": validation,
        "errors": errors,
        "visual_qa": {
            "required": True,
            "expected_page_count": 2,
            "status": "pending_external_render_and_human_or_agent_image_review",
        },
    }


def _apply_delivery_controls(document: Any, *, normalized_date: str, display_date: str) -> None:
    if len(document.sections) != 2 or len(document.tables) != 8:
        raise ValueError("Unexpected SOP base layout; expected 2 sections and 8 top-level tables")

    tables = document.tables
    landscape = document.sections[1]
    landscape.top_margin = Cm(0.3)
    landscape.bottom_margin = Cm(0.3)
    landscape.left_margin = Cm(0.6)
    landscape.right_margin = Cm(0.6)

    _set_word_cell(tables[0].cell(2, 3), "DRAFT")
    _set_word_cell(tables[0].cell(2, 5), display_date)
    _set_word_cell(tables[4].cell(1, 7), display_date)
    _set_word_cell(tables[4].cell(2, 3), "DRAFT")
    _set_word_cell(
        tables[5].cell(4, 4),
        f"变更内容\nA / 初版模板 / {normalized_date} / M2",
        size=7,
        align=0,
    )

    for row_index, height in enumerate([620, 620, 420, 620, 620, 420]):
        _set_row_height(tables[5].rows[row_index], height)
    for row_index in range(len(tables[6].rows)):
        _set_row_height(tables[6].rows[row_index], 220 if row_index < 2 else 180)
    _set_row_height(tables[7].rows[0], 280)
    _set_row_height(tables[7].rows[1], 420)

    for row_index in range(6):
        side_cell = tables[5].cell(row_index, 4)
        _set_word_cell_margins(side_cell, top=30, start=70, bottom=30, end=70)
        for paragraph in side_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)

    _set_word_cell(tables[7].cell(1, 3), "材料符合RoHS/REACH；发布前确认。", size=7, align=0)
    _set_word_cell_margins(tables[7].cell(1, 3), top=30, start=60, bottom=30, end=60)


def _build_manifest(
    *,
    output: Path,
    normalized_date: str,
    document_path: Path,
    center_flowchart: Path,
    format_check_path: Path,
    validation_path: Path,
    validation: dict[str, Any],
    profile: dict[str, Any],
) -> dict[str, Any]:
    artifacts = {
        "document_docx": _artifact_record(document_path, output),
        "center_flowchart_png": _artifact_record(center_flowchart, output),
        "format_check_json": _artifact_record(format_check_path, output),
        "validation_json": _artifact_record(validation_path, output),
    }
    return {
        "schema_version": "1.0",
        "template_id": profile["template_id"],
        "status": "demo_not_for_release",
        "document_date": normalized_date,
        "single_allowed_entrypoint": "python scripts/generate_sop_template_ai_handoff.py",
        "source_modules": [
            "cad_ai/sop_visual_template.py",
            "cad_ai/sop_agent.py",
        ],
        "generation_sequence": list(HANDOFF_GENERATION_SEQUENCE),
        "tables_filled_before_flowchart": True,
        "center_flowchart_target": "process_flow_body_table_cell_0_0",
        "shape_policy": SOP_FLOWCHART_SHAPE_POLICY,
        "fixed_template_profile": {
            "product": profile["product"],
            "content_profile": profile["name"],
            "flow_section": "A4 portrait",
            "work_instruction_section": "A4 landscape",
            "expected_rendered_pages": 2,
            "top_level_word_tables": 8,
            "step_order": "1,2,3 / 6,5,4",
        },
        "artifacts": artifacts,
        "validation": {
            "structural_pass": validation["structural_pass"],
            "visual_qa_required": True,
            "expected_rendered_pages": 2,
        },
        "guardrails": {
            "no_auto_release": True,
            "no_auto_signoff": True,
            "no_fabricated_site_photos": True,
            "no_fabricated_ie_measurement": True,
            "no_fabricated_equipment_status": True,
            "no_fabricated_ehs_oee_yield_trial_or_training_records": True,
        },
    }


def build_hdmi_cable_demo() -> dict[str, Any]:
    product_name = "HDMI 成品线（草案）"
    part_no = "DRAFT-HDMI-001"
    document_no = "SOP-HDMI-DRAFT-01"
    drawing_no = "SOP-HDMI-FMT-01"
    operations = ["物料与文件核对", "裁线与剥外被", "屏蔽及芯线整理", "端接作业", "外壳装配", "电气与外观检验"]
    process_flow = build_process_flow_page(
        product_name=product_name,
        part_no=part_no,
        document_no=document_no,
        drawing_no=drawing_no,
        operations=operations,
    )
    for node, node_type in zip(
        process_flow["flow_nodes"],
        ["inspection", "process", "process", "process", "process", "test"],
    ):
        node["type"] = node_type
        node["shape"] = node_type
    process_flow["center_flowchart"] = {"flowchart_title": "HDMI线制作与检验"}
    process_flow["ie_time_study"] = _default_process_ie_time_study(process_flow["flow_nodes"])
    for row in process_flow["ie_time_study"]["rows"]:
        row["工时来源"] = "demo_not_for_release / 待IE实测"

    work_instruction = build_work_instruction_page(
        product_name=product_name,
        part_no=part_no,
        station="组装与电测",
        document_no=document_no,
        drawing_no=drawing_no,
        version="DRAFT",
        page_no=1,
        page_total=1,
    )
    work_instruction["operation_order"] = "①备料==②裁剥线==③芯线整理==④端接==⑤装壳==⑥电气/外观检验"
    work_instruction["step_slots"] = [
        _demo_step(1, "物料核对", "按工单、BOM及受控图纸核对线材、连接器与辅料。", "materials"),
        _demo_step(2, "裁剥线", "按受控工艺卡裁线并剥外被；尺寸与设备参数待工程确认。", "coil"),
        _demo_step(3, "芯线整理", "按图纸整理屏蔽层、地线及芯线，保持线对结构且不得伤线。", "inspect"),
        _demo_step(4, "端接作业", "按受控接线图完成端接并逐点检查；方式与参数待确认。", "generic"),
        _demo_step(5, "外壳装配", "装配绝缘、屏蔽、外壳及应力消除件；成型条件待确认。", "tie"),
        _demo_step(6, "电气/外观检验", "按检验规范检查导通、短路、屏蔽与外观；限值待质量确认。", "inspect"),
    ]
    work_instruction["ie_time_study"] = _default_work_ie_time_study(work_instruction["step_slots"])
    for row in work_instruction["ie_time_study"]["rows"]:
        row["工时来源"] = "demo_not_for_release / 待IE实测"
    work_instruction["side_sections"] = [
        {
            "title": "作业标准",
            "lines": [
                "1. 脚位、线序及屏蔽结构以受控图纸为准。",
                "2. 未确认尺寸、参数及测试限值不得自行填写。",
                "3. 检验结论须由授权人员按规范判定。",
            ],
        },
        {
            "title": "设备/工具",
            "lines": ["1. 裁剥线设备（型号待确认）", "2. 端接工具（方式/型号待确认）", "3. HDMI测试仪（程序/限值待确认）"],
        },
        {"title": "辅助材料", "lines": ["热缩/绝缘辅材（以BOM为准）", "端接辅料（工艺路线待确认）", "外壳及应力消除件（以BOM为准）"]},
        {
            "title": "注意事项",
            "lines": [
                "1. ESD、高温及机械防护要求由现场确认。",
                "2. 不得损伤芯线、屏蔽层或连接器触点。",
                "3. 异常品隔离标识并提交人工评审。",
            ],
        },
        {"title": "变更内容", "lines": ["A / HDMI线SOP初版草案 / 待受控日期 / M2"]},
        {
            "title": "物料表",
            "lines": [
                "1 / 待编码 / HDMI线材 / 规格与用量待确认",
                "2 / 待编码 / HDMI连接器及外壳 / BOM待确认",
                "3 / 待编码 / 绝缘与屏蔽辅料 / BOM待确认",
            ],
        },
    ]
    return {
        "product_name": product_name,
        "part_no": part_no,
        "image_source": "local_generated_raster",
        "operations": operations,
        "process_flow": process_flow,
        "work_instruction": work_instruction,
    }


def _route_template_pages(
    payload: dict[str, Any],
    *,
    normalized_date: str,
) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    route = payload["route"]
    steps = list(payload.get("steps") or [])
    route_product_name = str(route.get("product_name") or "HDMI 成品线")
    product_name = "HDMI 成品线（草案）" if "HDMI" in route_product_name.upper() else route_product_name
    part_no = str(route.get("product_code") or "DRAFT-HDMI-001")
    version = str(route.get("version") or "DRAFT")
    document_no = f"SOP-{part_no}-V{version}"
    drawing_no = f"SOP-{part_no}-FMT"
    operations = [str(item.get("title") or item.get("step_code") or "工序") for item in steps]
    flow_page = build_process_flow_page(
        product_name=product_name,
        part_no=part_no,
        document_no=document_no,
        drawing_no=drawing_no,
        operations=operations,
    )
    flow_page["flow_nodes"] = [
        {
            "seq": index,
            "name": str(step.get("title") or step.get("step_code") or f"工序{index}"),
            "step_code": str(step.get("step_code") or ""),
            "type": _route_step_node_type(step),
            "shape": _route_step_node_type(step),
        }
        for index, step in enumerate(steps, start=1)
    ]
    flow_page["center_flowchart"] = {
        "flowchart_title": f"{product_name}完整制造流程",
        "nodes": flow_page["flow_nodes"],
    }
    flow_page["ie_time_study"] = _default_process_ie_time_study(flow_page["flow_nodes"])
    for row in flow_page["ie_time_study"]["rows"]:
        row["工时来源"] = "待IE实测/人工锁定"

    page_total = len(steps)
    work_pages: list[dict[str, Any]] = []
    for page_no, step in enumerate(steps, start=1):
        methods = [str(item).strip() for item in step.get("method_json", []) if str(item).strip()]
        slots = []
        for slot_no in range(1, 7):
            method = methods[slot_no - 1] if slot_no <= len(methods) else ""
            slots.append({
                "slot_no": slot_no,
                "image_placeholder": True,
                "image_label": "图片区（待人工上传确认）",
                "text_placeholder": f"{slot_no}. {method}" if method else f"{slot_no}. ",
            })
        page = build_work_instruction_page(
            product_name=product_name,
            part_no=part_no,
            station=str(step.get("title") or step.get("step_code") or "工序"),
            document_no=document_no,
            drawing_no=drawing_no,
            version="DRAFT",
            page_no=page_no,
            page_total=page_total,
        )
        page["operation_order"] = (
            f"{step.get('step_code', '')} | {step.get('title', '')} | "
            f"{step.get('action', '')}"
        )
        page["step_slots"] = slots
        page["ie_time_study"] = _default_work_ie_time_study(slots)
        for row in page["ie_time_study"]["rows"]:
            row["机器型号"] = "待工程确认"
            row["IE测量方法"] = "待IE现场实测"
            row["工时来源"] = "待IE实测/人工锁定"
        quality_lines = _compact_lines(
            list(step.get("quality_check_json") or []) + list(step.get("acceptance_criteria_json") or []),
            fallback="检查方法与合格判据待质量人员确认。",
        )
        equipment_lines = _compact_lines(
            list(step.get("tool_equipment_json") or []) + list(step.get("fixture_json") or []),
            fallback="设备、工具及治具型号待工程确认。",
        )
        material_lines = _compact_lines(
            list(step.get("material_json") or []),
            fallback="材料以受控BOM为准。",
        )
        caution_lines = _compact_lines(
            list(step.get("safety_json") or []) + list(step.get("exception_json") or []),
            fallback="异常时停止流转、隔离并提交人工判定。",
        )
        input_lines = _compact_lines(
            list(step.get("input_json") or []) + list(step.get("record_output_json") or []),
            fallback="输入与记录要求待责任人确认。",
        )
        page["side_sections"] = [
            {"title": "作业标准", "lines": quality_lines},
            {"title": "设备/工具", "lines": equipment_lines},
            {"title": "辅助材料", "lines": material_lines},
            {"title": "注意事项", "lines": caution_lines},
            {"title": "变更内容", "lines": [f"DRAFT / 路线v{version} / {normalized_date} / 待人工复核"]},
            {"title": "物料表", "lines": input_lines},
        ]
        work_pages.append(page)
    return flow_page, work_pages


def _route_step_node_type(step: dict[str, Any]) -> str:
    text = " ".join([
        str(step.get("title") or ""),
        " ".join(str(item) for item in step.get("quality_check_json", [])),
    ])
    return "test" if any(token in text for token in ("检", "测", "核对", "复核", "判定")) else "process"


def _compact_lines(values: list[Any], *, fallback: str, limit: int = 3) -> list[str]:
    lines = [str(item).strip() for item in values if str(item).strip()]
    return lines[:limit] or [fallback]


def _set_exact_row_height(row: Any, height_twips: int) -> None:
    row.height = Twips(height_twips)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _build_multi_page_document(
    flow_page: dict[str, Any],
    work_pages: list[dict[str, Any]],
) -> Any:
    if not work_pages:
        raise ValueError("At least one work instruction page is required")
    document = WordDocument()
    _set_document_defaults(document)
    _configure_word_section(document.sections[0], landscape=False)
    _render_process_flow_word_table(document, flow_page)
    _configure_word_section(document.add_section(WD_SECTION.NEW_PAGE), landscape=True)
    for index, page in enumerate(work_pages):
        if index:
            spacer = document.add_paragraph()
            spacer.paragraph_format.page_break_before = True
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(1)
        _render_work_instruction_word_table(document, page)
    return document


def _apply_multi_page_delivery_controls(
    document: Any,
    *,
    instruction_page_count: int,
    normalized_date: str,
    display_date: str,
) -> None:
    expected_tables = 4 + instruction_page_count * 4
    if len(document.sections) != 2 or len(document.tables) != expected_tables:
        raise ValueError(
            f"Unexpected SOP multi-page layout; expected 2 sections and {expected_tables} top-level tables"
        )
    tables = document.tables
    landscape = document.sections[1]
    landscape.top_margin = Cm(0.3)
    landscape.bottom_margin = Cm(0.3)
    landscape.left_margin = Cm(0.6)
    landscape.right_margin = Cm(0.6)
    _set_word_cell(tables[0].cell(2, 3), "DRAFT")
    _set_word_cell(tables[0].cell(2, 5), display_date)
    for page_index in range(instruction_page_count):
        base = 4 + page_index * 4
        header, body, ie_table, footer = tables[base : base + 4]
        for row, height in zip(header.rows, [520, 420, 600]):
            _set_exact_row_height(row, height)
        _set_word_cell(header.cell(1, 7), display_date)
        _set_word_cell(header.cell(2, 3), "DRAFT")
        _set_word_cell(
            body.cell(4, 4),
            f"变更内容\nDRAFT / 路线草稿 / {normalized_date} / 待人工复核",
            size=7,
            align=0,
        )
        for row_index, height in enumerate([900, 900, 500, 900, 900, 500]):
            _set_exact_row_height(body.rows[row_index], height)
        for row_index in range(len(ie_table.rows)):
            _set_exact_row_height(ie_table.rows[row_index], 260 if row_index < 2 else 220)
        _set_exact_row_height(footer.rows[0], 360)
        _set_exact_row_height(footer.rows[1], 560)
        for row_index in range(6):
            side_cell = body.cell(row_index, 4)
            _set_word_cell_margins(side_cell, top=30, start=70, bottom=30, end=70)
            for paragraph in side_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)
        _set_word_cell(footer.cell(1, 3), "材料符合RoHS/REACH；发布前确认。", size=7, align=0)
        _set_word_cell_margins(footer.cell(1, 3), top=30, start=60, bottom=30, end=60)


def validate_multi_page_document(
    document_path: str | Path,
    *,
    expected_date: str,
    expected_instruction_pages: int,
) -> dict[str, Any]:
    path = Path(document_path)
    validation = _validate_docx_package(path)
    document = Document(path)
    tables = document.tables
    expected_tables = 4 + expected_instruction_pages * 4
    signoffs_blank = True
    draft_headers = True
    dates_match = True
    visual_orders = True
    for page_index in range(expected_instruction_pages):
        base = 4 + page_index * 4
        if base + 3 >= len(tables):
            signoffs_blank = draft_headers = dates_match = visual_orders = False
            break
        header, body, _, footer = tables[base : base + 4]
        signoffs_blank = signoffs_blank and [footer.cell(1, index).text.strip() for index in range(3)] == ["", "", ""]
        draft_headers = draft_headers and header.cell(2, 3).text.strip() == "DRAFT"
        dates_match = dates_match and header.cell(1, 7).text.strip() == expected_date
        expected = [(0, 1, "1"), (0, 2, "2"), (0, 3, "3"), (3, 1, "6"), (3, 2, "5"), (3, 3, "4")]
        visual_orders = visual_orders and all(body.cell(row, col).text.strip().startswith(prefix) for row, col, prefix in expected)
    checks = {
        "sections": len(document.sections) == 2,
        "top_level_tables": len(tables) == expected_tables,
        "portrait_flow_section": document.sections[0].orientation == WD_ORIENT.PORTRAIT,
        "landscape_instruction_section": document.sections[1].orientation == WD_ORIENT.LANDSCAPE,
        "instruction_page_count": expected_instruction_pages > 1,
        "blank_signoff_cells_every_page": signoffs_blank,
        "version_is_draft_every_page": draft_headers,
        "document_date_every_page": dates_match,
        "visual_step_order_every_page": visual_orders,
        "has_png_media": validation.get("has_png_media") is True,
        "has_no_svg": validation.get("has_svg") is False,
        "has_no_vml": validation.get("has_vml_shape") is False,
        "has_no_replacement_char": validation.get("contains_replacement_char") is False,
    }
    errors = [name for name, passed in checks.items() if not passed]
    return {
        "template_id": HDMI_TEMPLATE_ID,
        "document": str(path.resolve()),
        "sha256": _sha256(path),
        "size_bytes": path.stat().st_size,
        "structural_pass": not errors,
        "checks": checks,
        "module_validation": validation,
        "errors": errors,
        "visual_qa": {
            "required": True,
            "expected_page_count": 1 + expected_instruction_pages,
            "status": "pending_external_render_and_human_or_agent_image_review",
        },
    }


def _resolve_content_profile(content_profile: str) -> dict[str, Any]:
    profiles = {
        CONTENT_PROFILE_USB_C: {
            "name": CONTENT_PROFILE_USB_C,
            "template_id": TEMPLATE_ID,
            "document_name": FINAL_DOCX_NAME,
            "product": "USB-C 简易数据线",
            "builder": build_usb_cable_packaging_demo,
        },
        CONTENT_PROFILE_HDMI: {
            "name": CONTENT_PROFILE_HDMI,
            "template_id": HDMI_TEMPLATE_ID,
            "document_name": HDMI_FINAL_DOCX_NAME,
            "product": "HDMI 成品线（草案）",
            "builder": build_hdmi_cable_demo,
        },
    }
    try:
        return profiles[content_profile]
    except KeyError as exc:
        raise ValueError(f"Unsupported content profile: {content_profile}") from exc


def _has_visual_step_order(tables: list[Any]) -> bool:
    if len(tables) < 6:
        return False
    body = tables[5]
    expected = [
        (0, 1, "1"),
        (0, 2, "2"),
        (0, 3, "3"),
        (3, 1, "6"),
        (3, 2, "5"),
        (3, 3, "4"),
    ]
    return all(body.cell(row, column).text.strip().startswith(prefix) for row, column, prefix in expected)


def _has_required_sections(tables: list[Any]) -> bool:
    if len(tables) < 8:
        return False
    text = "\n".join(cell.text for table in tables for row in table.rows for cell in row.cells)
    required = [
        "流程图",
        "标准作业指导书",
        "IE工时记录",
        "作业标准",
        "设备/工具",
        "辅助材料",
        "注意事项",
        "变更内容",
        "物料表",
        "批准",
        "审核",
        "制作",
        "材料环保要求",
        "管制文件（印章处）",
    ]
    return all(token in text for token in required)


def _artifact_record(path: Path, root: Path) -> dict[str, Any]:
    return {
        "path": path.relative_to(root).as_posix(),
        "sha256": _sha256(path),
        "size_bytes": path.stat().st_size,
    }


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize_date(value: str) -> str:
    try:
        return date.fromisoformat(value).isoformat()
    except ValueError as exc:
        raise argparse.ArgumentTypeError("document date must use YYYY-MM-DD") from exc


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate the governed SOP template, including route-backed multi-page work instructions."
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=Path("outputs/sop_template_ai_handoff"),
        help="Output directory. Defaults to outputs/sop_template_ai_handoff.",
    )
    parser.add_argument(
        "--document-date",
        default=date.today().isoformat(),
        help="Controlled document date in YYYY-MM-DD form.",
    )
    parser.add_argument(
        "--content-profile",
        choices=CONTENT_PROFILES,
        default=CONTENT_PROFILE_USB_C,
        help="Explicit content variant. The frozen USB-C profile remains the default.",
    )
    parser.add_argument(
        "--check-only",
        type=Path,
        help="Validate an existing DOCX instead of generating a new package.",
    )
    parser.add_argument("--route-db", type=Path, help="SOP knowledge SQLite database for multi-page route mode.")
    parser.add_argument("--route-id", type=int, help="Route id for multi-page route mode.")
    args = parser.parse_args()

    if args.check_only:
        profile = _resolve_content_profile(args.content_profile)
        result = validate_document(args.check_only, template_id=profile["template_id"])
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result["structural_pass"] else 1

    if (args.route_db is None) != (args.route_id is None):
        parser.error("--route-db and --route-id must be supplied together")
    if args.route_db is not None:
        result = generate_route_package(
            args.out_dir,
            document_date=args.document_date,
            db_path=args.route_db,
            route_id=args.route_id,
        )
    else:
        result = generate_package(
            args.out_dir,
            document_date=args.document_date,
            content_profile=args.content_profile,
        )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
