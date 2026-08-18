"""Capture the local SOP UI and build a plain-language worker manual."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from playwright.sync_api import sync_playwright


DEFAULT_URL = "http://127.0.0.1:8787/"
DEFAULT_OUTPUT = Path("outputs/worker_function_manual")
QUESTION = "当前路线共有多少道工序？请只回答，不要修改 SOP。"
CANDIDATE_QUESTION = "把检查工序的作业步骤补充完整。"


def _edge_executable() -> str:
    candidates = (
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    )
    for path in candidates:
        if path.is_file():
            return str(path)
    raise FileNotFoundError("Microsoft Edge was not found")


def capture_screenshots(base_url: str, output_dir: Path) -> tuple[dict[str, Path], dict[str, Any]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    shots = {
        "home": output_dir / "01-首页与DOCX预览.png",
        "question": output_dir / "02-输入问题.png",
        "answer": output_dir / "03-AI只回答.png",
        "candidate": output_dir / "04-选择目标工序.png",
        "route_editor": output_dir / "05-路线编辑.png",
        "route_menu": output_dir / "06-工序操作菜单.png",
    }
    console_issues: list[str] = []
    page_errors: list[str] = []
    responses: list[str] = []

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(executable_path=_edge_executable(), headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 900}, device_scale_factor=1)
        page.on("console", lambda message: console_issues.append(f"{message.type}: {message.text}") if message.type == "error" else None)
        page.on("pageerror", lambda error: page_errors.append(str(error)))
        page.on("response", lambda response: responses.append(f"{response.status} {response.url}") if response.status >= 400 else None)

        page.goto(base_url, wait_until="domcontentloaded", timeout=30_000)
        page.wait_for_selector("#product option", state="attached", timeout=60_000)
        page.wait_for_selector("#prompt", timeout=30_000)
        page.wait_for_timeout(1_500)
        page.screenshot(path=str(shots["home"]), full_page=True)

        page.locator("#worker").fill("示例操作员")
        page.locator("#prompt").fill(QUESTION)
        page.screenshot(path=str(shots["question"]), full_page=True)

        previous = page.locator(".message.assistant").count()
        page.locator("#send").click()
        page.wait_for_function(
            "previous => document.querySelectorAll('.message.assistant').length > previous",
            arg=previous,
            timeout=180_000,
        )
        page.wait_for_timeout(500)
        page.screenshot(path=str(shots["answer"]), full_page=True)

        page.locator("#prompt").fill(CANDIDATE_QUESTION)
        before_candidate = page.locator(".message.assistant").count()
        page.locator("#send").click()
        page.wait_for_function(
            "previous => document.querySelectorAll('.message.assistant').length > previous",
            arg=before_candidate,
            timeout=180_000,
        )
        page.wait_for_timeout(500)
        page.screenshot(path=str(shots["candidate"]), full_page=True)

        page.locator("#routeMode").click()
        page.wait_for_selector(".route-row", timeout=30_000)
        page.wait_for_timeout(300)
        page.screenshot(path=str(shots["route_editor"]), full_page=True)

        page.locator(".route-more").first.click()
        page.wait_for_selector("#routeMenu:not([hidden])", timeout=10_000)
        page.screenshot(path=str(shots["route_menu"]), full_page=True)

        browser.close()

    report = {
        "base_url": base_url,
        "questions": [QUESTION, CANDIDATE_QUESTION],
        "console_errors": console_issues,
        "failed_responses": responses,
        "page_errors": page_errors,
        "screenshots": {key: str(value.resolve()) for key, value in shots.items()},
    }
    return shots, report


def _shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_text(cell: Any, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)


def _add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.35
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        lead.font.size = Pt(10.5)
        rest = paragraph.add_run(text[len(bold_lead):])
        rest.font.size = Pt(10.5)
    else:
        run = paragraph.add_run(text)
        run.font.size = Pt(10.5)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.size = Pt(10.5)


def _add_screenshot(document: Document, image: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(16.8))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 98, 110)


def _add_page_break(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)


def build_manual(output_path: Path, screenshots: dict[str, Path], report: dict[str, Any]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    cover = document.add_table(rows=1, cols=1)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = cover.cell(0, 0)
    _shade(cell, "EAF1FB")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = cell.paragraphs[0].add_run("SOP 校样台\n工人操作手册")
    title.bold = True
    title.font.size = Pt(26)
    title.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("看图操作，先确认，再修改")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(74, 91, 110)
    document.add_paragraph()
    _add_body(document, "适用页面：SOP DOCX 校样台（本地网址：http://127.0.0.1:8787/）")
    _add_body(document, "这份手册只讲工人日常需要用到的功能：选路线、提问题、确认工序、看 DOCX、下载文件和调整工序。")
    _add_body(document, "重要提醒：AI 给出的所有修改都是“待人工核对”，不会自动批准或发布。")
    _add_page_break(document)

    _add_title(document, "1. 先认识这个页面")
    _add_body(document, "页面左边是最新 DOCX 的预览，右边是对话和路线编辑区。顶部可以选择产品路线、填写操作员姓名，并下载当前 DOCX。")
    _add_bullets(document, [
        "先在顶部选择要处理的产品路线。路线版本、状态和工序数量会显示在路线名称旁边。",
        "看到“AI 已连接”时，问题会交给 AI 理解；看到“离线解析”时，系统会按本地规则处理。",
        "DOCX 预览和下载来自同一份文件。预览正在刷新时请稍等，不要连续重复发送。",
    ])
    _add_screenshot(document, screenshots["home"], "图 1：页面总览。左侧看 DOCX，右侧提问或编辑路线。")
    _add_page_break(document)

    _add_title(document, "2. 先填写姓名，再输入问题")
    _add_body(document, "在右上角填写自己的姓名或工号。然后在右下角的输入框里，用日常说话的方式描述你要问的事。")
    _add_bullets(document, [
        "只想查看：加上“只回答，不要修改 SOP”。例如：当前路线共有多少道工序？",
        "要修改某道工序：尽量说出工序名称或编号，再说清楚要改什么。",
        "一次只处理一个问题。内容越具体，系统越容易找到正确位置。",
    ])
    _add_screenshot(document, screenshots["question"], "图 2：填写操作员并输入只查看的问题。")
    _add_page_break(document)

    _add_title(document, "3. 看懂 AI 的回答")
    _add_body(document, "发送后，系统会在对话卡片里给出结果。卡片上会说明这次是 AI 理解、离线解析，还是只做了查看。")
    _add_bullets(document, [
        "“未改动文档”表示这次只是问答，DOCX 不会重新生成。",
        "“DOCX 已重新生成”表示系统已经把有效修改写入草稿，并刷新预览。",
        "“待人工核对”表示需要人工确认，不能当成已批准的生产要求。",
        "点击“查看 AI 判断与实际写入”，可查看找到的工序、改了什么和需要注意的事项。",
    ])
    _add_screenshot(document, screenshots["answer"], "图 3：只查看问题的回答示例。回答不应改动 SOP。")
    _add_page_break(document)

    _add_title(document, "4. 系统找不准时，先选工序")
    _add_body(document, "当一句话可能指向多道工序时，系统会先列出候选工序。此时还没有修改任何内容。")
    _add_bullets(document, [
        "先看工序编号、名称和系统给出的原因。",
        "确认无误后，再点击对应工序。选错时不要继续，重新描述问题即可。",
        "候选卡片出现时，系统会先等待你的选择，避免把内容写错位置。",
    ])
    _add_screenshot(document, screenshots["candidate"], "图 4：需要确认的候选工序。点击前不会写入 SOP。")
    _add_page_break(document)

    _add_title(document, "5. DOCX 预览和下载")
    _add_body(document, "每次有效修改后，系统会重新生成 DOCX。左边预览会显示新页面，右上角“下载 DOCX”会指向同一份最新文件。")
    _add_bullets(document, [
        "修改后先看对话卡片是否显示“DOCX 已重新生成”。",
        "需要核对某项内容时，在回答卡片里点击“定位并高亮”，预览会跳到对应页。",
        "如果预览暂时不可用，DOCX 仍可下载；可稍后点击页面上的重试按钮。",
    ])
    _add_screenshot(document, screenshots["home"], "图 5：左侧 DOCX 预览与右上角下载入口。")
    _add_page_break(document)

    _add_title(document, "6. 路线编辑：增加、排序、拆分、合并和删除")
    _add_body(document, "点击顶部“路线编辑”后，可以查看所有工序。路线结构操作会影响流程图和后续指导书页数，因此请先核对再保存。")
    _add_bullets(document, [
        "新增工序：使用“新增工序”或工序之间的加号。",
        "调整顺序：拖动工序后，点击“保存顺序”。",
        "拆分和合并：从工序右侧菜单进入。合并前要确认相邻工序确实应当合在一起。",
        "删除后有短时间撤销入口，最近删除记录会保留一段时间，便于恢复误删内容。",
    ])
    _add_screenshot(document, screenshots["route_editor"], "图 6：路线编辑列表。这里可检查工序顺序和状态。")
    _add_screenshot(document, screenshots["route_menu"], "图 7：单道工序的操作菜单。先核对名称，再选择操作。")
    _add_page_break(document)

    _add_title(document, "7. 常用提问示例")
    examples = [
        ("只查看", "“裁线工序的检查方法在哪里？请只说明，不要修改 SOP。”"),
        ("修改安全要求", "“第 3 道工序的安全要求补充为：操作前确认设备状态。”"),
        ("修改记录要求", "“电测工序的记录要求补充：登记工单号和异常现象。”"),
        ("定位工序", "“电测对应第几道工序？请只回答。”"),
        ("请求候选", "“把检查工序的作业步骤补充完整。”系统会先请你确认是哪一道检查工序。"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    _set_cell_text(headers[0], "用途", bold=True, color="FFFFFF")
    _set_cell_text(headers[1], "可以这样说", bold=True, color="FFFFFF")
    _shade(headers[0], "1F4E79")
    _shade(headers[1], "1F4E79")
    for use, example in examples:
        row = table.add_row().cells
        _set_cell_text(row[0], use, bold=True)
        _set_cell_text(row[1], example)
    _add_body(document, "不要让 AI 猜测现场数据。设备型号、质量结论、工艺参数、单价、人数和工时等没有确认的内容，应当由负责人填写或核对。")

    _add_title(document, "8. 使用时请记住")
    _add_bullets(document, [
        "先选对路线，再提问。不同产品不要混用内容。",
        "看见候选工序时，先确认再点击；不确定就重新说得更具体。",
        "AI 修改后仍要人工核对。未确认、被退回或有未知项的内容不能正式发布。",
        "图片必须由人工上传并确认，系统不会虚构图片。",
        "遇到页面提示“离线解析”或“预览暂时不可用”时，记录提示内容并联系管理员处理。",
    ])
    _add_body(document, "本手册截图来自本地测试路线，仅用于说明按钮位置和操作过程，不代表生产标准或现场数据。")

    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Capture the SOP UI and create a worker manual DOCX.")
    parser.add_argument("--base-url", default=DEFAULT_URL)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    screenshots, report = capture_screenshots(args.base_url, args.out_dir / "screenshots")
    report_path = args.out_dir / "capture_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    output_path = args.out_dir / "SOP校样台_工人操作手册.docx"
    build_manual(output_path, screenshots, report)
    print(json.dumps({"manual": str(output_path.resolve()), "report": str(report_path.resolve()), **report}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
