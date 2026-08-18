"""Generate the SOP website demonstration video outline as a DOCX file."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "outputs" / "manuals" / "SOP网站功能演示视频大纲_草稿.docx"
SCREENSHOT_DIR = ROOT / "tests" / "screenshots"
SCREENSHOTS = {
    "overview": SCREENSHOT_DIR / "new-project-live-20260817" / "desktop-03-docx-preview.png",
    "new_project": SCREENSHOT_DIR / "new-project-live-20260817" / "desktop-02-result.png",
    "media": SCREENSHOT_DIR / "media-arrangement-order-controls-final-20260815.png",
}


BLUE = "2459C4"
DARK_BLUE = "1F4E79"
LIGHT_BLUE = "EAF1FB"
PALE_BLUE = "F4F7FC"
PALE_YELLOW = "FFF3C4"
PALE_GREEN = "E9F4ED"
INK = "202124"
MUTED = "59626E"
WHITE = "FFFFFF"


TIMELINE = [
    ("0:00-0:25", "开场效果", "快速展示一次自然语言修改，随后切到 DOCX 刷新和定位结果。", "不需要翻表格找位置。直接说明哪道工序要改，系统会定位、写入草稿并刷新 DOCX。"),
    ("0:25-1:05", "页面总览", "展示产品选择、路线版本、状态、工序数量、操作员和下载按钮。", "左边是最新 DOCX，右边是对话和路线编辑区。开始前先选对产品，并填写姓名或工号。"),
    ("1:05-2:05", "自然语言新建项目", "点击“+ 新建项目”，输入产品和大致工序，再点击“让 AI 整理”。", "AI 先整理预览，不会立即写入。缺少的参数、设备和标准会显示待补充，不会擅自猜测。"),
    ("2:05-2:35", "确认项目草稿", "核对产品名称、工序顺序和待补充项，点击“创建项目草稿”。", "创建后路线仍是草稿，全部工序等待人工核对，并自动生成第一版 DOCX。"),
    ("2:35-3:20", "只询问，不修改", "输入只查看问题，观察回答卡片和“未改动文档”标记。", "普通问答不会改 SOP，也不会重复生成 DOCX。"),
    ("3:20-4:15", "自然语言修改", "输入明确的工序修改要求，展开“查看 AI 判断与实际写入”。", "卡片会显示定位到的工序、修改字段、实际写入内容、警告和待人工核对状态。"),
    ("4:15-4:50", "候选工序确认", "输入可能对应多道工序的问题，展示候选列表，再人工选择目标。", "系统找不准时会先让人选，不会直接把内容写到猜测的工序。"),
    ("4:50-5:35", "DOCX 定位与下载", "点击“定位并高亮”，观察预览跳页；随后展示“下载 DOCX”。", "预览和下载来自同一份最新 DOCX。有效修改后会显示“DOCX 已重新生成”。"),
    ("5:35-6:45", "路线编辑", "展示新增、拖动排序、拆分、合并、删除和撤销入口。", "路线变化会同步更新流程图、指导书顺序和页数。误删可以立即撤销或从最近删除恢复。"),
    ("6:45-7:55", "项目图片管理", "打开完整版和“项目图片统一调整”，上传、绑定、排序并确认图片。", "图片必须由人工提供。只有经过人工确认的图片才会进入指导书。"),
    ("7:55-8:40", "人工确认与历史知识", "确认一条演示工序，再搜索相似内容并查看产品、工序和来源。", "保存草稿不等于确认。确认后的内容才能进入历史搜索，整条路线仍需独立批准。"),
    ("8:40-9:15", "异常与安全边界", "使用准备好的离线解析和预览失败截图进行说明。", "AI 不可用会降级提示；预览失败时 DOCX 仍可下载；系统不会自动批准或编造现场参数。"),
    ("9:15-9:30", "收尾", "回到页面总览，停留在草稿和待核对状态。", "系统负责把现场描述变成可核对的 SOP 草稿，最终确认权始终在人。"),
]


def _set_run_font(run, size: float = 10.5, *, bold: bool = False, color: str = INK) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_cell_text(cell, text: str, *, size: float = 9, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    _set_run_font(paragraph.add_run(text), size, bold=bold, color=color)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("SOP 网站功能演示大纲（草稿）  |  第 ")
    _set_run_font(prefix, 8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    suffix = paragraph.add_run(" 页")
    _set_run_font(suffix, 8, color=MUTED)


def _configure_section(section, *, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.clear()
    _add_page_number(footer)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    size = 18 if level == 1 else 13
    color = DARK_BLUE if level == 1 else BLUE
    _set_run_font(paragraph.add_run(text), size, bold=True, color=color)


def _add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.35
    if bold_lead and text.startswith(bold_lead):
        _set_run_font(paragraph.add_run(bold_lead), 10.5, bold=True)
        _set_run_font(paragraph.add_run(text[len(bold_lead):]), 10.5)
    else:
        _set_run_font(paragraph.add_run(text), 10.5)


def _add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.25
        _set_run_font(paragraph.add_run(item), 10)


def _add_callout(document: Document, title: str, text: str, *, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, fill)
    _set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    _set_run_font(paragraph.add_run(title), 11, bold=True, color=DARK_BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.3
    _set_run_font(paragraph.add_run(text), 10.5)


def _add_screenshot(document: Document, image: Path, caption: str, *, width: float = 24.0) -> None:
    if not image.is_file():
        _add_callout(document, "截图待补", str(image), fill=PALE_YELLOW)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.add_run().add_picture(str(image), width=Cm(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(5)
    run = caption_paragraph.add_run(caption)
    _set_run_font(run, 8.5, color=MUTED)
    run.italic = True


def _add_timeline(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(1.9), Cm(3.1), Cm(9.2), Cm(12.3))
    headers = ("时间", "演示章节", "页面操作", "口播重点")
    for index, (cell, width, header) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = width
        _shade(cell, DARK_BLUE)
        _set_cell_text(cell, header, size=9, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cant_split(table.rows[0])
    _repeat_table_header(table.rows[0])
    for row_index, values in enumerate(TIMELINE, start=1):
        row = table.add_row()
        _cant_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            cell.width = width
            if row_index % 2 == 0:
                _shade(cell, PALE_BLUE)
            _set_cell_text(cell, value, size=8.5, bold=False)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text(row.cells[1], values[1], size=8.5, bold=True, color=DARK_BLUE)


def _add_prompt(document: Document, label: str, prompt: str, expected: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(4.1)
    right.width = Cm(22.0)
    _shade(left, LIGHT_BLUE)
    _set_cell_text(left, label, size=9.5, bold=True, color=DARK_BLUE)
    right.text = ""
    _set_cell_margins(right, top=120, start=150, bottom=120, end=150)
    paragraph = right.paragraphs[0]
    _set_run_font(paragraph.add_run(prompt), 10, bold=True)
    paragraph = right.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    _set_run_font(paragraph.add_run(f"预期：{expected}"), 9, color=MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_document() -> Path:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "SOP 网站功能演示视频大纲（草稿）"
    document.core_properties.subject = "SOP 校样台客户与工人功能演示"
    document.core_properties.author = "Yunpai SOP 项目组"
    document.core_properties.keywords = "SOP, DOCX, 演示视频, 操作手册"

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    _configure_section(document.sections[0], landscape=False)

    cover = document.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover.cell(0, 0)
    _shade(cell, LIGHT_BLUE)
    _set_cell_margins(cell, top=450, start=300, bottom=450, end=300)
    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.add_run("SOP 网站功能演示视频大纲"), 25, bold=True, color=DARK_BLUE)
    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(10)
    _set_run_font(subtitle.add_run("从自然语言新建项目到生成最新 DOCX"), 14, color=BLUE)
    tag = cell.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_before = Pt(12)
    _set_run_font(tag.add_run("草稿 v0.1  |  预计时长 9 分 30 秒"), 10, bold=True, color=MUTED)

    document.add_paragraph()
    _add_callout(
        document,
        "整条视频只讲清一件事",
        "工作人员可以直接描述需求，AI 负责定位和整理，系统自动更新 DOCX；所有内容仍需人工核对，不会自动批准或发布。",
        fill=PALE_YELLOW,
    )
    document.add_paragraph()
    _add_body(document, "目标观众：客户、一线工作人员、现场审核人员", bold_lead="目标观众：")
    _add_body(document, "录制形式：Microsoft Edge 实机操作 + 口播，异常场景可插入已有截图", bold_lead="录制形式：")
    _add_body(document, f"准备日期：{date.today().isoformat()}  |  本文件仅为演示草稿", bold_lead="准备日期：")

    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(landscape, landscape=True)
    _add_heading(document, "1. 视频结构与时间轴")
    _add_body(document, "建议使用一条完整业务线进行演示：新建草稿 → 自然语言修改 → 定位 DOCX → 调整路线和图片 → 人工核对与复用。")
    _add_timeline(document)

    document.add_page_break()
    _add_heading(document, "2. 开场与收尾口径")
    _add_callout(
        document,
        "开场建议",
        "传统修改 SOP，要找文件、找工序、改表格，再重新导出。这个页面里，直接说清楚哪道工序要改，系统就会定位、写进草稿，并刷新 DOCX。",
    )
    document.add_paragraph()
    _add_callout(
        document,
        "收尾建议",
        "系统负责把现场描述变成可核对的 SOP 草稿，最终确认权始终在人。AI 不会自动批准、发布，也不会替现场人员编造参数和质量结论。",
        fill=PALE_GREEN,
    )
    _add_screenshot(document, SCREENSHOTS["overview"], "画面参考 1：主界面同时展示最新 DOCX、自然语言对话和路线状态。")

    document.add_page_break()
    _add_heading(document, "3. 重点演示流程")
    _add_heading(document, "3.1 自然语言新建项目", level=2)
    _add_bullets(document, [
        "点击产品选择器旁边的“+ 新建项目”。",
        "用日常说话的方式描述产品和大致工序，不确定的参数不用猜。",
        "点击“让 AI 整理”，先核对产品名称、工序顺序、待补充项和警告。",
        "只有点击“创建项目草稿”后才会写入数据库并生成第一版 DOCX。",
        "强调创建结果仍是 draft，全部工序保持 needs_revision。",
    ])
    _add_screenshot(document, SCREENSHOTS["new_project"], "画面参考 2：AI 整理结果先供人工核对，确认前不会创建项目。")

    _add_heading(document, "3.2 自然语言询问和修改", level=2)
    _add_bullets(document, [
        "先演示只查看问题，让观众看到“未改动文档”。",
        "再演示一次明确修改，展开“查看 AI 判断与实际写入”。",
        "说明卡片中的工序、字段、写入内容、警告和待人工核对状态。",
        "如果目标不明确，展示候选工序；人工选择前不得写入。",
    ])

    _add_heading(document, "3.3 DOCX 定位、路线编辑与图片", level=2)
    _add_bullets(document, [
        "有效修改后确认卡片出现“DOCX 已重新生成”。",
        "点击“定位并高亮”，让左侧预览跳到对应页，并展示最新下载入口。",
        "切换路线编辑，展示新增、排序、拆分、合并、删除和撤销。",
        "打开完整版和项目图片统一调整，展示上传、绑定、排序、解除和人工确认。",
        "说明流程图、指导书顺序、页数和已确认图片都会随路线重新生成。",
    ])
    _add_screenshot(document, SCREENSHOTS["media"], "画面参考 3：项目图片统一调整页面用于绑定工序、排序并人工确认图片。")

    document.add_page_break()
    _add_heading(document, "4. 现场演示问题")
    _add_prompt(
        document,
        "只查看",
        "电测对应第几道工序？请只回答，不要修改 SOP。",
        "回答目标位置，显示“未改动文档”，DOCX 不重新生成。",
    )
    _add_prompt(
        document,
        "明确修改",
        "把“裁线与长度补偿”工序的安全要求补充为：操作前确认设备状态。",
        "显示定位工序、修改字段和实际写入；标记待人工核对，并重新生成 DOCX。",
    )
    _add_prompt(
        document,
        "候选定位",
        "把检查工序的作业步骤补充完整。",
        "如果存在多道检查工序，先列出候选；人工选择前不写入。",
    )
    _add_prompt(
        document,
        "新建项目",
        "这是一个 USB-C 转接线项目。先检查来料，然后裁线、剥皮、焊接端子，最后进行通电检查和包装。具体参数还没有确认。",
        "AI 整理产品和工序，未提供的设备、参数和标准保持待补充。",
    )
    _add_callout(
        document,
        "录制提醒",
        "正式录制前先确认当前演示路线中确实存在对应工序。若名称不同，应直接使用页面显示的完整工序名称，避免现场出现不必要的候选或定位偏差。",
        fill=PALE_YELLOW,
    )

    document.add_page_break()
    _add_heading(document, "5. 录制前检查清单")
    _add_bullets(document, [
        "使用专门的演示草稿项目，不直接修改正式路线。",
        "填写演示用操作员姓名或工号，并提前确认所选产品和路线。",
        "等待 DOCX 首页实际显示后再开始录屏。",
        "准备一张不超过 10MB 的 PNG 或 JPEG 演示图片。",
        "把四条演示问题放在记事本中，现场复制，避免输入错误。",
        "提前演练一次删除和撤销，确保不会留下无用演示工序。",
        "AI 等待过程可以后期剪短，但保留“正在整理”和最终结果两个状态。",
        "离线解析、预览失败等异常场景使用已有截图说明，不必录制时主动切断服务。",
        "录制建议 1920×1080、浏览器缩放 100%，关闭通知和无关窗口。",
        "不得在画面中展示 API 密钥、数据库路径、客户隐私或未经确认的生产数据。",
    ])

    _add_heading(document, "6. 必须说清楚的安全边界")
    _add_callout(
        document,
        "人工确认不能省",
        "AI 写入后仍是待人工核对。未确认、被驳回、存在阻断未知项或整条路线未批准的内容，都不能当成正式生产模板直接使用。",
        fill=PALE_YELLOW,
    )
    _add_bullets(document, [
        "AI 不得编造设备型号、工艺参数、工时、人数、价格、质量结论或现场事实。",
        "图片只能来自人工实际上传；未确认图片不能进入正式指导书。",
        "已批准路线不能直接覆盖，修改时必须先创建新修订版。",
        "离线解析无法理解时应明确提示，不能胡乱写入。",
        "预览与下载必须来自同一份最新 DOCX。",
    ])

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
