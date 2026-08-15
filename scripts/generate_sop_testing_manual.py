from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuals" / "Yunpai_SOP_辅助测试说明书_对话与DOCX预览版.docx"


def shade(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    header = t.rows[0]
    repeat_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        cell.text = value
        cell.width = Cm(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "1F4E78")
        for run in cell.paragraphs[0].runs:
            run.font.name = "Microsoft YaHei"
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Cm(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(8)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    for name, size, color in (("Normal", 10, "000000"), ("Title", 23, "17365D"), ("Heading 1", 16, "17365D"), ("Heading 2", 12, "1F4E78")):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "Test Note" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Test Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string("7F6000")


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Test Note")
    p.add_run("注意：").bold = True
    p.add_run(text)


def case_rows(cases: list[tuple[str, str, str, str, str]]) -> list[list[str]]:
    return [[case_id, action, expected, evidence, priority] for case_id, action, expected, evidence, priority in cases]


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Yunpai SOP 辅助测试说明书")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("对话式 SOP 修改 + DOCX 预览版 | 测试基线 | 2026-08-13")
    doc.add_paragraph()
    note(doc, "本说明书用于人工验收本地服务 http://127.0.0.1:8787/。当前目标产品仅保留“自然语言对话修改”和“DOCX 预览/下载”两类能力；所有测试应使用测试路线和虚构数据，不得将 AI 产出视为生产批准。")

    doc.add_heading("1. 测试目标与判定", level=1)
    bullets(doc, [
        "确认首页可以选择产品路线、显示路线版本/状态/工序数量，并且只展示自然语言对话修改与 DOCX 预览。",
        "确认 AI 修改有明确定位、变更说明、人工核对标记和安全边界；问答、未知信息和不可理解输入不能误写入。",
        "确认有效修改会生成与页面预览一致的新 DOCX；流程、工序页、图片和草稿状态均符合版式与审批规则。",
        "确认工艺路线、图片、知识复用和版本审批的变更不会绕过人工确认或混用其他产品资料。",
    ])
    table(doc, ["等级", "判定与处理"], [
        ["P0 阻断", "误批准/误发布、写错产品或工序、数据丢失、DOCX 与页面不一致、已批准版本被直接覆盖。立即停止该路线继续写入。"],
        ["P1 严重", "核心修改、生成、确认、版本或知识隔离失败，导致测试结论不可信。"],
        ["P2 一般", "功能可替代但流程、提示、异常处理或展示不符合要求。"],
        ["P3 体验", "排版、文案、轻微交互或视觉问题，不影响数据和审批安全。"],
    ], [3.2, 15.6])

    doc.add_heading("2. 测试前准备", level=1)
    table(doc, ["项目", "要求"], [
        ["访问地址", "http://127.0.0.1:8787/；记录测试开始时间、浏览器版本、路线名称/ID、初始版本指纹和工序数量。"],
        ["测试数据", "准备两个不同产品型号的测试路线；两张小于 5MB 的 PNG/JPEG；测试人员姓名；可丢弃的描述文本。不得输入真实客户、人员或受控参数。"],
        ["AI 状态", "先记录页面显示“AI模型已连接”或“离线兜底模式”。分别完成在线与离线场景；切换服务配置前先导出或记录当前版本。"],
        ["证据", "每条失败用例至少保留：页面截图、输入原文、实际结果、路线与工序标识、DOCX 文件名/版本指纹；必要时保留浏览器控制台和网络请求。"],
        ["恢复", "涉及新增、删除、合并、批准的用例均在专用测试路线执行；结束后不要把测试路线作为正式模板复用。"],
    ], [3.2, 15.6])

    doc.add_heading("3. 页面与基础展示", level=1)
    base_cases = [
        ("UI-01", "访问 /。确认首页只显示：路线选择、路线摘要、自然语言对话区、DOCX 预览/下载区。", "页面正常加载；不出现旧填表工作台、工序编辑表单或其他旧模块。", "首页全屏截图", "P0"),
        ("UI-02", "直接访问 /workbench、刷新并用新窗口访问。", "旧工作台不存在，应返回 404/明确迁移页，而不是旧功能界面。", "地址栏和响应截图", "P0"),
        ("UI-03", "选择产品路线 A，再选择路线 B，然后切回 A。", "每次显示对应路线的产品、版本、状态、工序数量、DOCX 和对话上下文；不得串用。", "三次路线摘要截图", "P0"),
        ("UI-04", "检查路线摘要信息和 AI 状态。", "明确显示路线版本、状态、工序数量，以及“AI模型已连接”或“离线兜底模式”。", "摘要区域截图", "P1"),
        ("UI-05", "页面刷新、浏览器后退再前进。", "恢复当前路线及最新 DOCX，不显示陈旧版本或空白预览。", "刷新前后版本指纹", "P1"),
    ]
    table(doc, ["编号", "操作", "预期结果", "证据", "优先级"], case_rows(base_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("4. 自然语言修改 SOP", level=1)
    nl_cases = [
        ("NL-01", "输入：“将工序 2 名称改为 外观与端子检查”。", "准确定位工序 2，仅修改工序名称；结果卡列出目标、字段、写入文本与待核对状态。", "对话卡和路线版本", "P1"),
        ("NL-02", "输入：“把外观检查的检查方法改为目视检查插头外壳、端子、线身和护套。”", "可按工序名称定位，修改检查方法而非其他字段。", "对话卡和 DOCX", "P1"),
        ("NL-03", "逐次要求修改工序动作、作业步骤、工具/设备/治具、材料/输入、工艺参数、合格判据、安全要求、记录要求、异常处理。", "每次仅写入对应字段，完整保留多行内容；全部标记 needs_revision/待人工核对。", "每次的对话卡和字段截图", "P0"),
        ("NL-04", "输入：“把路线的包装标签章节补充为：标签内容按受控资料核对。”", "定位路线级章节，显示章节变更；不误写到任意工序。", "章节结果和 DOCX", "P1"),
        ("NL-05", "输入：“在外观检查和电测之间新增‘扫码绑定批次’工序，记录工单号和序列号。”", "新增工序及其顺序正确，流程图和指导书页同步增加。", "工序数、流程图和页数", "P0"),
        ("NL-06", "先输入“第 3 道工序增加异常隔离要求”，再输入“记录要求也补上”。", "第二句沿用已定位的工序；两次改动均清晰展示。", "连续对话截图", "P1"),
        ("NL-07", "输入：“第 2 道工序现在有什么检查要求？”或“解释一下这条路线”。", "作为问答回答，不产生版本变化、不写入字段、不重新生成 DOCX。", "前后版本指纹和对话卡", "P0"),
        ("NL-08", "输入：“把第 99 道工序改成通过。”以及模糊语句“给它优化一下”。", "无法定位时说明未确认/需澄清，不得猜测目标或写入。", "对话卡和前后内容", "P0"),
        ("NL-09", "输入无来源参数，例如“把平均工时定为 15 秒”“填入 XX-9000 设备型号”“直接判定合格”。", "不得编造或写入现场事实；必须标为未知、警告或要求受控来源/人工确认。", "警告内容与字段值", "P0"),
    ]
    table(doc, ["编号", "操作/输入", "预期结果", "证据", "优先级"], case_rows(nl_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("5. AI 路由与离线兜底", level=1)
    ai_cases = [
        ("AI-01", "在模型配置有效时提交一条明确修改。", "回答显示“AI已理解”；后台/结果记录 parser_kind=llm。", "对话卡、接口响应或日志", "P1"),
        ("AI-02", "临时使模型不可用后，提交可由规则理解的简单修改，例如“第 1 道工序改名为备料核对”。", "自动降级且明确显示“离线解析”；修改范围正确。", "状态与对话卡", "P1"),
        ("AI-03", "在离线模式输入复杂、含糊或无目标的修改。", "离线规则无法理解时拒绝写入并提示原因，不能假装理解或乱改。", "前后版本及提示", "P0"),
        ("AI-04", "恢复模型后再次提交明确修改。", "显示恢复到 AI 已理解；不遗留错误的离线状态。", "状态截图", "P2"),
    ]
    table(doc, ["编号", "操作", "预期结果", "证据", "优先级"], case_rows(ai_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("6. 修改结果与 DOCX 生成", level=1)
    docx_cases = [
        ("DX-01", "完成一条有效修改并等待生成完成。", "显示“DOCX已重新生成”；版本指纹变化，下载链接切换到最新文件。", "修改前后指纹和下载名", "P0"),
        ("DX-02", "仅提问，不执行写入。", "不重新生成 DOCX，版本指纹、文件链接和修改时间不变。", "修改前后截图", "P0"),
        ("DX-03", "下载最新 DOCX，并比对页面预览中被修改的工序文本。", "预览来自同一最新 DOCX 转换的 PDF；预览、下载文件和路线字段一致。", "PDF/PNG 与 DOCX 截图", "P0"),
        ("DX-04", "检查版式与页数：流程图页及所有工序指导书页。", "第 1 页 A4 纵向完整流程图；第 2 页起 A4 横向；每道工序独占一页；总页数 = 1 + 工序数量。", "页数、页面方向截图", "P0"),
        ("DX-05", "检查指导书六个动作区。", "动作区顺序必须是 1、2、3 / 6、5、4；不得被压缩成固定两页模板。", "单页截图", "P1"),
        ("DX-06", "未上传图片时生成；再上传并确认一张工序图片后重新生成。", "未提供图片位置保持空白；确认后仅使用用户实际图片；签名栏始终空白且文件仍为草稿。", "两版 DOCX 对比", "P0"),
        ("DX-07", "刷新页面并重新打开路线。", "仍加载最新 DOCX、PDF 和 PNG 页预览，三者页数相同。", "刷新后预览和页数", "P0"),
    ]
    table(doc, ["编号", "操作", "预期结果", "证据", "优先级"], case_rows(docx_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("7. 路线、图片与知识复用", level=1)
    knowledge_cases = [
        ("RT-01", "新增工序、删除该工序、改变相邻工序排序，分别生成 DOCX。", "流程图、指导书数量和顺序同步；删除工序的指导书必须消失。", "每一步页数和流程图", "P0"),
        ("RT-02", "对一条工序执行拆分子步骤、再测试合并工序。", "步骤关系和内容完整保留，工序数与 DOCX 页数按实际结果变化。", "变更前后内容", "P1"),
        ("IMG-01", "上传 PNG/JPEG 并关联到指定工序；尝试错误格式。", "支持 PNG/JPEG；关联正确；错误格式有明确拒绝提示。", "图片缩略图和提示", "P1"),
        ("IMG-02", "上传图片但不人工确认后生成 DOCX；随后人工确认再生成。", "未确认图片保持草稿且不进入正式指导书；确认后才可出现。", "两版 DOCX", "P0"),
        ("KN-01", "AI 修改一项后，立即搜索相关关键词。", "工序状态 needs_revision；未人工确认内容不能进入确认知识索引或正式复用结果。", "状态和搜索结果", "P0"),
        ("KN-02", "人工确认后搜索相似产品/工序。", "结果包含产品、工序、来源路线和版本；草稿/驳回默认排除。", "搜索结果截图", "P1"),
        ("KN-03", "在产品 A 写入唯一测试词，再到产品 B 搜索并尝试复用。", "不同型号不得内容串用；复用记录必须保留来源路线、来源工序和版本。", "A/B 搜索与复用记录", "P0"),
    ]
    table(doc, ["编号", "操作", "预期结果", "证据", "优先级"], case_rows(knowledge_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("8. 版本与审核安全", level=1)
    approval_cases = [
        ("AP-01", "对已批准路线提交一条修改。", "不能覆盖既有批准版本，系统必须创建新的修订版。", "旧/新版本标识", "P0"),
        ("AP-02", "保留一个阻断性未知项或未确认工序/章节，尝试正式批准。", "必须拒绝批准，并列出缺失的确认项；不产生部分批准副作用。", "拒绝提示和状态", "P0"),
        ("AP-03", "批准人留空、填写错误确认口令，分别执行批准。", "两种情况均拒绝，路线状态和知识索引保持不变。", "表单及前后状态", "P0"),
        ("AP-04", "将一项内容标为驳回后，搜索并尝试正式复用/批准。", "驳回内容不得进入批准索引或正式模板。", "搜索和审批记录", "P0"),
        ("AP-05", "检查所有 AI 写入后和异常处理相关状态。", "系统不得自动批准、自动发布或自动关闭异常；所有 AI 写入均待人工核对。", "状态历史和对话卡", "P0"),
    ]
    table(doc, ["编号", "操作", "预期结果", "证据", "优先级"], case_rows(approval_cases), [1.3, 5.1, 7.1, 3.1, 1.3])

    doc.add_heading("9. 缺陷记录模板", level=1)
    table(doc, ["字段", "填写说明"], [
        ["缺陷编号", "例如 SOP-20260813-001。"],
        ["严重级别", "P0/P1/P2/P3，并写明是否已停止该路线写入。"],
        ["环境", "浏览器版本、URL、产品型号、路线 ID、路线版本、AI 状态。"],
        ["前置条件", "例如：工序 3 未确认，当前 DOCX 指纹为 xxx。"],
        ["复现步骤", "按 1、2、3 编号，保留完整自然语言输入和点击对象。"],
        ["预期/实际", "分别填写规则要求与实际页面、状态、文件或接口结果。"],
        ["证据", "截图、录屏、DOCX/PDF 文件名、控制台或网络请求。"],
        ["影响", "是否造成误改、误批准、串用、数据丢失或 DOCX 不一致。"],
    ], [3.4, 15.4])

    doc.add_heading("10. 测试结论", level=1)
    table(doc, ["项目", "结果/备注"], [
        ["测试日期、人员与浏览器", ""],
        ["产品路线与初始/最终版本", ""],
        ["已执行用例数：通过/失败/阻塞", ""],
        ["P0/P1 缺陷列表", ""],
        ["AI 在线与离线兜底结论", ""],
        ["DOCX 预览与下载一致性结论", ""],
        ["是否满足批准与知识复用安全边界", ""],
        ["是否允许进入下一轮验收", ""],
    ], [6.0, 12.8])
    doc.add_paragraph("本手册为测试辅助材料，不构成生产放行、质量判定或正式 SOP 发布依据。")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
