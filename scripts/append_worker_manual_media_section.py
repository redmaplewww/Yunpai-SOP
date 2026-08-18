"""Add real screenshots and a plain-language image-arrangement guide to the worker manual."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any
from urllib.parse import urljoin

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from playwright.sync_api import sync_playwright


DEFAULT_BASE_URL = "http://127.0.0.1:8787/"
DEFAULT_OUTPUT = Path("outputs/worker_function_manual")


def edge_executable() -> str:
    candidates = (
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    )
    for path in candidates:
        if path.is_file():
            return str(path)
    raise FileNotFoundError("Microsoft Edge was not found")


def capture_media_screenshots(base_url: str, screenshot_dir: Path) -> dict[str, Any]:
    """Capture the real entry and media-arrangement page without changing route data."""
    screenshot_dir.mkdir(parents=True, exist_ok=True)
    entry_shot = screenshot_dir / "08-项目图片统一调整入口.png"
    arrangement_shot = screenshot_dir / "09-项目图片统一调整界面.png"
    console_errors: list[str] = []
    page_errors: list[str] = []
    failed_responses: list[str] = []

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(executable_path=edge_executable(), headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 900}, device_scale_factor=1)
        page.on(
            "console",
            lambda message: console_errors.append(f"{message.type}: {message.text}")
            if message.type == "error"
            else None,
        )
        page.on("pageerror", lambda error: page_errors.append(str(error)))
        page.on(
            "response",
            lambda response: failed_responses.append(f"{response.status} {response.url}")
            if response.status >= 400
            else None,
        )

        workbench_url = urljoin(base_url, "workbench")
        page.goto(workbench_url, wait_until="domcontentloaded", timeout=30_000)
        page.wait_for_selector("#mediaArrangeLink", timeout=30_000)
        page.wait_for_selector("#routeSelect option", state="attached", timeout=30_000)
        page.locator("#mediaArrangeLink").locator("xpath=..").screenshot(path=str(entry_shot))

        route_id = page.locator("#routeSelect option").first.get_attribute("value")
        if not route_id:
            raise RuntimeError("The workbench did not load a route")
        arrangement_url = urljoin(base_url, f"media-arrangement?route={route_id}")
        page.goto(arrangement_url, wait_until="domcontentloaded", timeout=30_000)
        page.wait_for_selector("#routeSelect option", state="attached", timeout=30_000)
        page.wait_for_selector("#flowRail .process-card", timeout=30_000)
        page.wait_for_timeout(500)
        page.screenshot(path=str(arrangement_shot), full_page=True)
        browser.close()

    return {
        "screenshots": {
            "media_entry": str(entry_shot.resolve()),
            "media_arrangement": str(arrangement_shot.resolve()),
        },
        "console_errors": console_errors,
        "page_errors": page_errors,
        "failed_responses": failed_responses,
    }


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.size = Pt(10.5)


def add_screenshot(document: Document, image: Path, caption: str, *, width_cm: float = 16.8) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 98, 110)


def append_manual_section(manual_path: Path, screenshots: dict[str, str]) -> None:
    document = Document(manual_path)
    existing_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "补充功能：项目图片统一调整" in existing_text:
        raise RuntimeError("The image-arrangement section is already in this manual")

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_page_break()
    add_title(document, "补充功能：项目图片统一调整")
    add_body(document, "这个功能用于把项目图片放到正确的工序，并整理同一工序内的图片顺序。它只处理实际上传的图片，不会自动生成或猜测图片。")
    add_bullets(
        document,
        [
            "在“图片素材”区域，先点击“项目图片统一调整”。这个按钮在“项目图片上传”下面。",
            "进入前请确认已选对产品路线。不同产品的图片不能混用。",
            "这里只是整理图片和工序的对应关系。图片没有人工确认前，仍然是草稿状态。",
        ],
    )
    add_screenshot(
        document,
        Path(screenshots["media_entry"]),
        "图 8：在“项目图片上传”下方点击“项目图片统一调整”。",
        width_cm=12.4,
    )

    document.add_page_break()
    add_title(document, "怎样整理图片并让 DOCX 更新")
    add_body(document, "进入后，左边是项目图片，中间是全部工序，右边是当前选中工序的图片清单。按下面顺序操作即可。")
    add_bullets(
        document,
        [
            "还没有图片时：点左边的“+”上传 PNG 或 JPEG 图片。单张图片不能超过 10MB。",
            "绑定图片：把左边的图片拖到中间对应的工序；也可以先点选图片，再点目标工序。每道工序最多绑定 6 张图片。",
            "改图或取消：在右边的当前工序清单里，可以解除不需要的图片；多张图片可用左侧拖动把手或上下箭头调整顺序。DOCX 会按这里的顺序放图。",
            "人工确认：新绑定或替换的图片会显示“待人工确认”。核对无误后，点右侧的确认按钮。只有已人工确认的图片才会进入正式指导书图片区。",
            "看 DOCX：每次绑定、替换、解除、调整顺序或确认图片后，系统都会重建最新 DOCX。页面顶部会更新版本信息，可直接下载最新文件。",
        ],
    )
    add_screenshot(
        document,
        Path(screenshots["media_arrangement"]),
        "图 9：项目图片统一调整页面。左边选图，中间选工序，右边核对、调整顺序并人工确认。",
    )
    add_body(document, "提醒：图片一旦确认后才会写进指导书；不确定图片是否适用时，先不要确认。已批准路线不能直接修改图片，应先创建新修订版。")
    document.save(manual_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Append the image-arrangement guide to the worker manual.")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    manual_path = args.out_dir / "SOP校样台_工人操作手册.docx"
    if not manual_path.is_file():
        raise FileNotFoundError(f"Worker manual not found: {manual_path}")
    capture = capture_media_screenshots(args.base_url, args.out_dir / "screenshots")
    append_manual_section(manual_path, capture["screenshots"])
    report_path = args.out_dir / "media_arrangement_capture_report.json"
    report_path.write_text(json.dumps(capture, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"manual": str(manual_path.resolve()), "report": str(report_path.resolve()), **capture}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
