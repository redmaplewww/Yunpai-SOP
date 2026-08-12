from __future__ import annotations

import argparse
import hashlib
import json
import sys
from datetime import date
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt

from cad_ai.sop_agent import _validate_docx_package
from cad_ai.sop_visual_template import (
    IE_TIME_STUDY_FIELDS,
    SOP_FLOWCHART_SHAPE_POLICY,
    _blank_ie_time_study_row,
    _build_sop_word_document,
    _demo_step,
    _render_center_flowchart_shape_image,
    _set_row_height,
    _set_word_cell,
    _write_word_format_check_json,
    build_process_flow_page,
    build_work_instruction_page,
    classify_sop_flow_node_shape,
)
from scripts.generate_sop_template_ai_handoff import _apply_delivery_controls


PROFILE_ID = "yunpai.sop.hdmi_finished_cable.inspection_packaging.draft.v1"
DOCX_NAME = "HDMI成品线检验与包装_SOP图_草案.docx"
FLOWCHART_NAME = "HDMI成品线_中心工艺流程图.png"
FORMAT_CHECK_NAME = "HDMI成品线_SOP格式检查.json"
VALIDATION_NAME = "HDMI成品线_SOP结构验证.json"
MANIFEST_NAME = "HDMI成品线_SOP生成清单.json"


def _ie_row(action: str, machine_model: str) -> dict[str, str]:
    row = _blank_ie_time_study_row(action)
    row.update(
        {
            "机器型号": machine_model,
            "IE测量方法": "待IE实测",
            "观测次数": "",
            "平均观测工时(s)": "",
            "评比系数": "",
            "宽放率": "",
            "标准工时(s)": "",
            "工时来源": "DRAFT",
            "动态调整": "人工复核后更新",
        }
    )
    return row


def _build_pages(document_date: str) -> tuple[dict[str, Any], dict[str, Any]]:
    product_name = "HDMI 成品线（规格待确认）"
    part_no = "TBD"
    document_no = "SOP-DRAFT-HDMI-001"
    drawing_no = "TBD"

    node_specs = [
        ("备料核对", "inspection"),
        ("外观检查", "inspection"),
        ("导通/短路测试", "test"),
        ("电气结果判定", "quality"),
        ("功能测试", "test"),
        ("功能结果判定", "quality"),
        ("盘线扎线", "process"),
        ("装袋贴标装箱", "process"),
    ]
    nodes: list[dict[str, Any]] = []
    for index, (name, node_type) in enumerate(node_specs, start=1):
        node = {
            "seq": index,
            "id": f"OP{index:02d}",
            "label": f"工序{index}",
            "name": name,
            "type": node_type,
            "station": "成品检验与包装",
            "note": "demo_not_for_release",
        }
        node["shape"] = classify_sop_flow_node_shape(node)
        nodes.append(node)

    flow_page = build_process_flow_page(
        product_name=product_name,
        part_no=part_no,
        document_no=document_no,
        drawing_no=drawing_no,
        operations=[name for name, _node_type in node_specs],
    )
    flow_page["flow_nodes"] = nodes
    flow_page["render_center_flowchart"] = True
    flow_page["center_flowchart_style"] = "pdf_reference_shape_blocks"
    flow_page["center_flowchart"] = {
        "flowchart_title": "HDMI 成品线检验与包装",
        "nodes": nodes,
        "edges": [
            {"from": nodes[index]["id"], "to": nodes[index + 1]["id"], "label": "next"}
            for index in range(len(nodes) - 1)
        ],
    }
    flow_page["ie_time_study"] = {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": [_ie_row("流程工时汇总", "TBD")],
        "policy": {
            "measurement_basis": "待现场IE实测",
            "release_requirement": "site_measured_or_human_locked",
        },
    }

    work_page = build_work_instruction_page(
        product_name=product_name,
        part_no=part_no,
        station="成品检验与包装",
        document_no=document_no,
        drawing_no=drawing_no,
        version="DRAFT",
        page_no=1,
        page_total=1,
    )
    work_page["operation_order"] = "①备料==②外观==③电测==④功能==⑤盘线==⑥包装"
    work_page["step_slots"] = [
        _demo_step(1, "备料核对", "按受控工单/BOM核对线材规格、版本、数量及包装辅料。", "materials"),
        _demo_step(2, "外观检查", "检查插头外壳、端子、线身和护套；无变形、破损、露铜和污染。", "inspect"),
        _demo_step(3, "导通/短路测试", "接受控测试仪，按批准程序检查针位导通、开短路及屏蔽连续性。", "test"),
        _demo_step(4, "功能测试", "接批准的信号源/显示端，按订单要求检查音视频输出和连接稳定性。", "test"),
        _demo_step(5, "盘线扎线", "按批准圈径盘线，禁扭结、过弯或挤压插头；扎带居中。", "tie"),
        _demo_step(6, "装袋贴标装箱", "按BOM装保护帽、入袋贴标并核对批次/数量；转待检区。", "carton"),
    ]
    work_page["ie_time_study"] = {
        "title": "IE工时记录",
        "fields": list(IE_TIME_STUDY_FIELDS),
        "rows": [
            _ie_row("备料核对", "扫码：TBD"),
            _ie_row("外观检查", "检具：TBD"),
            _ie_row("导通/短路测试", "测试仪：TBD"),
            _ie_row("功能测试", "信号治具：TBD"),
            _ie_row("盘线扎线", "扎线治具：TBD"),
            _ie_row("装袋贴标装箱", "扫码：TBD"),
        ],
        "policy": {
            "measurement_basis": "待现场IE实测",
            "release_requirement": "site_measured_or_human_locked",
        },
    }
    work_page["side_sections"] = [
        {
            "title": "作业标准",
            "lines": [
                "1. 规格、标签及包装配置与受控工单/BOM一致。",
                "2. 外观按受控限度样板；电测/功能测试按已批准程序。",
                "3. 不合格品隔离、标识并提交人工评审，不自动放行。",
            ],
        },
        {
            "title": "设备/工具",
            "lines": [
                "HDMI线缆测试仪：TBD",
                "信号源/显示端或治具：TBD",
                "扫码、盘线及扎线工具：TBD",
            ],
        },
        {
            "title": "辅助材料",
            "lines": ["保护帽（如BOM要求）", "扎带、PE袋、标签、纸箱：按受控BOM"],
        },
        {
            "title": "注意事项",
            "lines": [
                "1. HDMI版本、长度、测试程序及判定阈值发布前由工程确认。",
                "2. 禁止强拉、急折和挤压插头；热插拔按测试程序。",
                "3. 本文件为草案（demo_not_for_release），无现场实测、试产或签核结论。",
            ],
        },
        {"title": "变更内容", "lines": [f"A / HDMI成品线SOP草案 / {document_date} / M2"]},
        {
            "title": "物料表",
            "lines": ["1 / TBD / HDMI成品线 / 按工单", "2 / TBD / 包装辅料 / 按受控BOM"],
        },
    ]
    work_page["bottom_sections"] = [
        {"title": "批准", "value": ""},
        {"title": "审核", "value": ""},
        {"title": "制作", "value": ""},
        {"title": "材料环保要求", "value": "材料符合适用RoHS/REACH要求；发布前确认。"},
        {"title": "管制文件（印章处）", "value": ""},
        {"title": "图号", "value": drawing_no},
    ]
    return flow_page, work_page


def _validate(path: Path, expected_date: str) -> dict[str, Any]:
    module_validation = _validate_docx_package(path)
    document = Document(path)
    tables = document.tables
    all_text = "\n".join(cell.text for table in tables for row in table.rows for cell in row.cells)
    signoff_values = [tables[7].cell(1, index).text.strip() for index in range(3)] if len(tables) >= 8 else []
    checks = {
        "sections": len(document.sections) == 2,
        "top_level_tables": len(tables) == 8,
        "portrait_flow_section": len(document.sections) >= 1 and document.sections[0].orientation == WD_ORIENT.PORTRAIT,
        "landscape_instruction_section": len(document.sections) >= 2 and document.sections[1].orientation == WD_ORIENT.LANDSCAPE,
        "has_png_media": module_validation.get("has_png_media") is True,
        "has_no_svg": module_validation.get("has_svg") is False,
        "has_no_vml": module_validation.get("has_vml_shape") is False,
        "has_no_replacement_char": module_validation.get("contains_replacement_char") is False,
        "hdmi_profile": "HDMI 成品线" in all_text,
        "draft_boundary": "demo_not_for_release" in all_text and "DRAFT" in all_text,
        "blank_signoff": signoff_values == ["", "", ""],
        "date": len(tables) >= 5 and tables[0].cell(2, 5).text.strip() == expected_date and tables[4].cell(1, 7).text.strip() == expected_date,
        "six_steps": all(token in all_text for token in ["备料核对", "外观检查", "导通/短路测试", "功能测试", "盘线扎线", "装袋贴标装箱"]),
        "unmeasured_ie": "待IE实测" in all_text and "平均观测工时(s)" in all_text,
        "unknowns_marked": all(token in all_text for token in ["TBD", "受控工单/BOM", "已批准程序"]),
    }
    errors = [name for name, passed in checks.items() if not passed]
    return {
        "profile_id": PROFILE_ID,
        "document": str(path.resolve()),
        "sha256": _sha256(path),
        "size_bytes": path.stat().st_size,
        "structural_pass": not errors,
        "checks": checks,
        "module_validation": module_validation,
        "errors": errors,
        "visual_qa": {"required": True, "expected_page_count": 2, "status": "pending_render_review"},
    }


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def generate(out_dir: Path, document_date: str) -> dict[str, Any]:
    normalized_date = date.fromisoformat(document_date).isoformat()
    display_date = normalized_date.replace("-", "/").lstrip("0").replace("/0", "/")
    out_dir.mkdir(parents=True, exist_ok=True)
    flow_page, work_page = _build_pages(normalized_date)

    flowchart_path = out_dir / FLOWCHART_NAME
    flowchart_path.write_bytes(_render_center_flowchart_shape_image(flow_page))

    document = _build_sop_word_document(flow_page, work_page)
    _apply_delivery_controls(document, normalized_date=normalized_date, display_date=display_date)
    document.sections[1].top_margin = Cm(0.1)
    document.sections[1].bottom_margin = Cm(0.1)
    for row_index, height in enumerate([400, 400, 280, 400, 400, 280]):
        _set_row_height(document.tables[5].rows[row_index], height)
    for row_index in [2, 5]:
        for column_index in [1, 2, 3]:
            for paragraph in document.tables[5].cell(row_index, column_index).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
    _set_row_height(document.tables[7].rows[0], 180)
    _set_row_height(document.tables[7].rows[1], 180)
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        runs = paragraph.runs or [paragraph.add_run("")]
        for run in runs:
            run.font.size = Pt(1)
    _set_word_cell(document.tables[5].cell(4, 4), f"变更内容\nA / HDMI成品线SOP草案 / {normalized_date} / M2", size=6, align=0)
    document_path = out_dir / DOCX_NAME
    document.save(document_path)

    format_check_path = out_dir / FORMAT_CHECK_NAME
    _write_word_format_check_json(format_check_path, flow_page, work_page)
    validation = _validate(document_path, display_date)
    validation_path = out_dir / VALIDATION_NAME
    validation_path.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest = {
        "schema_version": "1.0",
        "profile_id": PROFILE_ID,
        "status": "demo_not_for_release",
        "document_date": normalized_date,
        "product": "HDMI 成品线（规格待确认）",
        "station": "成品检验与包装",
        "shape_policy": SOP_FLOWCHART_SHAPE_POLICY,
        "layout": {"flow_section": "A4 portrait", "instruction_section": "A4 landscape", "expected_pages": 2},
        "artifacts": {
            "document_docx": DOCX_NAME,
            "center_flowchart_png": FLOWCHART_NAME,
            "format_check_json": FORMAT_CHECK_NAME,
            "validation_json": VALIDATION_NAME,
        },
        "guardrails": {
            "no_auto_release": True,
            "no_auto_signoff": True,
            "no_fabricated_ie_measurement": True,
            "no_fabricated_test_result": True,
            "unknown_spec_equipment_program_threshold_marked_tbd": True,
        },
        "validation": {"structural_pass": validation["structural_pass"], "visual_qa_required": True},
    }
    manifest_path = out_dir / MANIFEST_NAME
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    if not validation["structural_pass"]:
        raise RuntimeError("HDMI SOP structural validation failed: " + "; ".join(validation["errors"]))
    return {
        "document_docx": str(document_path.resolve()),
        "center_flowchart_png": str(flowchart_path.resolve()),
        "manifest_json": str(manifest_path.resolve()),
        "format_check_json": str(format_check_path.resolve()),
        "validation_json": str(validation_path.resolve()),
        "structural_pass": True,
        "visual_qa_required": True,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a two-page draft SOP drawing for HDMI finished-cable inspection and packaging.")
    parser.add_argument("--out-dir", type=Path, default=Path("outputs/deliverables/hdmi_cable_sop_20260811"))
    parser.add_argument("--document-date", default=date.today().isoformat())
    args = parser.parse_args()
    print(json.dumps(generate(args.out_dir, args.document_date), ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
