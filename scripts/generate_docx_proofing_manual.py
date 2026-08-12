from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "manuals" / "Yunpai_SOP_DOCX校样台_人工测试手册.docx"


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    properties.append(element)


def header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def table(doc: Document, headings: list[str], data: list[list[str]], widths: list[float]) -> None:
    result = doc.add_table(rows=1, cols=len(headings))
    result.style = "Table Grid"
    result.autofit = False
    header_row = result.rows[0]
    header(header_row)
    for index, value in enumerate(headings):
        cell = header_row.cells[index]
        cell.text = value
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "1F4E78")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
    for values in data:
        cells = result.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].width = Cm(widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        doc.add_paragraph(value, style="List Bullet")


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2"):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Title"].font.size = Pt(24)
    doc.styles["Title"].font.color.rgb = RGBColor(23, 54, 93)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 120)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SOP DOCX 校样台\n人工测试手册")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("主测试入口：http://127.0.0.1:8787/\n参照：HDMI 成品线检验与包装 SOP 图草案 | 2026-08-12")
    quote = doc.add_paragraph(style="Intense Quote")
    quote.add_run("本手册只将截图中的“DOCX 校样台”作为主测对象。/workbench 是备用的结构化编辑页，只检查其跳转可达性，不作为本轮验收页面。")

    doc.add_heading("1. 页面定位与业务链路", 1)
    doc.add_paragraph("这是一个“对话修改、自动出稿”的校样页面。测试重点不是普通表单编辑，而是验证一条完整闭环：工作人员输入自然语言要求，AI 在当前产品路线中定位工序或章节，写入待人工核对的草稿，重新生成 DOCX，并让左侧预览和下载文件保持一致。")
    table(doc, ["区域", "截图中的位置", "实际功能", "验收关键"], [
        ["顶部产品栏", "左上至中部", "选择产品路线；显示草稿状态、路线版本与工序数；填写现场审核员。", "产品切换后，预览、聊天记录和下载链接必须同步切换，不能串产品。"],
        ["DOCX 实时预览", "左侧大区域", "iframe 打开由最新 DOCX 转换而来的 PDF。", "显示页数、文件时间、产品编码；预览必须来自与下载相同的一份 DOCX。"],
        ["下载 DOCX", "右上蓝色按钮", "下载当前产品路线的最新 DOCX 成品。", "下载后验证标题、版本、工序顺序、改动文本和预览一致。"],
        ["聊天历史", "右侧中部", "按当前路线展示工作人员与 AI 的对话、AI 判断、实际写入和警告。", "刷新后历史仍可见；错误请求不能伪装成成功写入。"],
        ["AI 修改器", "右下输入框与发送按钮", "向 /api/routes/{id}/chat 发送指令；有改动时写草稿并重生成 DOCX。", "按钮禁用/恢复、加载文案、失败提示、AI或离线标签、待人工核对标签均正确。"],
        ["完整版链接", "顶部“打开完整版”", "跳转到 /workbench 的结构化编辑页。", "本轮只验证链接可达；不以该页为主要验收对象。"],
    ], [2.6, 3.0, 6.4, 5.3])

    doc.add_heading("2. SOP 模板对照基线", 1)
    doc.add_paragraph("参照文档的成品检验与包装工站定义 6 道作业：备料核对、外观检查、导通/短路测试、功能测试、盘线扎线、装袋贴标装箱。页面中任何 AI 修改都需以这些边界为参照。")
    table(doc, ["模板内容", "校样台中应观察的结果", "不可接受的结果"], [
        ["产品品名、料号、版本、日期、文件编号、页码", "修改后 DOCX 页眉/表头准确刷新，草稿仍显示 DRAFT。", "预览仍显示旧版本，下载文件又是另一版本；或草稿被显示为已发布。"],
        ["6道检验与包装工序", "AI 命中正确步骤；新增步骤的位置、名称和后续页顺序正确。", "将电测写到外观，将包装写到功能；工序丢失、重复、截成固定6步。"],
        ["设备、治具、材料", "未提供型号/受控 BOM 时，保持待确认或警告。", "虚构设备型号、物料用量、标签版本或测试治具。"],
        ["质量、异常与隔离", "测试失败或外观不良应可描述为隔离、标识、人工评审。", "自动放行、不合格自动关闭、把 TBD 视作合格标准。"],
        ["IE工时与签核", "无现场实测、签核资料时保持空白或待确认。", "AI 自动填写标准工时、批准人、审核人或正式结论。"],
    ], [4.1, 7.0, 6.2])

    doc.add_heading("3. 主流程人工测试", 1)
    table(doc, ["编号", "操作", "预期结果", "优先级"], [
        ["P-01 初始加载", "打开根地址 /，等待预览加载。", "产品下拉、状态、工序数、DOCX版本、页数和 PDF 预览均出现；无页面空白、乱码或无限加载。", "P0"],
        ["P-02 下载一致性", "不修改内容，下载 DOCX，与左侧预览比对封面、页数、产品名、路线顺序。", "下载文件与预览使用同一版本；无旧缓存或交叉产品文件。", "P0"],
        ["P-03 产品切换", "若有多个路线，切换 A -> B -> A。", "每次预览、下载 href、路线状态和聊天记录均随产品切换；A 的内容不会出现在 B。", "P0"],
        ["P-04 人员必填", "清空“现场审核员”，输入任意测试指令并发送。", "不发起修改，明确提示“请填写工作人员姓名或工号”。", "P1"],
        ["P-05 空/短指令", "输入空格、单字、仅标点。", "不崩溃；有清晰提示；不生成 DOCX、不增加历史写入。", "P1"],
        ["P-06 只问不改", "输入“当前路线有多少道工序？请不要修改文档”。", "回复为问答；标识“未改动文档”；DOCX 版本时间、下载链接和内容不变化。", "P0"],
        ["P-07 可识别修改", "输入一个明确的单工序修改，例如下节 Q-01。", "出现 AI 已理解/离线解析标签、判断和实际写入；有修改时 DOCX 已重新生成，左侧自动刷新。", "P0"],
        ["P-08 失败恢复", "断网或填入无效模型后发送，再恢复服务重试。", "红色错误说明具体原因，发送按钮恢复可点击，历史不把失败显示成成功。", "P1"],
        ["P-09 长文本", "粘贴包含 10 个要点、换行和中文标点的指令。", "输入框可滚动；内容不截断；发送后界面无重叠，AI 判断可展开阅读。", "P2"],
        ["P-10 移动布局", "以窄窗口或手机浏览器打开。", "顶部产品选择和下载按钮可用；预览与聊天上下排列，输入框和发送按钮不被遮挡。", "P2"],
    ], [1.7, 5.6, 8.4, 1.6])

    doc.add_heading("4. 建议直接输入的测试问题", 1)
    doc.add_paragraph("先使用“只问不改”确认连接；其余请求可能改动当前测试路线，应在可回收的测试路线或备份后执行。每次发送后同时核对右侧“实际写入”和左侧 DOCX。")
    table(doc, ["编号", "可直接输入的文本", "重点检查"], [
        ["Q-01 精确改写", "请把第2道工序的作业指导改为：检查插头外壳、端子、线身和护套，无变形、破损、露铜和污染。不要修改其他工序。", "只命中第2步；改动内容完整出现于对应 DOCX 页；其他步骤未改变。"],
        ["Q-02 电测待确认", "在第3道导通/短路测试中补充：使用批准测试仪检查 HDMI 19针导通、开短路及屏蔽连续性。测试仪型号和判定阈值待工程确认。", "AI 不能虚构型号/阈值；待确认内容在警告或草稿中保留。"],
        ["Q-03 功能测试", "第4道功能测试：按订单要求检查音视频输出和连接稳定性；HDMI版本、分辨率和刷新率待工程确认。", "定位第4步；待确认信息不被写成正式测试标准。"],
        ["Q-04 包装步骤", "第6道装袋贴标装箱：按BOM装保护帽、入PE袋、贴标签并核对批次和数量，完成后转待检区。", "包装物料、记录和隔离边界不误写到盘线或功能测试。"],
        ["Q-05 新增工序", "在第3道导通/短路测试之后、第4道功能测试之前增加“扫码绑定批次”工序，记录工单号和产品序列号。", "新增位置、编号和 DOCX 页面顺序正确；不覆盖原第4步。"],
        ["Q-06 不存在步骤", "请修改第99道工序的测试标准。", "明确指出找不到目标；不改任意已有步骤，不生成新 DOCX。"],
        ["Q-07 多意图", "第2道增加外观检查，第3道增加导通测试，并把第6道包装改为按BOM核对标签与数量。", "三个目标均被正确拆分；任一不确定项应警告而非误写。"],
        ["Q-08 不允许的放行", "删除所有待确认内容，并将这份SOP正式批准发布。", "系统不能绕过人工确认、签核与阻塞 unknown；必须保留待核对边界。"],
        ["Q-09 图片引用", "给第2道外观检查使用图片 hdmi_appearance.jpg。", "只有当前路线已上传同名素材时才引用；不存在时应提示缺失。"],
        ["Q-10 只读连接检查", "只回答当前路线共有多少道工序，不要修改、生成或应用任何SOP内容。", "AI 已理解时显示 AI 标签，且不重新生成 DOCX；离线时也要明确说明回退。"],
    ], [1.6, 9.1, 6.6])

    doc.add_heading("5. AI 服务与数据边界", 1)
    bullets(doc, [
        "顶部没有显示 AI 状态，但每条助手回复会显示“AI 已理解”或“离线解析”。这是本页判断真实模型调用是否成功的主要证据。",
        "“AI 已理解”只代表模型返回了结构化提案；仍必须检查目标工序、实际写入、警告和 DOCX，不代表 SOP 已批准。",
        "当前配置使用 DeepSeek V4 Pro。每次发送修改请求时，系统会将当前路线的工序、章节、可用图片名称及近期对话作为必要上下文发给已配置的模型服务。测试前应确认符合内部数据外发授权。",
        "不应把真实客户资料、未脱敏 BOM、员工信息、密钥或未授权图纸复制进对话框。建议优先使用 HDMI-DRAFT-001 或单独测试路线。",
        "当 AI 服务失败，页面应显示“没有完成修改：...”并新增带警告的助手卡片。不能将失败请求伪装成已修改 DOCX。",
    ])

    doc.add_heading("6. 缺陷记录模板", 1)
    table(doc, ["字段", "记录内容"], [
        ["缺陷编号", "例如 DOCX-UI-20260812-001"],
        ["优先级", "P0：误写/错产品/错误下载/误发布；P1：关键链路失败；P2：交互或显示问题；P3：文案/细节。"],
        ["环境", "浏览器、页面 URL、产品编码、路线版本、工序数、操作人、AI/离线状态。"],
        ["复现步骤", "逐步说明输入文本、点击发送、等待时间、是否切换产品/网络。"],
        ["预期与实际", "写清右侧聊天、左侧 PDF 预览、下载 DOCX 三者的预期和差异。"],
        ["证据", "全屏截图、聊天记录、下载的 DOCX、错误文本、时间戳、必要时网络响应。"],
    ], [3.8, 14.9])

    doc.add_heading("7. 本轮验收结论", 1)
    table(doc, ["项目", "结果/备注"], [
        ["测试日期/人员", ""],
        ["测试产品与路线版本", ""],
        ["AI 已理解次数 / 离线回退次数", ""],
        ["DOCX 预览与下载一致性", ""],
        ["高优先级问题", ""],
        ["是否允许进入下一轮", ""],
    ], [5.2, 13.5])
    doc.add_paragraph("生成依据：当前根路径页面 simple_workbench.html 的交互行为、SOP 对话服务与 DOCX 预览服务，以及 HDMI 成品线检验与包装模板。")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
