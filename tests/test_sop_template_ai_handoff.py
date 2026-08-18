from __future__ import annotations

import json
import base64
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE

from scripts.generate_sop_template_ai_handoff import (
    CENTER_FLOWCHART_NAME,
    CONTENT_PROFILE_HDMI,
    FINAL_DOCX_NAME,
    FORMAT_CHECK_NAME,
    HDMI_FINAL_DOCX_NAME,
    HDMI_TEMPLATE_ID,
    MANIFEST_NAME,
    TEMPLATE_ID,
    VALIDATION_NAME,
    generate_package,
    generate_route_package,
    validate_document,
)
from cad_ai.sop_knowledge.store import SopKnowledgeStore
from cad_ai.sop_knowledge.models import RouteSectionDraft
from tests.test_sop_knowledge_workflow import make_identity, make_route


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class SopTemplateAiHandoffTests(unittest.TestCase):
    def test_frozen_handoff_entrypoint_generates_exact_two_section_template(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            result = generate_package(directory, document_date="2026-08-11")
            root = Path(directory)

            expected_files = {
                FINAL_DOCX_NAME,
                CENTER_FLOWCHART_NAME,
                MANIFEST_NAME,
                FORMAT_CHECK_NAME,
                VALIDATION_NAME,
            }
            self.assertEqual({path.name for path in root.iterdir()}, expected_files)
            self.assertEqual(result["template_id"], TEMPLATE_ID)
            self.assertTrue(result["structural_pass"])
            self.assertTrue(result["visual_qa_required"])

            document = Document(root / FINAL_DOCX_NAME)
            self.assertEqual(len(document.sections), 2)
            self.assertEqual(len(document.tables), 8)
            self.assertEqual(document.tables[0].cell(2, 3).text.strip(), "DRAFT")
            self.assertEqual(document.tables[4].cell(2, 3).text.strip(), "DRAFT")
            self.assertEqual(document.tables[0].cell(2, 5).text.strip(), "2026/8/11")
            self.assertEqual(document.tables[4].cell(1, 7).text.strip(), "2026/8/11")
            self.assertEqual(
                [document.tables[7].cell(1, index).text.strip() for index in range(3)],
                ["", "", ""],
            )

            validation = json.loads((root / VALIDATION_NAME).read_text(encoding="utf-8"))
            self.assertTrue(validation["structural_pass"])
            self.assertEqual(validation["visual_qa"]["expected_page_count"], 2)

            manifest = json.loads((root / MANIFEST_NAME).read_text(encoding="utf-8"))
            self.assertEqual(manifest["template_id"], TEMPLATE_ID)
            self.assertTrue(manifest["tables_filled_before_flowchart"])
            self.assertEqual(manifest["fixed_template_profile"]["step_order"], "1,2,3 / 6,5,4")
            self.assertTrue(manifest["guardrails"]["no_auto_release"])

    def test_hdmi_profile_rejects_the_obsolete_two_page_route(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaisesRegex(ValueError, "route-backed multi-page"):
                generate_package(
                    directory,
                    document_date="2026-08-11",
                    content_profile=CONTENT_PROFILE_HDMI,
                )

    def test_route_backed_hdmi_generates_one_flow_page_and_repeated_instruction_pages(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            store = SopKnowledgeStore(root / "knowledge.sqlite3")
            store.initialize()
            store.ensure_process_family("test_family", "测试工艺族")
            identity = make_identity("HDMI-ROUTE-TEST")
            store.upsert_product(identity, {"class": "cable"})
            route_id = store.create_route(make_route(identity, 3))
            store.create_route_section(
                route_id,
                RouteSectionDraft(
                    section_type="ie_timing",
                    content={"单价": "15.50", "人数": "2", "source": "人工填写"},
                ),
            )
            step = store.get_route(route_id)["steps"][1]
            store.update_step_field(
                step["id"],
                "record_output",
                ["首件记录", "巡检记录", "登记工单号和异常现象"],
                reviewer="worker-01",
                decision="needs_revision",
            )
            result = generate_route_package(
                root / "package",
                document_date="2026-08-12",
                db_path=store.path,
                route_id=route_id,
            )

            self.assertEqual(result["template_id"], HDMI_TEMPLATE_ID)
            self.assertEqual(result["instruction_page_count"], 3)
            self.assertEqual(result["expected_page_count"], 4)
            document = Document(result["document_docx"])
            self.assertEqual(len(document.sections), 2)
            self.assertEqual(len(document.tables), 16)
            flow_ie_time = document.tables[2]
            self.assertIn("单价", flow_ie_time.cell(1, 3).text)
            self.assertIn("人数", flow_ie_time.cell(1, 4).text)
            self.assertIn("15.50", flow_ie_time.cell(2, 3).text)
            self.assertIn("2", flow_ie_time.cell(2, 4).text)
            for page_index in range(3):
                base = 4 + page_index * 4
                self.assertEqual(document.tables[base].cell(2, 3).text.strip(), "DRAFT")
                work_ie_time = document.tables[base + 2]
                self.assertIn("单价", work_ie_time.cell(1, 3).text)
                self.assertIn("人数", work_ie_time.cell(1, 4).text)
                self.assertIn("15.50", work_ie_time.cell(2, 3).text)
                self.assertIn("2", work_ie_time.cell(2, 4).text)
                self.assertEqual(
                    [document.tables[base + 3].cell(1, index).text.strip() for index in range(3)],
                    ["", "", ""],
                )
            second_instruction_body = document.tables[9]
            self.assertIn("记录要求（最新）", second_instruction_body.cell(4, 7).text)
            self.assertIn("登记工单号和异常现象", second_instruction_body.cell(4, 7).text)

    def test_route_backed_hdmi_supports_every_work_image_layout(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            store = SopKnowledgeStore(root / "knowledge.sqlite3")
            store.initialize()
            store.ensure_process_family("test_family", "测试工艺族")
            identity = make_identity("HDMI-LAYOUT-TEST")
            store.upsert_product(identity, {"class": "cable"})
            route_id = store.create_route(make_route(identity, 6))
            steps = store.get_route(route_id)["steps"]
            store.update_step_field(
                steps[0]["id"],
                "method",
                ["准备物料", "核对方向", "执行作业", "记录结果"],
                reviewer="layout-tester",
            )
            for slot_count, step in enumerate(steps, start=1):
                store.set_step_work_image_slots(step["id"], slot_count, reviewer="layout-tester")

            result = generate_route_package(
                root / "package",
                document_date="2026-08-18",
                db_path=store.path,
                route_id=route_id,
            )

            document = Document(result["document_docx"])
            expected_positions = {
                1: [(0, 1, "1")],
                2: [(0, 1, "1"), (0, 4, "2")],
                3: [(0, 1, "1"), (0, 3, "2"), (0, 5, "3")],
                4: [(0, 1, "1"), (0, 4, "2"), (3, 4, "3"), (3, 1, "4")],
                5: [(0, 1, "1"), (0, 3, "2"), (0, 5, "3"), (3, 4, "4"), (3, 1, "5")],
                6: [(0, 1, "1"), (0, 3, "2"), (0, 5, "3"), (3, 5, "4"), (3, 3, "5"), (3, 1, "6")],
            }
            for page_index, slot_count in enumerate(range(1, 7)):
                base = 4 + page_index * 4
                body = document.tables[base + 1]
                ie_table = document.tables[base + 2]
                self.assertEqual(len(ie_table.rows), 2 + slot_count)
                self.assertEqual(ie_table.rows[0].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
                self.assertEqual(ie_table.rows[1].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
                self.assertGreaterEqual(ie_table.rows[0].height.twips, 360)
                self.assertGreaterEqual(ie_table.rows[1].height.twips, 400)
                for row, column, prefix in expected_positions[slot_count]:
                    self.assertTrue(body.cell(row, column).text.strip().startswith(prefix))
            first_body_text = "\n".join(
                cell.text for row in document.tables[5].rows for cell in row.cells
            )
            self.assertGreaterEqual(document.tables[5].rows[5].height.twips, 1300)
            for method in ("准备物料", "核对方向", "执行作业", "记录结果"):
                self.assertIn(method, first_body_text)

            manifest = json.loads((root / "package" / MANIFEST_NAME).read_text(encoding="utf-8"))
            self.assertEqual(
                [item["work_image_slots"] for item in manifest["layout"]["instruction_layouts"]],
                [1, 2, 3, 4, 5, 6],
            )
            validation = json.loads((root / "package" / VALIDATION_NAME).read_text(encoding="utf-8"))
            self.assertTrue(validation["checks"]["ie_action_rows_match_work_image_slots"])
            self.assertTrue(validation["checks"]["visual_step_order_every_page"])

    def test_route_backed_hdmi_embeds_only_confirmed_step_media(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            store = SopKnowledgeStore(root / "knowledge.sqlite3")
            store.initialize()
            store.ensure_process_family("test_family", "测试工艺族")
            identity = make_identity("HDMI-MEDIA-TEST")
            store.upsert_product(identity, {"class": "cable"})
            route_id = store.create_route(make_route(identity, 2))
            steps = store.get_route(route_id)["steps"]
            confirmed = store.upload_media_asset(
                route_id, original_name="confirmed.png", mime_type="image/png",
                data=PNG_1X1, uploaded_by="worker-01",
            )
            draft = store.upload_media_asset(
                route_id, original_name="draft.png", mime_type="image/png",
                data=PNG_1X1 + b"draft", uploaded_by="worker-01",
            )
            store.link_media_asset(steps[0]["id"], confirmed["id"], caption="确认图片")
            store.confirm_step(steps[0]["id"], reviewer="worker-01")
            store.link_media_asset(steps[1]["id"], draft["id"], caption="草稿图片")

            result = generate_route_package(
                root / "package", document_date="2026-08-12", db_path=store.path, route_id=route_id,
            )

            with zipfile.ZipFile(result["document_docx"]) as archive:
                embedded = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertEqual(len(embedded), 2)  # 流程图 + 1 张已确认工序图片
            document = Document(result["document_docx"])
            first_page_body = document.tables[5]
            second_page_body = document.tables[9]
            self.assertNotIn("待人工上传确认", first_page_body.cell(0, 1).text)
            self.assertIn("待人工上传确认", second_page_body.cell(0, 1).text)

    def test_check_only_validation_rejects_missing_document(self) -> None:
        result = validate_document(Path("does-not-exist.docx"))
        self.assertFalse(result["structural_pass"])
        self.assertEqual(result["errors"], ["document_not_found"])


if __name__ == "__main__":
    unittest.main()
