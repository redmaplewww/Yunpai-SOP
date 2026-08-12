from __future__ import annotations

import csv
import json
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook

from cad_ai.sop_visual_template import (
    REFERENCE_80806_129_FORMAT,
    SOP_FLOWCHART_SHAPE_POLICY,
    build_process_flow_page,
    build_sop_table_templates,
    build_usb_cable_packaging_demo,
    build_work_instruction_page,
    classify_sop_flow_node_shape,
    render_process_flow_svg,
    render_process_flow_table_markdown,
    render_work_instruction_table_markdown,
    render_work_instruction_svg,
    write_demo_sop_excel_package,
    write_demo_sop_word_package,
    write_demo_sop_table_package,
    write_demo_sop_package,
    write_sop_excel_template_package,
    write_sop_word_template_package,
    write_sop_table_template_package,
    write_sop_template_package,
)


class SopVisualTemplateTests(unittest.TestCase):
    def test_reference_profile_matches_80806_129_pdf_format(self) -> None:
        profile = REFERENCE_80806_129_FORMAT

        self.assertEqual(profile["page_count"], 42)
        self.assertEqual(profile["flow_page"]["page_number"], 1)
        self.assertEqual(profile["flow_page"]["orientation"], "portrait")
        self.assertEqual(profile["work_instruction_page"]["orientation"], "landscape")
        self.assertEqual(profile["work_instruction_page"]["image_slots"], 6)
        self.assertEqual(
            profile["work_instruction_page"]["right_sections"],
            ["作业标准", "设备/工具", "辅助材料", "注意事项", "变更内容", "物料表"],
        )
        self.assertEqual(
            profile["work_instruction_page"]["bottom_sections"],
            ["批准", "审核", "制作", "材料环保要求", "管制文件（印章处）", "图号"],
        )

    def test_process_flow_page_template_has_reference_header_and_flow_nodes(self) -> None:
        page = build_process_flow_page(
            product_name="USB3.1 TYPE C-C",
            part_no="80806-129",
            document_no="SOP-SD-4378",
            drawing_no="A-US22-0000-03",
            operations=[f"工序{i}" for i in range(1, 42)],
        )

        self.assertEqual(page["page_type"], "process_flow")
        self.assertEqual(page["orientation"], "portrait")
        self.assertEqual(page["title"], "流程图")
        self.assertIn("品名", page["header_fields"])
        self.assertIn("料号", page["header_fields"])
        self.assertIn("文件编号", page["header_fields"])
        self.assertEqual(len(page["flow_nodes"]), 41)
        self.assertEqual(page["flow_nodes"][0]["label"], "工序1")
        self.assertEqual(page["flow_nodes"][-1]["label"], "工序41")
        self.assertEqual(page["ie_time_study"]["title"], "IE工时记录")
        self.assertIn("机器型号", page["ie_time_study"]["fields"])
        self.assertIn("动态调整", page["ie_time_study"]["fields"])

    def test_work_instruction_page_template_has_six_blank_image_slots_and_side_sections(self) -> None:
        page = build_work_instruction_page(
            product_name="USB3.1 TYPE C-C",
            part_no="80806-129",
            station="裁线",
            document_no="SOP-SD-4378",
            drawing_no="A-US22-0000-03",
        )

        self.assertEqual(page["page_type"], "work_instruction")
        self.assertEqual(page["orientation"], "landscape")
        self.assertEqual(page["title"], "标准作业指导书")
        self.assertEqual(len(page["step_slots"]), 6)
        self.assertTrue(all(slot["image_placeholder"] for slot in page["step_slots"]))
        self.assertEqual([slot["slot_no"] for slot in page["step_slots"]], [1, 2, 3, 4, 5, 6])
        self.assertEqual(page["side_sections"][0]["title"], "作业标准")
        self.assertEqual(page["side_sections"][-1]["title"], "物料表")
        self.assertIn("批准", [item["title"] for item in page["bottom_sections"]])
        self.assertIn("管制文件（印章处）", [item["title"] for item in page["bottom_sections"]])

    def test_svg_renderers_emit_valid_reference_shaped_svg(self) -> None:
        flow = build_process_flow_page(operations=["裁线", "自动机前处理", "分线一"])
        work = build_work_instruction_page(station="裁线")

        flow_svg = render_process_flow_svg(flow)
        work_svg = render_work_instruction_svg(work)

        flow_root = ET.fromstring(flow_svg)
        work_root = ET.fromstring(work_svg)
        self.assertEqual(flow_root.attrib["viewBox"], "0 0 595.32 841.92")
        self.assertEqual(work_root.attrib["viewBox"], "0 0 841.92 595.32")
        self.assertIn("流程图", "".join(flow_root.itertext()))
        self.assertIn("标准作业指导书", "".join(work_root.itertext()))
        self.assertEqual("".join(work_root.itertext()).count("图片占位"), 6)

    def test_template_package_writes_svg_manifest_and_format_check_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            out_dir = Path(directory)
            paths = write_sop_template_package(out_dir)

            for key in ["process_flow_svg", "work_instruction_svg", "manifest_json", "format_check_csv"]:
                self.assertTrue(paths[key].exists(), key)
                self.assertGreater(paths[key].stat().st_size, 100, key)

            with paths["format_check_csv"].open(encoding="utf-8-sig") as handle:
                rows = list(csv.DictReader(handle))
            checks = {row["check_item"]: row for row in rows}
            self.assertEqual(checks["page_1_orientation"]["generated"], "portrait")
            self.assertEqual(checks["work_instruction_image_slots"]["generated"], "6")
            self.assertEqual(checks["work_instruction_side_sections"]["status"], "match")

    def test_demo_package_fills_text_and_images_for_simple_usb_cable_product(self) -> None:
        demo = build_usb_cable_packaging_demo()

        self.assertEqual(demo["product_name"], "USB-C 简易数据线")
        self.assertEqual(len(demo["operations"]), 6)
        self.assertEqual(len(demo["work_instruction"]["step_slots"]), 6)
        self.assertTrue(all(not slot["image_placeholder"] for slot in demo["work_instruction"]["step_slots"]))
        self.assertTrue(all(slot.get("visual") for slot in demo["work_instruction"]["step_slots"]))
        self.assertIn("扎线", demo["operations"])
        self.assertIn("装PE袋", demo["operations"])

        work_svg = render_work_instruction_svg(demo["work_instruction"])
        self.assertIn("USB-C 简易数据线", work_svg)
        self.assertIn("整理线材", work_svg)
        self.assertIn("generated-step-visual", work_svg)
        self.assertEqual(work_svg.count("generated-step-visual"), 6)
        self.assertNotIn("图片占位", work_svg)

        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_package(directory)
            for key in ["process_flow_svg", "work_instruction_svg", "manifest_json", "format_check_csv"]:
                self.assertTrue(paths[key].exists(), key)
                self.assertGreater(paths[key].stat().st_size, 100, key)
            manifest = json.loads(paths["manifest_json"].read_text(encoding="utf-8"))
            self.assertEqual(manifest["demo_product"]["product_name"], "USB-C 简易数据线")
            self.assertEqual(manifest["demo_product"]["image_source"], "local_generated_svg")

    def test_table_templates_match_reference_without_svg_outputs(self) -> None:
        tables = build_sop_table_templates()

        self.assertEqual(tables["format"], "table_only")
        self.assertIn("流程图首页-表头", tables["process_flow"])
        self.assertIn("流程图首页-工序流程", tables["process_flow"])
        self.assertIn("标准作业指导书-表头", tables["work_instruction"])
        self.assertIn("标准作业指导书-六步说明", tables["work_instruction"])
        self.assertIn("标准作业指导书-右侧栏", tables["work_instruction"])
        self.assertIn("标准作业指导书-底部签核", tables["work_instruction"])

        flow_md = render_process_flow_table_markdown(tables["process_flow"])
        work_md = render_work_instruction_table_markdown(tables["work_instruction"])

        self.assertIn("| 品名 |", flow_md)
        self.assertIn("| 工序序号 | 工序名称 | 节点类型 | 上道工序 | 下道工序 |", flow_md)
        self.assertIn("| 步骤序号 | 图片内容 | 文字说明 |", work_md)
        self.assertIn("| 作业标准 |", work_md)
        self.assertNotIn("<svg", flow_md)
        self.assertNotIn("<svg", work_md)

        with tempfile.TemporaryDirectory() as directory:
            paths = write_sop_table_template_package(directory)
            names = {path.name for path in paths.values()}
            self.assertIn("sop_process_flow_tables.md", names)
            self.assertIn("sop_work_instruction_tables.md", names)
            self.assertIn("sop_process_flow_header.csv", names)
            self.assertIn("sop_process_flow_routing.csv", names)
            self.assertIn("sop_work_instruction_header.csv", names)
            self.assertIn("sop_work_instruction_steps.csv", names)
            self.assertIn("sop_work_instruction_side_sections.csv", names)
            self.assertIn("sop_work_instruction_signoff.csv", names)
            self.assertIn("sop_table_format_check.csv", names)
            self.assertTrue(all(path.suffix in {".md", ".csv", ".json"} for path in paths.values()))
            self.assertFalse(any(path.suffix == ".svg" for path in paths.values()))

    def test_demo_table_package_outputs_simple_product_tables_only(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_table_package(directory)

            self.assertTrue(all(path.suffix in {".md", ".csv", ".json"} for path in paths.values()))
            self.assertFalse(any(path.suffix == ".svg" for path in paths.values()))
            work_md = paths["work_instruction_md"].read_text(encoding="utf-8")
            self.assertIn("USB-C 简易数据线", work_md)
            self.assertIn("| 1 | 备料示意 | 备料：确认数据线、扎带、PE袋、标签齐套。 |", work_md)
            self.assertIn("| 4 | 扎线示意 | 扎线：使用扎带固定线圈，扎带位置居中。 |", work_md)

            with paths["work_instruction_steps_csv"].open(encoding="utf-8-sig") as handle:
                step_rows = list(csv.DictReader(handle))
            self.assertEqual(step_rows[0]["图片内容"], "备料示意")
            self.assertEqual(step_rows[0]["文字说明"], "备料：确认数据线、扎带、PE袋、标签齐套。")
            self.assertEqual(step_rows[3]["图片内容"], "扎线示意")

            with paths["work_instruction_signoff_csv"].open(encoding="utf-8-sig") as handle:
                signoff_rows = list(csv.DictReader(handle))
            self.assertEqual(signoff_rows[0]["批准"], "")
            self.assertEqual(signoff_rows[0]["审核"], "")
            self.assertEqual(signoff_rows[0]["制作"], "")

    def test_excel_template_recreates_reference_layout_with_merged_tables_not_svg_or_csv(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_sop_excel_template_package(directory)

            self.assertEqual(paths["workbook_xlsx"].suffix, ".xlsx")
            self.assertTrue(all(path.suffix in {".xlsx", ".json"} for path in paths.values()))
            self.assertFalse(any(path.suffix in {".csv", ".svg"} for path in paths.values()))

            workbook = load_workbook(paths["workbook_xlsx"])
            self.assertEqual(workbook.sheetnames, ["01_流程图", "02_标准作业指导书"])

            flow_sheet = workbook["01_流程图"]
            self.assertEqual(flow_sheet.page_setup.orientation, "portrait")
            self.assertEqual(flow_sheet["H2"].value, "流程图")
            self.assertEqual(flow_sheet["B5"].value, "品名")
            self.assertEqual(flow_sheet["O5"].value, "文件编号")
            self.assertIn("B10:Q54", _merged_ranges(flow_sheet))
            self.assertIsNone(flow_sheet["B10"].value)
            self.assertNotIn("工序1", _sheet_values(flow_sheet))
            self.assertNotIn("裁线", _sheet_values(flow_sheet))
            self.assertGreaterEqual(len(flow_sheet.merged_cells.ranges), 12)

            work_sheet = workbook["02_标准作业指导书"]
            self.assertEqual(work_sheet.page_setup.orientation, "landscape")
            self.assertEqual(work_sheet["O2"].value, "标准作业指导书")
            self.assertEqual(work_sheet["A1"].value, "产品品名")
            self.assertEqual(work_sheet["AA6"].value, "作业标准")
            self.assertEqual(work_sheet["AA31"].value, "物料表")
            self.assertEqual(work_sheet["A32"].value, "批准")
            self.assertEqual(work_sheet["AA32"].value, "管制文件（印章处）")
            self.assertEqual(work_sheet["AA34"].value, "图号")
            self.assertIn("图片流程描述及说明", _sheet_values(work_sheet))
            self.assertIn("图片占位", _sheet_values(work_sheet))
            self.assertIn("1", str(work_sheet["C6"].value))
            self.assertIn("6", str(work_sheet["C19"].value))
            self.assertGreaterEqual(len(work_sheet.merged_cells.ranges), 35)

    def test_demo_excel_package_fills_same_visual_layout_for_usb_cable(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_excel_package(directory)

            self.assertEqual(paths["workbook_xlsx"].suffix, ".xlsx")
            self.assertFalse(any(path.suffix in {".csv", ".svg"} for path in paths.values()))

            workbook = load_workbook(paths["workbook_xlsx"])
            flow_sheet = workbook["01_流程图"]
            work_sheet = workbook["02_标准作业指导书"]
            values = _sheet_values(work_sheet)

            self.assertIn("USB-C 简易数据线", _sheet_values(flow_sheet))
            self.assertIn("扎线装袋", values)
            self.assertIn("备料", values)
            self.assertIn("扎线：使用扎带固定线圈，扎带位置居中。", values)
            self.assertEqual(work_sheet["A32"].value, "批准")
            self.assertEqual(work_sheet["A34"].value, None)
            self.assertEqual(work_sheet["M34"].value, None)

    def test_word_template_recreates_sop_with_native_tables_not_excel(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_sop_word_template_package(directory)

            self.assertEqual(paths["document_docx"].suffix, ".docx")
            self.assertTrue(all(path.suffix in {".docx", ".json"} for path in paths.values()))
            self.assertFalse(any(path.suffix in {".xlsx", ".csv", ".svg"} for path in paths.values()))

            document = Document(paths["document_docx"])
            self.assertEqual(len(document.sections), 2)
            self.assertEqual(document.sections[0].orientation, WD_ORIENT.PORTRAIT)
            self.assertEqual(document.sections[1].orientation, WD_ORIENT.LANDSCAPE)
            self.assertEqual(len(document.tables), 8)

            flow_header, flow_body, flow_ie_time, flow_footer = document.tables[:4]
            flow_text = _table_text(flow_header)
            self.assertIn("流程图", flow_text)
            self.assertIn("品名", flow_text)
            self.assertIn("页数", flow_text)
            self.assertIn("文件编号", flow_text)
            self.assertEqual(len(flow_body.rows), 1)
            self.assertEqual(_unique_cell_count(flow_body.rows[0]), 1)
            self.assertEqual(_table_text(flow_body), "")
            flow_ie_text = _table_text(flow_ie_time)
            self.assertIn("IE工时记录", flow_ie_text)
            self.assertIn("机器型号", flow_ie_text)
            self.assertIn("标准工时(s)", flow_ie_text)
            self.assertIn("动态调整", flow_ie_text)
            self.assertIn("待IE实测", flow_ie_text)
            self.assertIn("EF-42013-23", _table_text(flow_footer))
            self.assertNotIn("工序1", "\n".join(_table_text(table) for table in document.tables[:4] if table is not flow_ie_time))
            self.assertNotIn("裁线", "\n".join(_table_text(table) for table in document.tables[:4] if table is not flow_ie_time))

            work_header, work_body, work_ie_time, work_footer = document.tables[4:8]
            work_text = "\n".join(_table_text(table) for table in [work_header, work_body, work_ie_time, work_footer])
            self.assertIn("标准作业指导书", work_text)
            self.assertIn("产品品名", _table_text(work_header))
            self.assertIn("作业顺序", _table_text(work_header))
            self.assertIn("图片流程描述及说明", work_text)
            self.assertIn("图片占位", work_text)
            work_ie_text = _table_text(work_ie_time)
            self.assertIn("IE工时记录", work_ie_text)
            self.assertIn("动作", work_ie_text)
            self.assertIn("机器型号", work_ie_text)
            self.assertIn("宽放率", work_ie_text)
            self.assertIn("标准工时(s)", work_ie_text)
            self.assertIn("动态调整", work_ie_text)
            self.assertIn("作业标准", work_text)
            self.assertIn("物料表", work_text)
            self.assertIn("管制文件（印章处）", work_text)
            self.assertIn("批准", work_text)

    def test_demo_word_package_fills_work_instruction_without_signoff_names(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_word_package(directory)

            self.assertEqual(paths["document_docx"].suffix, ".docx")
            self.assertFalse(any(path.suffix in {".xlsx", ".csv", ".svg"} for path in paths.values()))

            document = Document(paths["document_docx"])
            combined_text = "\n".join(_table_text(table) for table in document.tables)
            self.assertIn("USB-C 简易数据线", combined_text)
            self.assertIn("扎线装袋", combined_text)
            self.assertIn("备料：确认数据线、扎带、PE袋、标签齐套。", combined_text)
            self.assertIn("扎线：使用扎带固定线圈，扎带位置居中。", combined_text)
            self.assertNotIn("唐华伟", combined_text)


    def test_demo_word_package_can_embed_pdf_reference_center_flowchart(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_word_package(directory, center_flowchart=True)

            document = Document(paths["document_docx"])
            self.assertEqual(len(document.tables), 8)
            flow_body = document.tables[1]
            self.assertEqual(len(flow_body.rows), 1)
            self.assertEqual(_unique_cell_count(flow_body.rows[0]), 1)
            self.assertEqual(len(flow_body.cell(0, 0).tables), 1)

            center_chart = flow_body.cell(0, 0).tables[0]
            center_text = _table_text(center_chart)
            process_prefix = "\u5de5\u5e8f"
            self.assertIn(f"{process_prefix}1", center_text)
            self.assertIn(f"{process_prefix}6", center_text)
            self.assertIn("\u2192", center_text)
            self.assertTrue("\u2191" in center_text or "\u2193" in center_text)

            with zipfile.ZipFile(paths["document_docx"]) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")
            self.assertNotIn("<svg", document_xml.lower())
            self.assertNotIn("<v:shape", document_xml.lower())

            check = json.loads(paths["format_check_json"].read_text(encoding="utf-8"))
            checks = {item["check_item"]: item for item in check["checks"]}
            self.assertEqual(checks["center_flowchart_style"]["generated"], "pdf_reference_compact")

    def test_demo_word_package_can_embed_reference_shape_flowchart_image(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = write_demo_sop_word_package(
                directory,
                center_flowchart=True,
                center_flowchart_style="pdf_reference_shape_blocks",
            )

            document = Document(paths["document_docx"])
            self.assertEqual(len(document.tables), 8)
            flow_body = document.tables[1]
            self.assertEqual(len(flow_body.rows), 1)
            self.assertEqual(_unique_cell_count(flow_body.rows[0]), 1)
            self.assertEqual(len(flow_body.cell(0, 0).tables), 0)

            with zipfile.ZipFile(paths["document_docx"]) as package:
                names = set(package.namelist())
                document_xml = package.read("word/document.xml").decode("utf-8")
            self.assertTrue(any(name.startswith("word/media/") and name.endswith(".png") for name in names))
            self.assertIn("pic:pic", document_xml)
            self.assertNotIn("<svg", document_xml.lower())
            self.assertNotIn("<v:shape", document_xml.lower())

            check = json.loads(paths["format_check_json"].read_text(encoding="utf-8"))
            checks = {item["check_item"]: item for item in check["checks"]}
            self.assertEqual(checks["center_flowchart_style"]["generated"], "pdf_reference_shape_blocks")
            self.assertEqual(checks["center_flowchart_shape_policy"]["generated"], "testing=diamond; processing=ellipse")

    def test_reference_shape_policy_classifies_testing_as_diamond_and_processing_as_ellipse(self) -> None:
        self.assertEqual(SOP_FLOWCHART_SHAPE_POLICY, "testing=diamond; processing=ellipse")
        self.assertEqual(classify_sop_flow_node_shape({"name": "PCBA上电测试"}), "diamond")
        self.assertEqual(classify_sop_flow_node_shape({"name": "CCD检查焊点"}), "diamond")
        self.assertEqual(classify_sop_flow_node_shape({"name": "ICT Test"}), "diamond")
        self.assertEqual(classify_sop_flow_node_shape({"type": "measurement", "name": "量测外露尺寸"}), "diamond")
        self.assertEqual(classify_sop_flow_node_shape({"name": "组装下盖"}), "ellipse")
        self.assertEqual(classify_sop_flow_node_shape({"name": "包装入库"}), "ellipse")
        self.assertEqual(classify_sop_flow_node_shape({"name": "返工处理", "note": "功能测试不合格品维修"}), "ellipse")
 
    def test_batch_closing_samples_fill_tables_before_center_flowchart_png(self) -> None:
        from scripts.generate_sop_batch_samples import generate_sop_batch_samples

        with tempfile.TemporaryDirectory() as directory:
            result = generate_sop_batch_samples(directory, sample_count=4)

            manifest_path = Path(result["manifest_json"])
            self.assertTrue(manifest_path.exists())
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["status"], "demo_not_for_release")
            self.assertEqual(
                manifest["generation_sequence"],
                [
                    "parse_requirement",
                    "fill_word_tables",
                    "build_structured_flowchart",
                    "render_center_flowchart_png",
                    "insert_png_into_process_flow_body_cell",
                    "validate_docx",
                ],
            )
            self.assertEqual(len(result["samples"]), 4)

            for sample in result["samples"]:
                sample_manifest = json.loads(Path(sample["manifest_json"]).read_text(encoding="utf-8"))
                self.assertEqual(sample_manifest["status"], "demo_not_for_release")
                self.assertEqual(sample_manifest["center_flowchart_target"], "process_flow_body_table_cell_0_0")
                self.assertTrue(sample_manifest["tables_filled_before_flowchart"])

                docx_path = Path(sample["document_docx"])
                preview_path = Path(sample["preview_png"])
                check_path = Path(sample["format_check_json"])
                self.assertTrue(docx_path.exists(), sample["slug"])
                self.assertTrue(preview_path.exists(), sample["slug"])
                self.assertTrue(check_path.exists(), sample["slug"])

                document = Document(docx_path)
                self.assertEqual(len(document.sections), 2)
                self.assertEqual(document.sections[0].orientation, WD_ORIENT.PORTRAIT)
                self.assertEqual(document.sections[1].orientation, WD_ORIENT.LANDSCAPE)
                self.assertEqual(len(document.tables), 8)

                flow_header, flow_body, flow_ie_time, _flow_footer = document.tables[:4]
                work_header, work_body, work_ie_time, work_footer = document.tables[4:8]
                self.assertIn(sample_manifest["product_name"], _table_text(flow_header))
                self.assertIn(sample_manifest["product_name"], _table_text(work_header))
                self.assertIn(sample_manifest["station"], _table_text(work_header))
                self.assertGreater(len(_table_text(work_body)), 40)
                self.assertEqual(len(flow_body.rows), 1)
                self.assertEqual(_unique_cell_count(flow_body.rows[0]), 1)
                self.assertEqual(len(flow_body.cell(0, 0).tables), 0)
                self.assertIn("IE工时记录", _table_text(flow_ie_time))
                self.assertIn("IE工时记录", _table_text(work_ie_time))
                for expected_ie_field in ["动作", "机器型号", "IE测量方法", "平均观测工时(s)", "标准工时(s)", "动态调整"]:
                    self.assertIn(expected_ie_field, _table_text(flow_ie_time))
                    self.assertIn(expected_ie_field, _table_text(work_ie_time))
                self.assertIn("demo_not_for_release", _table_text(flow_ie_time))
                self.assertIn("demo_not_for_release", _table_text(work_ie_time))
                self.assertNotIn("IE宸ユ椂", _table_text(flow_ie_time))
                self.assertNotIn("IE宸ユ椂", _table_text(work_ie_time))
                self.assertEqual(work_footer.cell(1, 0).text.strip(), "")
                self.assertEqual(work_footer.cell(1, 1).text.strip(), "")
                self.assertEqual(work_footer.cell(1, 2).text.strip(), "")

                with zipfile.ZipFile(docx_path) as package:
                    names = set(package.namelist())
                    document_xml = package.read("word/document.xml").decode("utf-8")
                self.assertTrue(any(name.startswith("word/media/") and name.endswith(".png") for name in names))
                self.assertIn("pic:pic", document_xml)
                self.assertNotIn("<svg", document_xml.lower())
                self.assertNotIn("<v:shape", document_xml.lower())

                format_check = json.loads(check_path.read_text(encoding="utf-8"))
                checks = {item["check_item"]: item for item in format_check["checks"]}
                self.assertEqual(checks["center_flowchart_style"]["generated"], "pdf_reference_shape_blocks")
                self.assertEqual(checks["center_flowchart_shape_policy"]["generated"], SOP_FLOWCHART_SHAPE_POLICY)
                for node in sample_manifest["center_flowchart"]["nodes"]:
                    self.assertEqual(node["shape"], classify_sop_flow_node_shape(node))


def _sheet_values(sheet: object) -> str:
    values: list[str] = []
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value is not None:
                values.append(str(cell.value))
    return "\n".join(values)


def _merged_ranges(sheet: object) -> set[str]:
    return {str(cell_range) for cell_range in sheet.merged_cells.ranges}


def _table_text(table: object) -> str:
    values: list[str] = []
    seen: set[object] = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            text = cell.text.strip()
            if text:
                values.append(text)
    return "\n".join(values)


def _non_empty_cell_count(row: object) -> int:
    seen: set[object] = set()
    count = 0
    for cell in row.cells:
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        if cell.text.strip():
            count += 1
    return count


def _unique_cell_count(row: object) -> int:
    return len({cell._tc for cell in row.cells})


if __name__ == "__main__":
    unittest.main()
