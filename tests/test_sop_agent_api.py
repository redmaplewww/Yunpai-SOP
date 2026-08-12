from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from cad_ai.sop_agent import SOP_GENERATION_SEQUENCE, SopGenerateRequest, SopRoutingStep, generate_sop_package
from cad_ai.sop_agent import _ollama_native_chat_payload, _ollama_native_chat_url
from cad_ai.sop_api import create_sop_fastapi_app


class SopAgentApiTests(unittest.TestCase):
    def test_ollama_native_payload_disables_thinking_for_qwen35b_json(self) -> None:
        url = _ollama_native_chat_url("http://127.0.0.1:11434/v1/chat/completions")
        payload = _ollama_native_chat_payload(
            {
                "model": "qwen3.6:35b",
                "messages": [{"role": "user", "content": "Return JSON."}],
                "max_tokens": 1900,
                "temperature": 0.1,
            }
        )

        self.assertEqual(url, "http://127.0.0.1:11434/api/chat")
        self.assertEqual(payload["model"], "qwen3.6:35b")
        self.assertFalse(payload["think"])
        self.assertEqual(payload["format"], "json")
        self.assertEqual(payload["options"]["num_predict"], 1900)

    def test_generate_sop_package_from_structured_request(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            request = _sample_request(Path(directory))

            response = generate_sop_package(request)

            self.assertEqual(response.status, "demo_not_for_release")
            self.assertEqual(response.generation_sequence, SOP_GENERATION_SEQUENCE)
            self.assertTrue(response.tables_filled_before_flowchart)
            self.assertEqual(response.center_flowchart_target, "process_flow_body_table_cell_0_0")
            self.assertTrue(response.artifacts.document_docx.exists())
            self.assertTrue(response.artifacts.center_flowchart_png.exists())
            self.assertTrue(response.artifacts.manifest_json.exists())
            self.assertEqual(response.validation["sections"], 2)
            self.assertEqual(response.validation["top_level_tables"], 8)
            self.assertTrue(response.validation["has_png_media"])
            self.assertFalse(response.validation["has_svg"])
            self.assertFalse(response.validation["has_vml_shape"])
            self.assertFalse(response.validation["contains_replacement_char"])
            self.assertTrue(response.validation["contains_ie_time_title"])
            self.assertTrue(response.validation["contains_machine_model_field"])

            document = Document(response.artifacts.document_docx)
            self.assertEqual(len(document.tables), 8)
            document_text = _document_table_text(document)
            self.assertNotIn("\ufffd", document_text)
            self.assertIn("车载摄像头模组", document_text)

    def test_fastapi_generate_and_download_artifacts(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except ImportError as exc:  # pragma: no cover - depends on local test extras.
            self.skipTest(f"FastAPI test client unavailable: {exc}")

        with tempfile.TemporaryDirectory() as directory:
            app = create_sop_fastapi_app(default_out_dir=Path(directory))
            client = TestClient(app)

            health = client.get("/api/sop/health")
            self.assertEqual(health.status_code, 200)
            self.assertEqual(health.json()["draft_status"], "demo_not_for_release")

            payload = _sample_request(Path(directory)).model_dump(mode="json")
            response = client.post("/api/sop/generate", json=payload)
            self.assertEqual(response.status_code, 200, response.text)
            body = response.json()
            self.assertEqual(body["run_id"], "sop_api_test")
            self.assertEqual(body["validation"]["top_level_tables"], 8)
            self.assertFalse(body["validation"]["contains_replacement_char"])

            manifest = client.get("/api/sop/runs/sop_api_test")
            self.assertEqual(manifest.status_code, 200)
            self.assertEqual(manifest.json()["manifest"]["run_id"], "sop_api_test")

            docx = client.get("/api/sop/runs/sop_api_test/artifacts/document_docx")
            self.assertEqual(docx.status_code, 200)
            self.assertGreater(len(docx.content), 1000)
            self.assertIn("application/vnd.openxmlformats-officedocument", docx.headers["content-type"])

            png = client.get("/api/sop/runs/sop_api_test/artifacts/center_flowchart_png")
            self.assertEqual(png.status_code, 200)
            self.assertEqual(png.headers["content-type"], "image/png")
            self.assertGreater(len(png.content), 1000)


def _sample_request(out_dir: Path) -> SopGenerateRequest:
    return SopGenerateRequest(
        product_name="车载摄像头模组",
        part_no="CAM-API-DEMO-001",
        document_no="SOP-CAM-API-001",
        drawing_no="DWG-CAM-API-001",
        station="无尘装配及EOL测试",
        requirement_text="来料核对 -> PCBA装入前壳 -> 镜头组件装配 -> 点胶固化 -> 气密检查 -> EOL测试 -> 合格判定 -> 外观清洁贴标包装",
        routing_steps=[
            SopRoutingStep(name="来料核对", type="inspection", visual_type="inspection", machine_model="IQC-BENCH-DEMO", average_observed_time_s=8),
            SopRoutingStep(name="PCBA装入前壳", type="process", visual_type="assembly", machine_model="ESD-BENCH-DEMO", average_observed_time_s=12),
            SopRoutingStep(name="镜头组件装配", type="process", visual_type="assembly", machine_model="ASSY-JIG-DEMO", average_observed_time_s=14),
            SopRoutingStep(name="点胶固化", type="process", visual_type="process", machine_model="DISPENSE-UV-DEMO", average_observed_time_s=20),
            SopRoutingStep(name="EOL测试", type="test", visual_type="test", machine_model="EOL-TEST-DEMO", average_observed_time_s=18),
            SopRoutingStep(name="外观清洁贴标包装", type="process", visual_type="pack", machine_model="PACK-BENCH-DEMO", average_observed_time_s=10),
        ],
        machine_hints=["IQC-BENCH-DEMO", "ESD-BENCH-DEMO", "ASSY-JIG-DEMO", "DISPENSE-UV-DEMO", "EOL-TEST-DEMO", "PACK-BENCH-DEMO"],
        run_id="sop_api_test",
        out_dir=out_dir,
    )


def _document_table_text(document: Document) -> str:
    values: list[str] = []
    seen: set[object] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen:
                    continue
                seen.add(cell._tc)
                if cell.text.strip():
                    values.append(cell.text.strip())
    return "\n".join(values)


if __name__ == "__main__":
    unittest.main()
